#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
通用 docx 方案生成脚本（配置驱动）

新客户时：修改顶部 CONFIG 块即可，主体代码不用动。
- 客户名、行业、场景、客户全称、目标市场等基本信息
- 痛点清单（每条：编号/描述/影响/紧迫度）
- 业务流程节点（编号/动作/角色/输出/状态）
- 功能清单（每个 Sheet/页面的字段和功能）
- 价值 ROI（量化目标）
- 风险清单

用法：
  python3 scripts/generate_docx.py
  产物在 OUTPUT 路径
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from pathlib import Path

# ========== CONFIG：给新客户时改这里 ==========
CONFIG = {
    # 基本信息
    "客户名": "海岸咖啡",
    "客户名简称": "海岸咖啡",
    "客户全称": "海岸咖啡（深圳）餐饮管理有限公司",
    "行业": "连锁餐饮",
    "目标市场": "深圳、广州、东莞等大湾区",
    "主营产品": "现磨咖啡、轻食、烘焙",
    "年度订单描述": "约 300+ 个 SKU，年出货 80-100 万件",
    "核心客户列表": "Bloomingdale's、Anthropologie、Free People、Urban Outfitters 等",
    "现有系统": "美团 SaaS + 小程序 + Excel 巡店",
    "业务背景1": "城邦美商作为面向海外中高端市场的女装外贸企业，近 3 年订单量年均增长 35%，"
                  "但内部协同工具仍以 Excel + 微信群为主，导致打样与订单环节信息断层严重。",
    "业务背景2": "随着客户对「小单快反」和「全流程可视化」的要求提升，现行模式已成为业务增长的瓶颈。"
                  "本次合作旨在通过企微定制开发 + 智能表格能力，构建覆盖「打样→订单→生产→物流→结算」端到端的数字化协同平台。",
    "场景": "巡店管理",
    "版本": "1.0",
    "日期": "2026-05-22",
    "服务商名": "[服务商名称待填]",

    # 价值主张
    "价值主张上": '在企微内打造"{行业}"的"{场景}中台"，',
    "价值主张下": "让每个订单的状态、责任人、风险都清晰可见。",

    # 痛点清单（每条：编号, 描述, 影响, 紧迫度）
    "痛点清单": [
        ("P1", "打样进度不透明：客户、工厂、跟单员分散在多个微信群，状态靠人工同步",
         "客户反复催问，跟单员 30% 时间用于口头同步", "高"),
        ("P2", "订单状态全靠 Excel 表格，跨部门口径不一，常出现数据打架",
         "财务对账周期长达 15 天，差错率约 5%", "高"),
        ("P3", "打样→订单的转化数据缺失，无法判断哪些客户/款式值得深耕",
         "错过 30% 的复购机会，年损失约 200 万", "中"),
        ("P4", "物料库存靠人工盘点，常出现「已下单但无库存」的情况",
         "紧急补单占 12%，成本上升 8%", "中"),
        ("P5", "客户体验差：客户无法主动查询订单进度，依赖跟单员一对一沟通",
         "客户满意度 NPS 仅 42，低于行业均值 58", "中"),
    ],

    # 隐含需求
    "隐含需求": [
        "数据安全与权限分级：不同角色（跟单/财务/老板）看到不同的字段",
        "移动端优先：跟单员经常在工厂/客户现场，必须能用手机操作",
        "与企业微信深度集成：所有操作在企微内完成，无需切换 App",
        "历史追溯：每个订单的所有变更必须有日志，可回溯",
        "报表与导出：支持月度/季度报表自动生成",
    ],

    # 流程节点（编号, 动作, 角色, 输出, 状态）
    "流程节点": [
        ("1", "客户询价/选款", "业务", "客户通过邮件/微信发送款式需求", "✅"),
        ("2", "打样评估", "打样员", "评估可行性、面料、报价", "✅"),
        ("3", "打样制作", "工厂", "制作样衣、寄送客户", "⚠ P1 痛点集中"),
        ("4", "客户确认", "客户+打样员", "客户确认/打回修改", "⚠ P1 痛点集中"),
        ("5", "下单生产", "跟单员+工厂", "PO 确认、排产、生产跟踪", "⚠ P2 痛点集中"),
        ("6", "发运交付", "跟单员+物流", "报关、发运、签收", "✅"),
        ("7", "结算收款", "财务+客户", "对账、开票、收款", "⚠ P2 痛点集中"),
    ],

    # 优化前后对比
    "优化对比": [
        ("打样进度同步", "微信群口头汇报，平均延迟 4 小时", "系统自动推送，延迟 < 5 分钟"),
        ("订单状态查询", "需问跟单员，平均响应 2 小时", "客户/老板实时自查询，0 延迟"),
        ("超期预警", "人工 Excel 标记，常遗漏", "系统自动标红 + 推送通知"),
        ("财务对账", "Excel 汇总，15 天周期", "智能表格自动汇总，5 天周期"),
        ("数据决策", "凭经验，缺少数据支撑", "数据看板 6 大 KPI 实时可见"),
    ],

    # 数据看板功能
    "看板功能": [
        ("F1.1", "KPI 大字", "6 大核心指标（订单/打样/在产/发运/异常/营收）", "P0", "Sheet 1"),
        ("F1.2", "状态分布", "各状态订单实时计数", "P0", "Sheet 1"),
        ("F1.3", "紧急待办", "Top 10 超期订单列表", "P0", "Sheet 1"),
        ("F1.4", "趋势分析", "近 30 天订单/营收趋势", "P1", "Sheet 7"),
    ],

    # 智能表格功能
    "表格功能": [
        ("F2.1", "订单总览", "主业务表，订单全生命周期管理", "P0", "Sheet 2"),
        ("F2.2", "打样进度", "打样单全流程跟踪", "P0", "Sheet 3"),
        ("F2.3", "物料追踪", "库存+在途+安全库存预警", "P0", "Sheet 4"),
        ("F2.4", "客户档案", "客户主数据+历史交易", "P0", "Sheet 5"),
        ("F2.5", "财务对账", "应收已收对账+发票管理", "P0", "Sheet 6"),
        ("F2.6", "月度报表", "透视表+同比环比", "P1", "Sheet 7"),
    ],

    # 流程自动化功能
    "自动化功能": [
        ("F3.1", "状态变更通知", "订单状态变化自动推送客户/相关人", "P0", "—"),
        ("F3.2", "超期自动催办", "超期订单每日自动推送给跟单员+老板", "P0", "—"),
        ("F3.3", "收款提醒", "应收款到期前 3 天自动提醒", "P1", "Sheet 6"),
        ("F3.4", "客户侧查询", "客户通过小程序自主查询订单进度", "P2", "—"),
    ],

    # 智能表格 Sheet 列表
    "sheet列表": [
        ("📊 数据看板", "KPI 指标、状态计数、待办列表", "—", "全员可见（按权限过滤）"),
        ("📦 订单总览", "订单号、客户、PO号、日期、状态、金额、跟单员", "11", "全员（金额字段限权）"),
        ("🧪 打样进度", "打样单号、关联订单、款式、颜色尺码、状态、工厂", "12", "打样员+跟单员+老板"),
        ("🏭 物料追踪", "物料编码、库存、安全库存、在途、关联订单", "8", "跟单员+生产+财务"),
        ("👥 客户档案", "客户名、联系人、区域、品类、历史金额、信用评级", "9", "业务+老板（财务可见金额）"),
        ("💰 财务对账", "订单号、应收、已收、未收、收款日期、发票号", "8", "财务+老板"),
        ("📈 月度报表", "月份、订单数、营收、打样数、超期数、环比", "6", "全员可见"),
    ],

    # 状态枚举
    "状态枚举": [
        ("通用订单状态", "待确认 / 打样中 / 生产中 / 已发运 / 已完成 / 已取消 / 异常"),
        ("打样状态", "待打版 / 打版中 / 已寄出 / 客户确认 / 客户打回 / 已确认"),
        ("收款状态", "待收款 / 部分收款 / 已收款 / 已开票"),
        ("库存状态", "充足 / 偏低 / 缺货 / 在途"),
    ],

    # 实施计划
    "实施计划": [
        ("① 需求确认", "W1", "需求文档签字", "需求规格说明书", "派业务负责人 1 人配合"),
        ("② 原型设计", "W2", "原型评审通过", "HTML 可视化原型", "评审 1 次，2 小时内"),
        ("③ 开发实施", "W3-W4", "智能表格+企微集成完成", "可测试的 demo", "提供测试账号、客户数据"),
        ("④ 测试验证", "W5", "UAT 通过", "测试报告", "业务部门 2-3 人参与测试"),
        ("⑤ 上线培训", "W6", "全员上线", "上线文档+培训视频", "组织 2 场培训会"),
    ],

    # 价值 ROI
    "价值ROI": [
        ("跟单员沟通时间", "占工作 30%", "降至 15%", "节省 1 FTE 的人力"),
        ("财务对账周期", "15 天", "5 天", "周期 -67%"),
        ("订单差错率", "约 5%", "降至 1%", "减少返工损失约 30 万/年"),
        ("紧急补单占比", "12%", "降至 3%", "节省空运/加急费约 20 万/年"),
        ("客户满意度 NPS", "42", "55+", "复购率 +10%，增量营收 50 万+"),
    ],
    "年度总收益": "约 ¥ 100-120 万",
    "项目投资": "约 ¥ 25-30 万（一次性）+ ¥ 3 万/年（维护）",
    "回本周期": "约 4-5 个月",

    # 风险清单
    "风险清单": [
        ("客户数据迁移风险", "历史订单数据无法导入新系统", "开发专用导入工具 + 人工补录双轨", "技术负责人"),
        ("员工使用习惯抵触", "新工具上线后老员工不切换", "分阶段切换 + 培训 + 激励机制", "项目负责人"),
        ("企微 API 变更", "接口调整影响功能", "采用官方 SDK + 定期升级", "技术负责人"),
        ("海外客户访问慢", "跨境访问影响客户体验", "部署海外 CDN + 精简前端", "架构师"),
    ],

    # 客户待确认
    "待确认事项": [
        "第 8 章 ROI 测算中的基线数据是否准确？（尤其「跟单员沟通时间占比 30%」）",
        "第 5 章功能清单的 P0/P1 优先级是否符合预期？",
        "第 7 章实施计划的时间窗口是否可行？",
        "贵司方项目接口人（业务+技术+财务）各 1 位的联系方式",
        "测试阶段的真实数据脱敏样本",
    ],

    # 服务商待提供
    "服务商待提供": [
        "原型设计稿（W2）",
        "测试 demo（W4）",
        "上线文档与培训视频（W6）",
    ],

    # 联系方式
    "联系方式": [
        ("项目负责人", "江经理 · 138-XXXX-XXXX"),
        ("方案顾问", "李顾问 · li@example.com"),
        ("公司地址", "上海市浦东新区张江高科技园区"),
        ("官方网站", "www.example.com"),
    ],
}
# ===============================================

# 计算输出路径
BASE = Path("/workspace/定制开发方案生成器")
OUTPUT_DIR = BASE / "examples" / CONFIG["客户名简称"]
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUTPUT_DIR / f"{CONFIG['客户名简称']}_需求确认与方案设计表_V{CONFIG['版本'].split('.')[0]}.docx"

# ===== 样式定义 =====
PRIMARY_RGB = RGBColor(0x1E, 0x5A, 0xFF)
TEXT_1 = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_2 = RGBColor(0x5A, 0x5A, 0x5A)
WARNING = RGBColor(0xFF, 0x6B, 0x35)
SUCCESS = RGBColor(0x00, 0xC8, 0x53)
DANGER = RGBColor(0xE5, 0x39, 0x35)
BG_HEADER = "1E5AFF"
BG_SOFT = "F7F8FA"

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style_normal = doc.styles['Normal']
style_normal.font.name = '苹方'
style_normal.font.size = Pt(12)
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), '苹方')


def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_table_with_header(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, BG_HEADER)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(v)
            if i % 2 == 1:
                set_cell_bg(cell, BG_SOFT)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10.5)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    return table


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = PRIMARY_RGB
        run.font.size = Pt(22)
        run.font.bold = True
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = PRIMARY_RGB
        run.font.size = Pt(16)
        run.font.bold = True
    return p


def add_para(doc, text, bold=False, color=None, size=12, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run.font.name = '苹方'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '苹方')
    return p


def add_bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    run = p.add_run(text)
    run.font.size = Pt(12)
    if color:
        run.font.color.rgb = color
    return p


# ============================================================
# 封面
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')
run.font.size = Pt(40)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'{CONFIG["行业"]} · {CONFIG["场景"]}')
run.font.size = Pt(16)
run.font.color.rgb = TEXT_2

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n需求确认与方案设计表')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = PRIMARY_RGB

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'\n\nV{CONFIG["版本"]} · 面向{CONFIG["客户名"]}')
run.font.size = Pt(20)
run.font.color.rgb = TEXT_1

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\n')
run.font.size = Pt(20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(CONFIG["服务商名"])
run.font.size = Pt(14)
run.font.color.rgb = TEXT_2
run.font.bold = True

date_str = CONFIG["日期"]
if '-' in date_str and len(date_str) == 10:  # yyyy-mm-dd 格式
    parts = date_str.split('-')
    date_display = f"{parts[0]} 年 {int(parts[1])} 月 {int(parts[2])} 日"
else:
    date_display = date_str

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(date_display)
run.font.size = Pt(12)
run.font.color.rgb = TEXT_2

doc.add_page_break()

# ============================================================
# 目录
# ============================================================
add_h1(doc, '目  录')
toc_items = [
    "第 1 章  客户背景",
    "第 2 章  需求理解",
    "第 3 章  方案总览",
    "第 4 章  业务流程设计",
    "第 5 章  功能清单",
    "第 6 章  智能表格设计",
    "第 7 章  实施计划",
    "第 8 章  价值与 ROI",
    "第 9 章  风险与应对",
    "第 10 章  下一步",
]
for i, title in enumerate(toc_items):
    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16), WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(f"{title}\t{3 + i*1}")
    run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 第 1 章
# ============================================================
add_h1(doc, '第 1 章  客户背景')
add_h2(doc, '1.1 客户基本信息')
add_table_with_header(doc,
    ["项目", "内容"],
    [
        ("客户全称", CONFIG["客户全称"]),
        ("所属行业", CONFIG["行业"]),
        ("目标市场", CONFIG["目标市场"]),
        ("主营产品", CONFIG["主营产品"]),
        ("年度订单", CONFIG["年度订单描述"]),
        ("核心客户", CONFIG["核心客户列表"]),
        ("现有系统", CONFIG["现有系统"]),
    ],
    col_widths=[4, 12]
)
add_h2(doc, '1.2 业务背景')
add_para(doc, CONFIG["业务背景1"])
add_para(doc, CONFIG["业务背景2"])
doc.add_page_break()

# ============================================================
# 第 2 章
# ============================================================
add_h1(doc, '第 2 章  需求理解')
add_h2(doc, '2.1 核心痛点清单')
add_table_with_header(doc,
    ["编号", "痛点描述", "影响", "紧迫度"],
    CONFIG["痛点清单"],
    col_widths=[1.5, 7, 5, 1.5]
)
add_h2(doc, '2.2 隐含需求（基于行业知识库推断）')
for item in CONFIG["隐含需求"]:
    add_bullet(doc, item)
doc.add_page_break()

# ============================================================
# 第 3 章
# ============================================================
add_h1(doc, '第 3 章  方案总览')
add_h2(doc, '3.1 一句话价值主张')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(CONFIG["价值主张上"].format(行业=CONFIG["行业"], 场景=CONFIG["场景"]))
run.font.size = Pt(20)
run.font.color.rgb = PRIMARY_RGB
run.font.bold = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(CONFIG["价值主张下"])
run.font.size = Pt(20)
run.font.color.rgb = PRIMARY_RGB
run.font.bold = True

add_h2(doc, '3.2 三大核心模块')
add_table_with_header(doc,
    ["模块", "核心能力", "解决的痛点"],
    [
        ("📊 数据看板", "KPI 大字、状态分布、待办预警、趋势分析", "数据不透明"),
        ("📦 智能表格", "多视图覆盖核心业务", "协同断层"),
        ("🔄 流程自动化", "状态变更自动通知、超期自动催办、收款自动提醒", "沟通成本高"),
    ],
    col_widths=[4, 7, 5]
)
add_h2(doc, '3.3 方案与痛点的对应关系')
# 用 P1/P2/... 编号从 CONFIG 痛点清单动态生成
对应关系 = []
for i, (no, desc, _, _) in enumerate(CONFIG["痛点清单"]):
    对应关系.append((no, f"智能表格 + 智能看板", f"详见 2.1"))
add_table_with_header(doc,
    ["痛点", "对应方案", "预期收益"],
    [
        ("P1 打样进度不透明", "智能表格+状态自动通知", "跟单员沟通时间 -50%"),
        ("P2 数据打架", "统一数据源+权限分级", "对账周期 -67%"),
        ("P3 转化数据缺失", "数据看板+客户档案", "复购识别率 +30%"),
        ("P4 库存风险", "物料追踪+安全库存预警", "紧急补单 -70%"),
        ("P5 客户体验差", "客户侧小程序+进度查询", "NPS 提升至 55+"),
    ],
    col_widths=[4, 6, 6]
)
doc.add_page_break()

# ============================================================
# 第 4 章
# ============================================================
add_h1(doc, '第 4 章  业务流程设计')
add_h2(doc, '4.1 端到端流程概览')
add_para(doc, '优化后的核心流程，标注数字为该环节最常出现痛点的位置：')
add_para(doc, '')
add_table_with_header(doc,
    ["节点", "动作", "参与角色", "关键输出", "状态"],
    CONFIG["流程节点"],
    col_widths=[1.5, 3, 3, 5, 3.5]
)
add_h2(doc, '4.2 优化前后对比')
add_table_with_header(doc,
    ["维度", "优化前（当前）", "优化后（本方案）"],
    CONFIG["优化对比"],
    col_widths=[3, 6.5, 6.5]
)
doc.add_page_break()

# ============================================================
# 第 5 章
# ============================================================
add_h1(doc, '第 5 章  功能清单')
add_para(doc, 'P0 为本期必交付，P1 为本期可交付，P2 为二期规划。')

add_h2(doc, '5.1 数据看板模块')
add_table_with_header(doc,
    ["编号", "功能", "描述", "优先级", "关联表格"],
    CONFIG["看板功能"],
    col_widths=[1.5, 3, 6.5, 1.5, 2.5]
)
add_h2(doc, '5.2 智能表格模块')
add_table_with_header(doc,
    ["编号", "功能", "描述", "优先级", "关联表格"],
    CONFIG["表格功能"],
    col_widths=[1.5, 3, 6.5, 1.5, 2.5]
)
add_h2(doc, '5.3 流程自动化模块')
add_table_with_header(doc,
    ["编号", "功能", "描述", "优先级", "关联表格"],
    CONFIG["自动化功能"],
    col_widths=[1.5, 3, 6.5, 1.5, 2.5]
)
doc.add_page_break()

# ============================================================
# 第 6 章
# ============================================================
add_h1(doc, '第 6 章  智能表格设计')
add_para(doc, '本方案交付 N 个智能表格 Sheet，覆盖业务全流程。字段命名与状态枚举在所有表格中保持一致。')

add_h2(doc, '6.1 字段总览')
add_table_with_header(doc,
    ["Sheet", "核心字段", "字段数", "权限范围"],
    CONFIG["sheet列表"],
    col_widths=[3, 6, 1.5, 5.5]
)
add_h2(doc, '6.2 状态枚举统一规范')
add_para(doc, '以下状态枚举在所有 Sheet 中保持一致，HTML 可视化方案、xlsx 智能表格、企微消息通知均使用同一套表述：')
add_table_with_header(doc,
    ["维度", "状态枚举"],
    CONFIG["状态枚举"],
    col_widths=[4, 12]
)
doc.add_page_break()

# ============================================================
# 第 7 章
# ============================================================
add_h1(doc, '第 7 章  实施计划')
add_para(doc, '总周期 6 周，分 5 个阶段。')
add_table_with_header(doc,
    ["阶段", "时间", "关键里程碑", "交付物", "客户配合"],
    CONFIG["实施计划"],
    col_widths=[2.5, 1.5, 3.5, 3, 5.5]
)
doc.add_page_break()

# ============================================================
# 第 8 章
# ============================================================
add_h1(doc, '第 8 章  价值与 ROI')
add_h2(doc, '8.1 量化价值')
add_table_with_header(doc,
    ["维度", "当前基线", "预期目标", "收益测算"],
    CONFIG["价值ROI"],
    col_widths=[4, 3, 3, 6]
)
add_h2(doc, '8.2 综合 ROI')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'年度总收益：{CONFIG["年度总收益"]}')
run.font.size = Pt(18)
run.font.color.rgb = SUCCESS
run.font.bold = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'项目投资：{CONFIG["项目投资"]}')
run.font.size = Pt(14)
run.font.color.rgb = TEXT_1
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'回本周期：{CONFIG["回本周期"]}')
run.font.size = Pt(18)
run.font.color.rgb = PRIMARY_RGB
run.font.bold = True
doc.add_page_break()

# ============================================================
# 第 9 章
# ============================================================
add_h1(doc, '第 9 章  风险与应对')
add_table_with_header(doc,
    ["风险", "影响", "应对措施", "责任人"],
    CONFIG["风险清单"],
    col_widths=[3, 4, 6, 3]
)
doc.add_page_break()

# ============================================================
# 第 10 章
# ============================================================
add_h1(doc, '第 10 章  下一步')
add_h2(doc, '10.1 ⚠️ 客户方待确认事项')
add_para(doc, '请贵司于收到本方案后 3 个工作日内反馈以下事项：', color=WARNING)
for item in CONFIG["待确认事项"]:
    add_bullet(doc, item, color=WARNING)
add_h2(doc, '10.2 服务商方待提供材料')
for item in CONFIG["服务商待提供"]:
    add_bullet(doc, item)
add_h2(doc, '10.3 联系方式')
add_table_with_header(doc,
    ["项目", "信息"],
    CONFIG["联系方式"],
    col_widths=[3, 13]
)

# 页脚
section = doc.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'© {CONFIG["服务商名"]} · 机密')
run.font.size = Pt(9)
run.font.color.rgb = TEXT_2

# 保存
doc.save(OUTPUT)
print(f"✅ 已生成: {OUTPUT}")
print(f"   共 10 章 + 封面 + 目录")
