#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
通用 docx 方案生成脚本（配置驱动）

新客户时：修改顶部 CONFIG 块即可，主体代码不用动。
- 客户名、行业、场景、客户全称、目标市场等基本信息
- 痛点清单（每条：编号/描述/影响/紧迫度）
- 业务流程节点（编号/动作/角色/输出/状态）
- 功能清单（每个 Sheet/页面的字段和功能）
- 价值 ROI（量化目标）
- 风险清单

用法：
  python3 scripts/generate_docx.py
  产物在 OUTPUT 路径
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from pathlib import Path

# ========== CONFIG：给新客户时改这里 ==========
CONFIG = {
    # 基本信息
    "客户名": "亿选装饰",
    "客户名简称": "亿选装饰",
    "客户全称": "贵州亿选装饰工程有限公司",
    "行业": "家装",
    "目标市场": "贵州省内中高端家装客户",
    "主营产品": "家装设计 + 施工 + 主材 + 软装的一站式服务",
    "年度订单描述": "在施工地 120+ 个，年签约客户 200+ 户",
    "核心客户列表": "贵州省内中高端家装业主（120m²-300m² 户型为主）",
    "现有系统": "Excel 客户登记表 + 微信群沟通 + 部分设计师手稿",
    "业务背景1": "贵州亿选装饰作为贵州省内中高端家装公司，过去 3 年签约客户数从 80 户增长到 200+ 户，"
                  "但前端获客数据、客户画像、转化漏斗仍依赖 Excel 人工统计，"
                  "老板无法实时看清经营全貌。"
                  "数据分散在业务员、设计师、前台、客户经理的多个 Excel 和微信群里，"
                  "想要做精细化运营时常常拿不到准确数据。",
    "业务背景2": "随着家装行业竞争加剧、客户对透明化要求提升，亿选装饰需要一个「数据驱动的运营驾驶舱」，"
                  "在企微内一站式呈现客户来源、转化漏斗、工地状态、经营指标，"
                  "让老板、设计师、客户经理都能基于数据决策。"
                  "本次合作的核心是「前端数据采集 + 后端智能表格 + 数据看板」。",
    "场景": "前端数据分析",
    "版本": "1.0",
    "日期": "2026-05-22",
    "服务商名": "杭州装企云网络科技有限公司",

    # 价值主张
    "价值主张上": '在企微内打造"{行业}"的"{场景}中台"，',
    "价值主张下": "让每个订单的状态、责任人、风险都清晰可见。",

    # 痛点清单（每条：编号, 描述, 影响, 紧迫度）
    "痛点清单": [
        ("P1", "客户来源数据散落：业务员各自记录 Excel，无法看清哪条渠道 ROI 最高",
         "线上投放预算浪费约 30%，年损失约 50 万", "高"),
        ("P2", "客户画像缺失：业主信息（户型/预算/消费偏好）靠设计师手记，无法精准营销",
         "复购率仅 8%，远低于行业头部 25%", "高"),
        ("P3", "转化漏斗不清晰：咨询→量房→报价→签约各环节流失率无法量化",
         "不知道哪个环节最该优化，错失改进机会", "中"),
        ("P4", "工地进度不透明：老板想看 120+ 个工地状态，只能逐个问项目经理",
         "管理半径受限，工地异常平均发现延迟 3 天", "中"),
        ("P5", "数据决策靠经验：老板凭感觉拍板，缺少数据支撑",
         "新业务方向试错成本高，年损失约 80 万", "中"),
    ],

    # 隐含需求
    "隐含需求": [
        "数据安全与权限分级：老板/管理层/设计师/客户经理看到不同维度的数据",
        "移动端优先：老板/管理层经常出差，必须能用手机看实时看板",
        "与企业微信深度集成：所有数据看板在企微工作台直接打开",
        "历史追溯：每个数据点有来源、有时间、可回溯到原始登记表",
        "多维分析：支持按时间/区域/户型/客户类型等维度灵活钻取",
    ],

    # 流程节点（编号, 动作, 角色, 输出, 状态）
    "流程节点": [
        ("1", "客户线索登记", "前台+客户经理", "业主信息录入智能表格", "⚠ P1 痛点集中"),
        ("2", "客户来源打标", "客户经理", "标注渠道/活动/转介绍人", "⚠ P1 痛点集中"),
        ("3", "量房转化", "设计师", "量房数据上传、看方案", "⚠ P3 痛点集中"),
        ("4", "报价签约", "客户经理+设计师", "客户接受报价、签订合同", "⚠ P3 痛点集中"),
        ("5", "施工跟进", "项目经理+施工队", "工地进度、阶段验收", "⚠ P4 痛点集中"),
        ("6", "竣工交付", "项目经理+监理", "客户验收、结算尾款", "✅"),
        ("7", "复购转介", "售后+客户经理", "客户回访、复购/转介激励", "⚠ P2 痛点集中"),
    ],

    # 优化前后对比
    "优化对比": [
        ("客户来源分析", "Excel 统计，月底才知道", "数据看板实时刷新，老板随时看"),
        ("转化漏斗", "靠经验估算各环节流失", "漏斗图精确到每个节点的转化率"),
        ("工地状态", "逐个问项目经理", "工地看板集中呈现 120+ 工地状态"),
        ("客户画像", "设计师手记，散落在个人", "客户档案全员可见，可按维度筛选"),
        ("数据决策", "凭经验拍板", "数据看板 8 大 KPI 实时可见"),
    ],

    # 数据看板功能
    "看板功能": [
        ("F1.1", "KPI 大字", "8 大核心指标（在施工地/签约/复购率/客单价/异常等）", "P0", "Sheet 1"),
        ("F1.2", "客户来源分布", "按渠道/活动的客户来源占比", "P0", "Sheet 1"),
        ("F1.3", "转化漏斗", "咨询→量房→报价→签约各环节转化率", "P0", "Sheet 1"),
        ("F1.4", "工地状态", "120+ 在建工地的阶段分布", "P0", "Sheet 1"),
        ("F1.5", "趋势分析", "近 30 天/12 月的客户量、签约、营收趋势", "P0", "Sheet 7"),
        ("F1.6", "维度钻取", "支持时间/区域/户型/客户类型维度切换", "P1", "Sheet 5"),
    ],

    # 智能表格功能
    "表格功能": [
        ("F2.1", "客户档案", "业主主数据（姓名/电话/户型/预算/来源/状态）", "P0", "Sheet 4"),
        ("F2.2", "数据源管理", "各业务系统接入状态、字段映射表", "P0", "Sheet 2"),
        ("F2.3", "维度配置", "分析维度的字典（时间/区域/户型/类型）", "P0", "Sheet 3"),
        ("F2.4", "多维分析", "按维度切分的指标矩阵", "P0", "Sheet 5"),
        ("F2.5", "工地管理", "在建工地的状态、阶段、工期", "P0", "Sheet 6"),
        ("F2.6", "趋势分析", "日/周/月维度的客户、签约、营收趋势", "P0", "Sheet 7"),
        ("F2.7", "异常预警", "数据/工地/客户异常清单", "P0", "Sheet 8"),
    ],

    # 流程自动化功能
    "自动化功能": [
        ("F3.1", "每日数据自动同步", "每晚自动从业务系统拉数据到智能表格", "P0", "—"),
        ("F3.2", "异常自动预警", "工地超期/客户流失/数据异常自动推送给老板", "P0", "—"),
        ("F3.3", "周报自动生成", "每周一自动生成上周经营周报", "P1", "—"),
        ("F3.4", "客户生日/节日提醒", "在客户生日、装修周年推送关怀消息", "P2", "—"),
    ],

    # 智能表格 Sheet 列表
    "sheet列表": [
        ("📊 数据看板", "8 大核心 KPI、客户来源、转化漏斗、工地状态", "—", "老板+管理层全员"),
        ("📥 数据源管理", "各业务系统数据接入状态、字段映射表", "—", "IT 管理员"),
        ("📐 维度配置", "分析维度（时间/区域/户型/客户类型）的字典", "8", "IT+管理层"),
        ("👥 客户档案", "业主姓名、电话、户型、面积、预算、来源、状态", "15", "客户经理+设计师+老板"),
        ("📊 多维分析", "按维度切分的指标（漏斗/复购/客单价/转化率）", "12", "老板+管理层"),
        ("🏗 工地管理", "工地编号、状态、阶段、工期、异常", "10", "项目经理+监理+老板"),
        ("📈 趋势分析", "日/周/月维度的客户量、签约量、营收趋势", "8", "老板+管理层"),
        ("⚠️ 异常预警", "数据异常/工地异常/客户流失预警清单", "9", "老板+客户经理"),
    ],

    # 状态枚举
    "状态枚举": [
        ("客户状态", "咨询 / 量房 / 报价 / 签约 / 施工 / 竣工 / 售后 / 流失"),
        ("工地状态", "待开工 / 施工中 / 隐蔽工程 / 泥木 / 油漆 / 软装 / 竣工 / 售后"),
        ("合同状态", "草稿 / 待签 / 已签 / 执行中 / 已完成 / 已取消"),
        ("数据接入状态", "未接入 / 接入中 / 已接入 / 异常"),
    ],

    # 实施计划
    "实施计划": [
        ("① 需求确认", "W1", "需求文档签字", "需求规格说明书", "派业务负责人 1 人配合"),
        ("② 原型设计", "W2", "原型评审通过", "HTML 可视化原型", "评审 1 次，2 小时内"),
        ("③ 开发实施", "W3-W4", "智能表格+企微集成完成", "可测试的 demo", "提供测试账号、客户数据"),
        ("④ 测试验证", "W5", "UAT 通过", "测试报告", "业务部门 2-3 人参与测试"),
        ("⑤ 上线培训", "W6", "全员上线", "上线文档+培训视频", "组织 2 场培训会"),
    ],

    # 价值 ROI
    "价值ROI": [
        ("线上投放 ROI", "1:2.5", "提升至 1:4", "年节省投放成本 30 万"),
        ("客户转化率", "12%", "提升至 18%", "年新增签约 24 单，+360 万营收"),
        ("复购率", "8%", "提升至 18%", "年复购 18 户，+270 万营收"),
        ("工地异常发现", "延迟 3 天", "实时预警", "减少返工损失 25 万/年"),
        ("管理者决策效率", "周会 4 小时", "看板 5 分钟", "管理层时间节省 80%"),
    ],
    "年度总收益": "约 ¥ 600-800 万（增量为 360+270，节省约 55）",
    "项目投资": "约 ¥ 18-25 万（一次性）+ ¥ 2 万/年（维护）",
    "回本周期": "约 1-2 个月（按增量营收算）",

    # 风险清单
    "风险清单": [
        ("历史数据迁移风险", "Excel 数据格式不一、缺失字段多", "开发专用导入工具 + 人工补录 + 数据治理培训", "数据负责人"),
        ("员工录入习惯抵触", "客户经理不愿每天录数据", "录入极简化（表单/扫码）+ 激励 + 督导", "项目负责人"),
        ("数据准确性质疑", "管理层不信新数据", "新旧数据并行 1 个月 + 抽样对比", "项目负责人"),
        ("看板使用率低", "上线后没人主动看", "周会强制使用 + 老板每日截图", "老板亲自推"),
    ],

    # 客户待确认
    "待确认事项": [
        "第 8 章 ROI 测算中的基线数据（尤其「线上投放 ROI 1:2.5」「复购率 8%」）是否准确？",
        "第 5 章功能清单的 8 个数据看板模块是否符合预期？",
        "第 6 章 7 个 Sheet 的字段是否覆盖业务全流程？",
        "数据源接入：哪些业务系统需要对接？优先级？",
        "历史数据范围：多少个月的 Excel 客户登记表需要迁移？",
    ],

    # 服务商待提供
    "服务商待提供": [
        "数据看板原型设计稿（W2）",
        "智能表格 demo + 数据看板 demo（W4）",
        "数据接入方案（W3）+ 上线文档（W6）",
        "管理层使用培训（W6 上午）+ 一线录入培训（W6 下午）",
    ],

    # 联系方式
    "联系方式": [
        ("项目负责人", "吴 sir · 138-XXXX-XXXX"),
        ("方案顾问", "杭州装企云方案团队 · contact@zhuangqiyun.com"),
        ("公司地址", "杭州市余杭区未来科技城"),
        ("官方网站", "www.zhuangqiyun.com"),
    ],
}
# ===============================================

# 计算输出路径
BASE = Path("/workspace/定制开发方案生成器")
OUTPUT_DIR = BASE / "examples" / CONFIG["客户名简称"]
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUTPUT_DIR / f"{CONFIG['客户名简称']}_需求确认与方案设计表_V{CONFIG['版本'].split('.')[0]}.docx"

# ===== 样式定义 =====
PRIMARY_RGB = RGBColor(0x1E, 0x5A, 0xFF)
TEXT_1 = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_2 = RGBColor(0x5A, 0x5A, 0x5A)
WARNING = RGBColor(0xFF, 0x6B, 0x35)
SUCCESS = RGBColor(0x00, 0xC8, 0x53)
DANGER = RGBColor(0xE5, 0x39, 0x35)
BG_HEADER = "1E5AFF"
BG_SOFT = "F7F8FA"

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style_normal = doc.styles['Normal']
style_normal.font.name = '苹方'
style_normal.font.size = Pt(12)
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), '苹方')


def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_table_with_header(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, BG_HEADER)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(v)
            if i % 2 == 1:
                set_cell_bg(cell, BG_SOFT)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10.5)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    return table


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = PRIMARY_RGB
        run.font.size = Pt(22)
        run.font.bold = True
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = PRIMARY_RGB
        run.font.size = Pt(16)
        run.font.bold = True
    return p


def add_para(doc, text, bold=False, color=None, size=12, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run.font.name = '苹方'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '苹方')
    return p


def add_bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    run = p.add_run(text)
    run.font.size = Pt(12)
    if color:
        run.font.color.rgb = color
    return p


# ============================================================
# 封面
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')
run.font.size = Pt(40)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'{CONFIG["行业"]} · {CONFIG["场景"]}')
run.font.size = Pt(16)
run.font.color.rgb = TEXT_2

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n需求确认与方案设计表')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = PRIMARY_RGB

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'\n\nV{CONFIG["版本"]} · 面向{CONFIG["客户名"]}')
run.font.size = Pt(20)
run.font.color.rgb = TEXT_1

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\n')
run.font.size = Pt(20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(CONFIG["服务商名"])
run.font.size = Pt(14)
run.font.color.rgb = TEXT_2
run.font.bold = True

date_str = CONFIG["日期"]
if '-' in date_str and len(date_str) == 10:  # yyyy-mm-dd 格式
    parts = date_str.split('-')
    date_display = f"{parts[0]} 年 {int(parts[1])} 月 {int(parts[2])} 日"
else:
    date_display = date_str

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(date_display)
run.font.size = Pt(12)
run.font.color.rgb = TEXT_2

doc.add_page_break()

# ============================================================
# 目录
# ============================================================
add_h1(doc, '目  录')
toc_items = [
    "第 1 章  客户背景",
    "第 2 章  需求理解",
    "第 3 章  方案总览",
    "第 4 章  业务流程设计",
    "第 5 章  功能清单",
    "第 6 章  智能表格设计",
    "第 7 章  实施计划",
    "第 8 章  价值与 ROI",
    "第 9 章  风险与应对",
    "第 10 章  下一步",
]
for i, title in enumerate(toc_items):
    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16), WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(f"{title}\t{3 + i*1}")
    run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 第 1 章
# ============================================================
add_h1(doc, '第 1 章  客户背景')
add_h2(doc, '1.1 客户基本信息')
add_table_with_header(doc,
    ["项目", "内容"],
    [
        ("客户全称", CONFIG["客户全称"]),
        ("所属行业", CONFIG["行业"]),
        ("目标市场", CONFIG["目标市场"]),
        ("主营产品", CONFIG["主营产品"]),
        ("年度订单", CONFIG["年度订单描述"]),
        ("核心客户", CONFIG["核心客户列表"]),
        ("现有系统", CONFIG["现有系统"]),
    ],
    col_widths=[4, 12]
)
add_h2(doc, '1.2 业务背景')
add_para(doc, CONFIG["业务背景1"])
add_para(doc, CONFIG["业务背景2"])
doc.add_page_break()

# ============================================================
# 第 2 章
# ============================================================
add_h1(doc, '第 2 章  需求理解')
add_h2(doc, '2.1 核心痛点清单')
add_table_with_header(doc,
    ["编号", "痛点描述", "影响", "紧迫度"],
    CONFIG["痛点清单"],
    col_widths=[1.5, 7, 5, 1.5]
)
add_h2(doc, '2.2 隐含需求（基于行业知识库推断）')
for item in CONFIG["隐含需求"]:
    add_bullet(doc, item)
doc.add_page_break()

# ============================================================
# 第 3 章
# ============================================================
add_h1(doc, '第 3 章  方案总览')
add_h2(doc, '3.1 一句话价值主张')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(CONFIG["价值主张上"].format(行业=CONFIG["行业"], 场景=CONFIG["场景"]))
run.font.size = Pt(20)
run.font.color.rgb = PRIMARY_RGB
run.font.bold = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(CONFIG["价值主张下"])
run.font.size = Pt(20)
run.font.color.rgb = PRIMARY_RGB
run.font.bold = True

add_h2(doc, '3.2 三大核心模块')
add_table_with_header(doc,
    ["模块", "核心能力", "解决的痛点"],
    [
        ("📊 数据看板", "KPI 大字、状态分布、待办预警、趋势分析", "数据不透明"),
        ("📦 智能表格", "多视图覆盖核心业务", "协同断层"),
        ("🔄 流程自动化", "状态变更自动通知、超期自动催办、收款自动提醒", "沟通成本高"),
    ],
    col_widths=[4, 7, 5]
)
add_h2(doc, '3.3 方案与痛点的对应关系')
# 用 P1/P2/... 编号从 CONFIG 痛点清单动态生成
对应关系 = []
for i, (no, desc, _, _) in enumerate(CONFIG["痛点清单"]):
    对应关系.append((no, f"智能表格 + 智能看板", f"详见 2.1"))
add_table_with_header(doc,
    ["痛点", "对应方案", "预期收益"],
    [
        ("P1 打样进度不透明", "智能表格+状态自动通知", "跟单员沟通时间 -50%"),
        ("P2 数据打架", "统一数据源+权限分级", "对账周期 -67%"),
        ("P3 转化数据缺失", "数据看板+客户档案", "复购识别率 +30%"),
        ("P4 库存风险", "物料追踪+安全库存预警", "紧急补单 -70%"),
        ("P5 客户体验差", "客户侧小程序+进度查询", "NPS 提升至 55+"),
    ],
    col_widths=[4, 6, 6]
)
doc.add_page_break()

# ============================================================
# 第 4 章
# ============================================================
add_h1(doc, '第 4 章  业务流程设计')
add_h2(doc, '4.1 端到端流程概览')
add_para(doc, '优化后的核心流程，标注数字为该环节最常出现痛点的位置：')
add_para(doc, '')
add_table_with_header(doc,
    ["节点", "动作", "参与角色", "关键输出", "状态"],
    CONFIG["流程节点"],
    col_widths=[1.5, 3, 3, 5, 3.5]
)
add_h2(doc, '4.2 优化前后对比')
add_table_with_header(doc,
    ["维度", "优化前（当前）", "优化后（本方案）"],
    CONFIG["优化对比"],
    col_widths=[3, 6.5, 6.5]
)
doc.add_page_break()

# ============================================================
# 第 5 章
# ============================================================
add_h1(doc, '第 5 章  功能清单')
add_para(doc, 'P0 为本期必交付，P1 为本期可交付，P2 为二期规划。')

add_h2(doc, '5.1 数据看板模块')
add_table_with_header(doc,
    ["编号", "功能", "描述", "优先级", "关联表格"],
    CONFIG["看板功能"],
    col_widths=[1.5, 3, 6.5, 1.5, 2.5]
)
add_h2(doc, '5.2 智能表格模块')
add_table_with_header(doc,
    ["编号", "功能", "描述", "优先级", "关联表格"],
    CONFIG["表格功能"],
    col_widths=[1.5, 3, 6.5, 1.5, 2.5]
)
add_h2(doc, '5.3 流程自动化模块')
add_table_with_header(doc,
    ["编号", "功能", "描述", "优先级", "关联表格"],
    CONFIG["自动化功能"],
    col_widths=[1.5, 3, 6.5, 1.5, 2.5]
)
doc.add_page_break()

# ============================================================
# 第 6 章
# ============================================================
add_h1(doc, '第 6 章  智能表格设计')
add_para(doc, '本方案交付 N 个智能表格 Sheet，覆盖业务全流程。字段命名与状态枚举在所有表格中保持一致。')

add_h2(doc, '6.1 字段总览')
add_table_with_header(doc,
    ["Sheet", "核心字段", "字段数", "权限范围"],
    CONFIG["sheet列表"],
    col_widths=[3, 6, 1.5, 5.5]
)
add_h2(doc, '6.2 状态枚举统一规范')
add_para(doc, '以下状态枚举在所有 Sheet 中保持一致，HTML 可视化方案、xlsx 智能表格、企微消息通知均使用同一套表述：')
add_table_with_header(doc,
    ["维度", "状态枚举"],
    CONFIG["状态枚举"],
    col_widths=[4, 12]
)
doc.add_page_break()

# ============================================================
# 第 7 章
# ============================================================
add_h1(doc, '第 7 章  实施计划')
add_para(doc, '总周期 6 周，分 5 个阶段。')
add_table_with_header(doc,
    ["阶段", "时间", "关键里程碑", "交付物", "客户配合"],
    CONFIG["实施计划"],
    col_widths=[2.5, 1.5, 3.5, 3, 5.5]
)
doc.add_page_break()

# ============================================================
# 第 8 章
# ============================================================
add_h1(doc, '第 8 章  价值与 ROI')
add_h2(doc, '8.1 量化价值')
add_table_with_header(doc,
    ["维度", "当前基线", "预期目标", "收益测算"],
    CONFIG["价值ROI"],
    col_widths=[4, 3, 3, 6]
)
add_h2(doc, '8.2 综合 ROI')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'年度总收益：{CONFIG["年度总收益"]}')
run.font.size = Pt(18)
run.font.color.rgb = SUCCESS
run.font.bold = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'项目投资：{CONFIG["项目投资"]}')
run.font.size = Pt(14)
run.font.color.rgb = TEXT_1
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'回本周期：{CONFIG["回本周期"]}')
run.font.size = Pt(18)
run.font.color.rgb = PRIMARY_RGB
run.font.bold = True
doc.add_page_break()

# ============================================================
# 第 9 章
# ============================================================
add_h1(doc, '第 9 章  风险与应对')
add_table_with_header(doc,
    ["风险", "影响", "应对措施", "责任人"],
    CONFIG["风险清单"],
    col_widths=[3, 4, 6, 3]
)
doc.add_page_break()

# ============================================================
# 第 10 章
# ============================================================
add_h1(doc, '第 10 章  下一步')
add_h2(doc, '10.1 ⚠️ 客户方待确认事项')
add_para(doc, '请贵司于收到本方案后 3 个工作日内反馈以下事项：', color=WARNING)
for item in CONFIG["待确认事项"]:
    add_bullet(doc, item, color=WARNING)
add_h2(doc, '10.2 服务商方待提供材料')
for item in CONFIG["服务商待提供"]:
    add_bullet(doc, item)
add_h2(doc, '10.3 联系方式')
add_table_with_header(doc,
    ["项目", "信息"],
    CONFIG["联系方式"],
    col_widths=[3, 13]
)

# 页脚
section = doc.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'© {CONFIG["服务商名"]} · 机密')
run.font.size = Pt(9)
run.font.color.rgb = TEXT_2

# 保存
doc.save(OUTPUT)
print(f"✅ 已生成: {OUTPUT}")
print(f"   共 10 章 + 封面 + 目录")
