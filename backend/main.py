"""
Provider Assist 后端API - FastAPI
轻量化版本：SQLite + JWT认证 + 知识库管理
"""
import os
import json
import re
import zipfile
import io
from pathlib import Path
from string import Template
from datetime import datetime
from typing import List, Optional
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse

from database import get_db, init_db, init_kb_db
from auth import (
    get_password_hash, verify_password, create_access_token,
    UserRegister, UserLogin, UserResponse, require_auth,
    validate_invitation_code, mark_invitation_code_used, seed_invitation_codes, seed_dev_user
)

# 自动加载 .env 文件
load_dotenv(Path(__file__).parent / ".env")

BASE_DIR = Path(__file__).parent.parent
KNOWLEDGE_DIR = BASE_DIR / "knowledge"
KB_DIR = BASE_DIR / "data" / "kb"
KB_DIR.mkdir(parents=True, exist_ok=True)

# API Keys - 从环境变量读取
MINIMAX_API_KEY = os.environ.get("MINIMAX_API_KEY", "")
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")
TAVILY_API_KEY = os.environ.get("TAVILY_API_KEY", "")


def estimate_tokens(text: str) -> int:
    """
    使用 tiktoken 估算文本的 token 数量（兼容 GPT-4/Claude 等）
    """
    try:
        import tiktoken
        # 使用 cl100k_base 编码器（GPT-4/Claude/大多数模型通用）
        enc = tiktoken.get_encoding("cl100k_base")
        return len(enc.encode(text))
    except Exception:
        # fallback：粗略估算（中文字符约 2 token，英文约 0.25 token）
        chinese_chars = sum(1 for c in text if '一' <= c <= '鿿')
        other_chars = len(text) - chinese_chars
        return int(chinese_chars * 2 + other_chars * 0.25)


def estimate_cost(tokens: int, model: str = "deepseek-chat") -> float:
    """
    估算 API 调用费用（单位：元）
    model: deepseek-chat / gpt-4 / claude-3-sonnet
    """
    # DeepSeek 价格：输入 ¥1/百万token，输出 ¥2/百万token（估算）
    if "deepseek" in model.lower():
        return tokens / 1_000_000 * 1.5  # 平均约 ¥1.5/百万
    # 其他模型价格可扩展
    return tokens / 1_000_000 * 10  # 默认估算


def call_deepseek(system_prompt: str, user_prompt: str, max_tokens: int = 4000) -> dict:
    """调用 DeepSeek API，300秒超时，最多一次重试
    返回 {"content": str, "usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}}"""
    import httpx

    def _do_request():
        response = httpx.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
            },
            json={
                "model": "deepseek-chat",
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                "max_tokens": max_tokens,
                "temperature": 0.7
            },
            timeout=httpx.Timeout(300.0, connect=30.0)
        )
        if response.status_code != 200:
            return {"content": f"Error: DeepSeek API 返回 {response.status_code}: {response.text}", "usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}}
        data = response.json()
        content = data.get("choices", [{}])[0].get("message", {}).get("content", "")
        usage = data.get("usage", {})
        return {"content": content, "usage": {
            "prompt_tokens": usage.get("prompt_tokens", 0) or 0,
            "completion_tokens": usage.get("completion_tokens", 0) or 0,
            "total_tokens": usage.get("total_tokens", 0) or 0,
        }}

    try:
        return _do_request()
    except httpx.TimeoutException:
        try:
            return _do_request()
        except httpx.TimeoutException:
            return {"content": "Error: DeepSeek API 请求超时（300秒），请稍后重试", "usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}}
        except Exception as e:
            return {"content": f"Error: {str(e)}", "usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}}
    except Exception as e:
        return {"content": f"Error: {str(e)}", "usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}}

def tavily_search(query: str, max_results: int = 8) -> dict:
    """使用 Tavily API 进行实时搜索，返回搜索结果摘要"""
    try:
        from tavily import TavilyClient
        client = TavilyClient(api_key=TAVILY_API_KEY)
        results = client.search(query, max_results=max_results, search_depth="advanced")
        # 提取关键信息供 AI 分析
        search_summary = []
        for r in results.get("results", []):
            search_summary.append({
                "title": r.get("title", ""),
                "content": r.get("content", ""),
                "url": r.get("url", "")
            })
        return {"success": True, "results": search_summary, "raw": results}
    except Exception as e:
        return {"success": False, "error": str(e)}

# ==================== 通用辅助函数 ====================

def call_mcp(tool_name: str, arguments: dict) -> dict:
    """调用企微 API（通过本地认证的 wecom-cli）"""
    import subprocess, json as jsonmod

    # 使用新机器人凭证（解决旧机器人创建的文档无权限访问问题）
    env = {
        "WECOM_BOT_ID": "aibmNuLSHYId2jvHy68XtU6HWhwCHZnvkpD",
        "WECOM_BOT_SECRET": "qE6ePDNGRv8PZZft7a1aqaQpNHQLtw8lIcjfJmwc55A",
    }

    cmd = ["wecom-cli", "doc", tool_name, "--json", jsonmod.dumps(arguments, ensure_ascii=False)]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60, env={**__import__("os").environ, **env})
        output = result.stdout
        if result.returncode != 0 and result.stderr:
            if "authorization expired" in result.stderr.lower():
                return {"errcode": 851014, "errmsg": "authorization expired", "help_message": "请重新运行 wecom-cli init 授权"}
        try:
            outer = jsonmod.loads(output)
            content = outer.get("result", {}).get("content", [{}])[0].get("text", "{}")
            return jsonmod.loads(content) if isinstance(content, str) else content
        except:
            return jsonmod.loads(output) if output else {"errcode": -1, "errmsg": "parse error"}
    except subprocess.TimeoutExpired:
        return {"errcode": -1, "errmsg": "wecom-cli timeout"}
    except Exception as e:
        return {"error": str(e)}


def extract_mcp(mcp_resp: dict):
    """从 MCP 响应中提取实际数据"""
    if not mcp_resp:
        return None
    result = mcp_resp.get("result", mcp_resp)
    if isinstance(result, dict):
        content = result.get("content", [])
        if isinstance(content, list):
            for item in content:
                if isinstance(item, dict) and item.get("type") == "text":
                    try:
                        return json.loads(item["text"])
                    except:
                        return item.get("text")
        return result
    return result


MINIMAX_API_KEY = "sk-cp-FxfZUSUHnTWn7eCtl1V-5CI1jFpfF3XLI0jxHZJ7U0p16_cea_FTQqxOaOYavdfwiS9DDN4pomf4CxLZlQYqIyvJJK_eaKR7tbh4d77_1dGK8DwQtwwjLDc"

def call_minimax(system_prompt: str, user_prompt: str, max_tokens: int = 4000) -> dict:
    """调用 MiniMax API，300秒超时，最多一次重试
    返回 {"content": str, "usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}}"""
    import httpx

    def _do_request():
        response = httpx.post(
            "https://api.minimax.chat/v1/text/chatcompletion_v2",
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {MINIMAX_API_KEY}"
            },
            json={
                "model": "abab6.5s-chat",
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                "max_tokens": max_tokens,
                "temperature": 0.7
            },
            timeout=httpx.Timeout(300.0, connect=15.0)
        )
        result = response.json()
        content = result["choices"][0]["message"]["content"]
        # 提取 usage（MiniMax 返回结构）
        usage = result.get("usage", {}) or {}
        return {
            "content": content,
            "usage": {
                "prompt_tokens": usage.get("prompt_tokens", 0) or 0,
                "completion_tokens": usage.get("completion_tokens", 0) or 0,
                "total_tokens": usage.get("total_tokens", 0) or 0,
            }
        }

    try:
        return _do_request()
    except httpx.TimeoutException:
        try:
            return _do_request()
        except httpx.TimeoutException:
            return {"content": "Error: MiniMax API 请求超时（300秒），请稍后重试", "usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}}
        except Exception as e:
            return {"content": f"Error: {str(e)}", "usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}}
    except Exception as e:
        return {"content": f"Error: {str(e)}", "usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}}


def call_codebuddy(system_prompt: str, user_prompt: str, max_tokens: int = 4000, timeout: int = 600) -> dict:
    """调用 CodeBuddy CLI（单行模式），默认600秒超时，最多一次重试
    如果第一次输出不是 HTML（不以 < 开头），自动重试并强制要求只输出 HTML
    返回 {"content": str, "usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}}"""
    import subprocess

    full_prompt = f"{system_prompt}\n\n---\n\n{user_prompt}"

    def _do_request(extra_directive=""):
        prompt_to_send = full_prompt + ("\n\n" + extra_directive if extra_directive else "")
        result = subprocess.run(
            ["codebuddy", prompt_to_send, "-p", "-y", "--output-format", "text"],
            capture_output=True, text=True, timeout=timeout
        )
        content = result.stdout.strip() if result.stdout.strip() else result.stderr.strip()
        return content

    try:
        content = _do_request()
        # 如果输出不是 HTML（不以 < 开头），重试并强制要求只输出 HTML
        if not content.startswith("<"):
            content = _do_request(
                "【重要】上一次输出不是 HTML。请直接在消息中输出完整的 HTML 代码（以 <html 开头），不要输出任何解释文字、分析说明或 markdown 代码块包裹。只需 HTML 代码本身。"
            )
        return {"content": content, "usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}}
    except subprocess.TimeoutExpired:
        return {"content": f"Error: CodeBuddy CLI 执行超时（{timeout}秒），请稍后重试", "usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}}
    except Exception as e:
        return {"content": f"Error: {str(e)}", "usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}}


def record_ai_tokens(client_id: int, tokens: int):
    """将 AI 消耗的 token 数累加到客户的 token_count"""
    if not client_id or not tokens:
        return
    try:
        conn = get_db()
        cursor = conn.cursor()
        cursor.execute("UPDATE clients SET token_count = COALESCE(token_count, 0) + ? WHERE id = ?", (tokens, client_id))
        conn.commit()
        conn.close()
    except Exception:
        pass


app = FastAPI(title="Provider Assist API", version="1.0.0")

# 所有 OPTIONS 请求直接返回 200（处理 preflight）
@app.middleware("http")
async def cors_options_middleware(request, call_next):
    if request.method == "OPTIONS":
        from starlette.responses import JSONResponse
        return JSONResponse(
            {"status": "ok"},
            status_code=200,
            headers={
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "GET, POST, PUT, DELETE, OPTIONS",
                "Access-Control-Allow-Headers": "*",
                "Access-Control-Max-Age": "600",
            },
        )
    return await call_next(request)

# CORS 允许所有来源（含 localhost:9090）
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
    allow_credentials=True,
)

# 初始化数据库
init_db()
init_kb_db()

# 挂载 public 静态文件目录（用于 Agent Demo H5）
public_dir = Path(__file__).parent / "public"
public_dir.mkdir(exist_ok=True)
app.mount("/public", StaticFiles(directory=str(public_dir)), name="public")

# outputs 目录（存放生成的 Word/HTML 等产物）
outputs_dir = Path(__file__).parent.parent / "outputs"
outputs_dir.mkdir(exist_ok=True)
app.mount("/outputs", StaticFiles(directory=str(outputs_dir)), name="outputs")

# reports 目录（Skill 生成的 HTML 报告托管）
reports_dir = Path(__file__).parent.parent / "data" / "reports"
reports_dir.mkdir(parents=True, exist_ok=True)
app.mount("/reports", StaticFiles(directory=str(reports_dir)), name="reports")

# 初始化测试用受邀码
seed_invitation_codes()
seed_dev_user()

# ==================== 认证相关 ====================

@app.post("/api/auth/register")
async def register(user: dict):
    """用户注册 - 需要有效的受邀码"""
    import re

    provider_name = user.get("provider_name", "").strip()
    invitation_code = user.get("invitation_code", "").strip()
    username = user.get("username", "").strip()
    password = user.get("password", "")

    # 字段校验
    if not provider_name: raise HTTPException(status_code=400, detail="服务商名称不能为空")
    if not invitation_code: raise HTTPException(status_code=400, detail="受邀码不能为空")
    if not username: raise HTTPException(status_code=400, detail="用户名不能为空")
    if not password: raise HTTPException(status_code=400, detail="密码不能为空")

    # 密码格式校验
    if len(password) < 8 or len(password) > 25:
        raise HTTPException(status_code=400, detail="密码长度必须为8-25位")
    types = sum([
        bool(re.search(r'[0-9]', password)),
        bool(re.search(r'[a-z]', password)),
        bool(re.search(r'[A-Z]', password)),
        bool(re.search(r'[!@#$%^&*()_+\-=\[\]{};:\'",.<>?/\\|`~]', password))
    ])
    if types < 2:
        raise HTTPException(status_code=400, detail="密码必须包含至少2种不同字符类型")

    # 校验受邀码（返回 provider_name 以邀请码为准）
    is_valid, error_msg, prov_from_code = validate_invitation_code(invitation_code, provider_name)
    if not is_valid:
        raise HTTPException(status_code=400, detail=error_msg)
    # 企业名称以邀请码为准（防止注册页伪造）
    provider_name = prov_from_code

    conn = get_db()
    cursor = conn.cursor()

    # 检查是否已存在
    cursor.execute("SELECT id FROM users WHERE username = ?", (username,))
    if cursor.fetchone():
        raise HTTPException(status_code=400, detail="用户名已存在")

    # 创建用户
    password_hash = get_password_hash(password)
    cursor.execute(
        "INSERT INTO users (username, password_hash, provider_name) VALUES (?, ?, ?)",
        (username, password_hash, provider_name)
    )
    conn.commit()

    user_id = cursor.lastrowid

    # 标记受邀码已使用
    mark_invitation_code_used(invitation_code, user_id)

    conn.close()

    # 生成token
    token = create_access_token({"sub": username, "user_id": user_id})

    return {
        "success": True,
        "access_token": token,
        "token_type": "bearer",
        "user": {"id": user_id, "username": username, "provider_name": provider_name}
    }

@app.get("/api/invitation-code/{code}")
async def get_invitation_code_info(code: str):
    """查询邀请码对应的企业名称（注册页预填用）"""
    is_valid, error_msg, provider_name = validate_invitation_code(code)
    if not is_valid:
        raise HTTPException(status_code=400, detail=error_msg)
    return {"provider_name": provider_name}

@app.post("/api/auth/login", response_model=dict)
async def login(user: UserLogin):
    """用户登录 - 已注册用户直接登录"""
    conn = get_db()
    cursor = conn.cursor()

    cursor.execute("SELECT id, username, password_hash, provider_name FROM users WHERE username = ?", (user.username,))
    row = cursor.fetchone()
    conn.close()

    if not row or not verify_password(user.password, row["password_hash"]):
        raise HTTPException(status_code=401, detail="用户名或密码错误")

    # 再次确认服务商名称匹配
    if row["provider_name"] != user.provider_name:
        raise HTTPException(status_code=401, detail="服务商名称与受邀码不匹配")

    token = create_access_token({"sub": row["username"], "user_id": row["id"]})

    return {
        "success": True,
        "access_token": token,
        "token_type": "bearer",
        "user": {"id": row["id"], "username": row["username"], "provider_name": row["provider_name"] or ""}
    }

@app.get("/api/debug-saved")
async def debug_saved(user: dict = Depends(require_auth)):
    """Debug: 检查 _saved 列"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT id, name, _completed, _saved FROM clients WHERE id = 71")
    row = cursor.fetchone()
    conn.close()
    d = dict(row)
    # 直接返回确认有 _saved
    import sys
    print(f"_saved in d: {'_saved' in d}, value: {d.get('_saved')}", file=sys.stderr)
    return {"_saved_val": d.get("_saved"), "completed_val": d.get("_completed"), "raw": d}

@app.post("/api/test-login")
async def test_login():
    """极简测试登录 - 一键登录测试账号"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT id, username, provider_name FROM users WHERE username = 'devuser'")
    row = cursor.fetchone()
    conn.close()

    if not row:
        return {"error": "测试用户不存在"}

    token = create_access_token({"sub": row["username"], "user_id": row["id"]})
    return {
        "success": True,
        "access_token": token,
        "token_type": "bearer",
        "user": {"id": row["id"], "username": row["username"], "provider_name": row["provider_name"]}
    }


@app.post("/api/auth/dev-login", response_model=dict)
async def dev_login(user: dict):
    """开发者登录 - 不校验受邀码，直接登录已注册用户"""
    username = user.get("username")
    password = user.get("password")
    if not username or not password:
        raise HTTPException(status_code=400, detail="缺少用户名或密码")

    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT id, username, password_hash, provider_name FROM users WHERE username = ?", (username,))
    row = cursor.fetchone()
    conn.close()

    if not row or not verify_password(password, row["password_hash"]):
        raise HTTPException(status_code=401, detail="用户名或密码错误")

    token = create_access_token({"sub": row["username"], "user_id": row["id"]})
    return {
        "success": True,
        "access_token": token,
        "token_type": "bearer",
        "user": {"id": row["id"], "username": row["username"], "provider_name": row["provider_name"] or ""}
    }

@app.post("/api/auth/auto-login")
async def auto_login(body: dict, user: dict = Depends(require_auth)):
    """自动登录 - 检查token是否有效"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT username, provider_name FROM users WHERE id = ?", (user["user_id"],))
    row = cursor.fetchone()
    conn.close()
    if row:
        return {"success": True, "user": {"id": user["user_id"], "username": row[0], "provider_name": row[1]}}
    return {"success": True, "user": {"id": user["user_id"], "username": user["sub"]}}

@app.get("/api/auth/me")
async def get_me(user: dict = Depends(require_auth)):
    """获取当前用户信息"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT id, username, provider_name FROM users WHERE id = ?", (user["user_id"],))
    row = cursor.fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="用户不存在")

    return {"id": row["id"], "username": row["username"], "provider_name": row["provider_name"] or ""}

# ==================== 知识库 ====================

def load_global_knowledge():
    """加载全局知识库"""
    industries = {}
    ind_dir = KNOWLEDGE_DIR / "industries"
    if ind_dir.exists():
        for f in ind_dir.glob("*.json"):
            try:
                data = json.loads(f.read_text(encoding="utf-8"))
                industries[f.stem] = data
            except:
                pass

    cases = []
    cases_dir = KNOWLEDGE_DIR / "cases"
    if cases_dir.exists():
        for f in cases_dir.glob("*.json"):
            try:
                data = json.loads(f.read_text(encoding="utf-8"))
                cases.append(data)
            except:
                pass

    templates = []
    tpl_dir = KNOWLEDGE_DIR / "field_templates"
    if tpl_dir.exists():
        for f in tpl_dir.glob("*.json"):
            try:
                data = json.loads(f.read_text(encoding="utf-8"))
                templates.append(data)
            except:
                pass

    return {"industries": industries, "cases": cases, "templates": templates}

@app.get("/api/knowledge/global")
async def get_global_knowledge():
    """获取全局知识库摘要"""
    kb = load_global_knowledge()
    return {
        "industries_count": len(kb["industries"]),
        "cases_count": len(kb["cases"]),
        "templates_count": len(kb["templates"]),
        "industries": [
            {"key": k, "name": v.get("industry_name", k), "tags": v.get("tags", [])}
            for k, v in list(kb["industries"].items())[:10]
        ]
    }

def _do_knowledge_search(industry: str, keywords: List[str], user_id: int = None):
    """内部知识库搜索逻辑（供路由和报告生成共用）"""
    kb = load_global_knowledge()

    # 1. 匹配行业
    industry_lower = industry.lower()
    matched_industry = None
    for key, data in kb["industries"].items():
        tags = [t.lower() for t in data.get("tags", [])]
        name = data.get("industry_name", "").lower()
        if industry_lower in name or any(industry_lower in t or t in industry_lower for t in tags):
            matched_industry = data
            break

    # 2. 匹配案例
    query = industry + " " + " ".join(keywords)
    query_lower = query.lower()
    scored_cases = []
    for case in kb["cases"]:
        score = 0
        meta = case.get("meta", {})
        case_industry = meta.get("industry", "").lower()
        case_scene = meta.get("scene", "").lower()
        if case_industry in query_lower or query_lower in case_industry:
            score += 5
        else:
            for word in re.split(r'[，,、。/\s]+', query_lower):
                if len(word) >= 2 and word in case_industry:
                    score += 4
                    break
        if case_scene in query_lower:
            score += 3
        if score > 0:
            scored_cases.append((score, case))

    scored_cases.sort(key=lambda x: x[0], reverse=True)
    matched_cases = [c for _, c in scored_cases[:3]]

    # 3. 匹配模板
    scored_templates = []
    for tpl in kb["templates"]:
        score = 0
        meta = tpl.get("meta", {})
        tpl_industry = meta.get("industry", "").lower()
        applicable = meta.get("applicable_when", "").lower()
        if tpl_industry in query_lower:
            score += 5
        for word in keywords:
            if word.lower() in applicable:
                score += 2
        if score >= 4:
            scored_templates.append((score, tpl))

    scored_templates.sort(key=lambda x: x[0], reverse=True)
    matched_templates = [t for _, t in scored_templates[:2]]

    # 4. 服务商私有知识库
    conn = get_db()
    cursor = conn.cursor()
    user_kb = []
    if user_id:
        if industry:
            cursor.execute(
                "SELECT * FROM provider_knowledge WHERE user_id = ? AND (industry = ? OR category = 'industry_knowledge')",
                (user_id, industry)
            )
        else:
            cursor.execute("SELECT * FROM provider_knowledge WHERE user_id = ?", (user_id,))
        for row in cursor.fetchall():
            user_kb.append(dict(row))
    conn.close()

    return {
        "industry_knowledge": matched_industry.get("content", "")[:3000] if matched_industry else "",
        "matched_cases": matched_cases,
        "matched_templates": matched_templates,
        "user_knowledge": user_kb,
        "matched": bool(matched_industry or matched_cases or matched_templates)
    }

@app.post("/api/knowledge/search")
async def search_knowledge(body: dict, user: dict = Depends(require_auth)):
    """搜索知识库（API路由）"""
    industry = body.get("industry", "")
    keywords = body.get("keywords", [])
    return _do_knowledge_search(industry, keywords, user["user_id"])

# ==================== 服务商知识库管理 ====================

@app.post("/api/provider-knowledge")
async def add_provider_knowledge(
    body: dict,
    user: dict = Depends(require_auth)
):
    """添加服务商知识库条目"""
    category = body.get("category", "")
    title = body.get("title", "")
    content = body.get("content", "")
    industry = body.get("industry", "")
    tags = body.get("tags", "")
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO provider_knowledge (user_id, category, title, content, industry, tags) VALUES (?, ?, ?, ?, ?, ?)",
        (user["user_id"], category, title, content, industry, tags)
    )
    conn.commit()
    item_id = cursor.lastrowid
    conn.close()

    return {"success": True, "id": item_id}

@app.get("/api/provider-knowledge")
async def list_provider_knowledge(category: str = "", user: dict = Depends(require_auth)):
    """获取服务商知识库列表"""
    conn = get_db()
    cursor = conn.cursor()

    if category:
        cursor.execute(
            "SELECT * FROM provider_knowledge WHERE user_id = ? AND category = ? ORDER BY created_at DESC",
            (user["user_id"], category)
        )
    else:
        cursor.execute(
            "SELECT * FROM provider_knowledge WHERE user_id = ? ORDER BY created_at DESC",
            (user["user_id"],)
        )

    items = [dict(row) for row in cursor.fetchall()]
    conn.close()

    return items

@app.delete("/api/provider-knowledge/{item_id}")
async def delete_provider_knowledge(item_id: int, user: dict = Depends(require_auth)):
    """删除服务商知识库条目"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM provider_knowledge WHERE id = ? AND user_id = ?", (item_id, user["user_id"]))
    conn.commit()
    conn.close()
    return {"success": True}

@app.get("/api/provider-knowledge/stats")
async def get_knowledge_stats(user: dict = Depends(require_auth)):
    """获取服务商知识库统计"""
    conn = get_db()
    cursor = conn.cursor()

    stats = {}
    categories = ["case", "template", "qa", "sales_tool", "industry_knowledge"]
    for cat in categories:
        cursor.execute(
            "SELECT COUNT(*) FROM provider_knowledge WHERE user_id = ? AND category = ?",
            (user["user_id"], cat)
        )
        stats[cat] = cursor.fetchone()[0]

    conn.close()
    return stats

# ==================== 管理后台 API ====================
import secrets
import string

@app.get("/api/admin/stats")
async def admin_stats(user: dict = Depends(require_auth)):
    """管理后台统计数据"""
    conn = get_db()
    cursor = conn.cursor()

    # 注册用户数
    cursor.execute("SELECT COUNT(*) FROM users")
    total_users = cursor.fetchone()[0]

    # 今日新增
    cursor.execute("SELECT COUNT(*) FROM users WHERE date(created_at) = date('now')")
    today_users = cursor.fetchone()[0]

    # 客户总数
    cursor.execute("SELECT COUNT(*) FROM clients")
    total_clients = cursor.fetchone()[0]

    # 已分配受邀码
    cursor.execute("SELECT COUNT(*) FROM invitation_codes WHERE used > 0")
    assigned_codes = cursor.fetchone()[0]

    # 可用受邀码
    cursor.execute("SELECT COUNT(*) FROM invitation_codes WHERE used < max_users OR max_users IS NULL")
    available_codes = cursor.fetchone()[0]

    # 漏斗：已完成客户数
    cursor.execute("SELECT COUNT(*) FROM clients WHERE is_completed = 1")
    completed = cursor.fetchone()[0]

    # 7日活跃客户
    cursor.execute("SELECT COUNT(*) FROM clients WHERE updated_at >= datetime('now', '-7 days')")
    active_7d = cursor.fetchone()[0]

    # 服务商平均客户数
    avg_clients = round(total_clients / total_users, 1) if total_users > 0 else 0

    # AI 调用相关（改用 token_count）
    cursor.execute("""
        SELECT
            COUNT(*) as ai_calls,
            SUM(COALESCE(token_count, 0)) as total_tokens
        FROM clients
        WHERE token_count > 0
    """)
    ai_row = cursor.fetchone()
    total_ai_calls = ai_row["ai_calls"] or 0
    total_tokens = ai_row["total_tokens"] or 0

    # Step S1-S5 漏斗（真实数量）
    cursor.execute("SELECT COUNT(*) FROM clients WHERE step1_result IS NOT NULL AND step1_result != ''")
    s1_count = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT(*) FROM clients WHERE step2_report IS NOT NULL AND step2_report != ''")
    s2_count = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT(*) FROM clients WHERE step3_summary IS NOT NULL AND step3_summary != ''")
    s3_count = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT(*) FROM clients WHERE step4_input_draft IS NOT NULL AND step4_input_draft != ''")
    s4_count = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT(*) FROM clients WHERE step5_schema IS NOT NULL AND step5_schema != ''")
    s5_count = cursor.fetchone()[0]

    # 按服务商 Token 排行
    cursor.execute("""
        SELECT u.provider_name,
               SUM(COALESCE(c.token_count, 0)) as tokens,
               COUNT(c.id) as client_count
        FROM clients c JOIN users u ON c.user_id = u.id
        WHERE c.token_count > 0
        GROUP BY u.provider_name
        ORDER BY tokens DESC
        LIMIT 10
    """)
    provider_rank = [dict(r) for r in cursor.fetchall()]

    # 按客户 Token 排行（只看有消耗的）
    cursor.execute("""
        SELECT c.id, c.name, u.provider_name,
               COALESCE(c.token_count, 0) as tokens
        FROM clients c JOIN users u ON c.user_id = u.id
        WHERE c.token_count > 0
        ORDER BY tokens DESC
        LIMIT 10
    """)
    client_rank = [dict(r) for r in cursor.fetchall()]

    conn.close()
    return {
        "total_users": total_users,
        "today_users": today_users,
        "total_clients": total_clients,
        "assigned_codes": assigned_codes,
        "available_codes": available_codes,
        "funnel": {
            "registered": total_users,
            "clients": total_clients,
            "completed": completed,
            "s1": s1_count,
            "s2": s2_count,
            "s3": s3_count,
            "s4": s4_count,
            "s5": s5_count,
        },
        "active_clients_7d": active_7d,
        "clients_per_provider_avg": avg_clients,
        "total_ai_calls": total_ai_calls,
        "total_tokens": total_tokens,
        "provider_ai_rank": provider_rank,
        "client_ai_rank": client_rank,
    }

@app.get("/api/admin/users")
async def admin_list_users(user: dict = Depends(require_auth)):
    """服务商用户列表"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT u.id, u.username, u.provider_name, u.created_at, u.grade,
               (SELECT COUNT(*) FROM clients WHERE user_id = u.id) as client_count
        FROM users u ORDER BY u.id DESC
    """)
    rows = cursor.fetchall()
    conn.close()
    return [{"id": r["id"], "username": r["username"], "provider_name": r["provider_name"],
             "created_at": r["created_at"], "client_count": r["client_count"],
             "grade": r["grade"] or "普通"} for r in rows]

@app.put("/api/admin/users/{uid}")
async def admin_update_user(uid: int, body: dict, user: dict = Depends(require_auth)):
    """更新服务商用户（分级等）"""
    conn = get_db()
    cursor = conn.cursor()
    grade = body.get("grade")
    if grade:
        cursor.execute("UPDATE users SET grade = ? WHERE id = ?", (grade, uid))
        conn.commit()
    conn.close()
    return {"success": True}

@app.delete("/api/admin/users/{uid}")
async def admin_delete_user(uid: int, user: dict = Depends(require_auth)):
    """删除服务商用户"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM users WHERE id = ?", (uid,))
    conn.commit()
    conn.close()
    return {"success": True}

@app.get("/api/admin/clients")
async def admin_list_clients(user: dict = Depends(require_auth)):
    """客户列表（不含大字段，速度快）"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT c.id, c.name, c.industry, c.status, c.created_at, c.updated_at,
               c.user_id, c.is_completed, c.is_saved,
               c.step1_result, c.step2_report, c.step3_summary,
               c.step4_input_draft, c.step5_schema, c.step5_agent_suggestions,
               c.scale, c.tags, c.initial_demand,
               (c.step1_result IS NOT NULL AND c.step1_result != '') AS has_step1,
               (c.step2_report IS NOT NULL AND c.step2_report != '') AS has_step2,
               ((c.step3_summary IS NOT NULL AND c.step3_summary != '') OR (c.uploaded_files IS NOT NULL AND c.uploaded_files != '')) AS has_step3,
               ((c.step4_presales_versions IS NOT NULL AND c.step4_presales_versions != '') OR (c.step4_technical_versions IS NOT NULL AND c.step4_technical_versions != '') OR (c.step4_input_draft IS NOT NULL AND c.step4_input_draft != '')) AS has_step4,
               (c.step5_schema IS NOT NULL AND c.step5_schema != '') AS has_step5,
               u.provider_name
        FROM clients c JOIN users u ON c.user_id = u.id
        ORDER BY c.id DESC LIMIT 500
    """)
    rows = cursor.fetchall()
    conn.close()
    return [{
        "id": r["id"],
        "name": r["name"],
        "industry": r["industry"],
        "status": r["status"],
        "created_at": r["created_at"],
        "updated_at": r["updated_at"],
        "user_id": r["user_id"],
        "is_completed": bool(r["is_completed"]),
        "is_saved": bool(r["is_saved"]),
        "provider_name": r["provider_name"],
        "scale": r["scale"] or "",
        "tags": r["tags"] or "",
        "initial_demand": r["initial_demand"] or "",
        "hasStep1": bool(r["has_step1"]),
        "hasStep2": bool(r["has_step2"]),
        "hasStep3": bool(r["has_step3"]),
        "hasStep4": bool(r["has_step4"]),
        "hasStep5": bool(r["has_step5"]),
        # 大字段：空字符串占位，详情从 /api/admin/clients/{id} 获取
        "step1_result": r["step1_result"] or "",
        "step2_report": r["step2_report"] or "",
        "step3_summary": r["step3_summary"] or "",
        "step4_input_draft": r["step4_input_draft"] or "",
        "step5_schema": r["step5_schema"] or "",
        "step5_agent_suggestions": r["step5_agent_suggestions"] or "",
    } for r in rows]


@app.get("/api/admin/clients/{client_id}")
async def admin_get_client(client_id: int, user: dict = Depends(require_auth)):
    """客户详情（含所有大字段）"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT c.*, u.provider_name
        FROM clients c JOIN users u ON c.user_id = u.id
        WHERE c.id = ?
    """, (client_id,))
    row = cursor.fetchone()
    conn.close()
    if not row:
        return {"error": "客户不存在"}
    return {k: (v or "") for k, v in dict(row).items()}


@app.get("/api/admin/kb/files/{file_id}/download")
async def admin_download_kb_file(file_id: str, user: dict = Depends(require_auth)):
    """管理员下载任意知识库文件（不受 user_id 限制）"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT filepath, original_filename FROM kb_files WHERE id = ?",
        (file_id,))
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="文件不存在")
    filepath = row["filepath"]
    filename = row["original_filename"]
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(filepath, filename=filename, media_type="application/octet-stream")


@app.get("/api/admin/kb/files/{file_id}/content")
async def admin_get_kb_file_content(file_id: str, user: dict = Depends(require_auth)):
    """获取 KB 文件内容（用于预览）"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT filepath, original_filename FROM kb_files WHERE id = ?",
        (file_id,))
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="文件不存在")
    filepath = row["filepath"]
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="文件不存在")
    # Try to read as text
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        return {"content": content, "filename": row["original_filename"]}
    except:
        raise HTTPException(status_code=422, detail="该文件不支持预览")


@app.delete("/api/admin/kb/files/{file_id}")
async def admin_delete_kb_file(file_id: str, user: dict = Depends(require_auth)):
    """删除 KB 文件"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT id, filepath FROM kb_files WHERE id = ?", (file_id,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="文件不存在")
    cursor.execute("DELETE FROM kb_files WHERE id = ?", (file_id,))
    conn.commit()
    conn.close()
    # Optionally delete the physical file
    # os.remove(row["filepath"])  # 保留原文件，只删数据库记录
    return {"success": True}


@app.patch("/api/admin/kb/files/{file_id}")
async def admin_rename_kb_file(file_id: str, display_name: str = Form(...), user: dict = Depends(require_auth)):
    """重命名 KB 文件"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM kb_files WHERE id = ?", (file_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="文件不存在")
    cursor.execute("UPDATE kb_files SET display_name = ? WHERE id = ?", (display_name, file_id))
    conn.commit()
    conn.close()
    return {"success": True}


@app.get("/api/admin/kb/providers/{provider_name}/download-all")
async def admin_download_provider_kb(provider_name: str, user: dict = Depends(require_auth)):
    """管理员一键下载指定服务商的所有知识库文件（打包成zip）"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT k.filepath, k.original_filename, k.display_name, k.category,
               k.industry, k.created_at, u.provider_name
        FROM kb_files k
        JOIN users u ON k.user_id = u.id
        WHERE u.provider_name = ?
        ORDER BY k.created_at DESC
    """, (provider_name,))
    rows = cursor.fetchall()
    conn.close()

    if not rows:
        raise HTTPException(status_code=404, detail="该服务商暂无知识库文件")

    # Create zip in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for row in rows:
            filepath = row["filepath"]
            original_filename = row["original_filename"] or "未命名"
            if os.path.exists(filepath):
                # Read file content and add to zip
                with open(filepath, 'rb') as f:
                    file_data = f.read()
                # Use display_name if available, otherwise original filename
                display_name = row["display_name"] or original_filename
                # Preserve original extension
                ext = os.path.splitext(original_filename)[1]
                if ext and not display_name.endswith(ext):
                    display_name += ext
                # Ensure filename is ASCII-safe for zip compatibility
                safe_name = str(display_name).encode('ascii', 'replace').decode('ascii')
                zf.writestr(safe_name, file_data)
            else:
                # File missing on disk — add a placeholder text file
                placeholder = f"【提示】原始文件已丢失: {original_filename}\n路径: {filepath}"
                zf.writestr(f"[丢失]_{original_filename}.txt", placeholder.encode('utf-8'))

    zip_buffer.seek(0)
    import urllib.parse
    safe_name = re.sub(r'[^\w\-_. ]', '_', provider_name)
    zip_name = f"kb_{safe_name}.zip"
    cd_value = f"attachment; filename*=UTF-8''{urllib.parse.quote(zip_name)}"
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": cd_value}
    )


@app.get("/api/kb/files/{file_id}/download")
async def download_kb_file(file_id: str, user: dict = Depends(require_auth)):
    """下载知识库文件"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT filepath, original_filename FROM kb_files WHERE id = ? AND user_id = ?",
        (file_id, user["user_id"])
    )
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="文件不存在")
    filepath = row["filepath"]
    filename = row["original_filename"]
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(filepath, filename=filename, media_type="application/octet-stream")


@app.get("/api/admin/knowledge")
async def admin_list_knowledge(user: dict = Depends(require_auth)):
    """所有服务商知识库文件（按服务商分组）"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT k.id, k.original_filename, k.display_name, k.category, k.industry,
               k.status, k.progress, k.char_count, k.created_at,
               k.user_id, u.provider_name
        FROM kb_files k
        JOIN users u ON k.user_id = u.id
        ORDER BY u.provider_name, k.created_at DESC
    """)
    rows = cursor.fetchall()
    conn.close()
    # 按 provider_name 分组
    grouped = {}
    for r in rows:
        pname = r["provider_name"] or "未知服务商"
        item = {
            "id": r["id"],
            "original_filename": r["original_filename"],
            "display_name": r["display_name"],
            "category": r["category"],
            "industry": r["industry"] or "",
            "status": r["status"],
            "progress": r["progress"],
            "char_count": r["char_count"],
            "created_at": r["created_at"],
            "user_id": r["user_id"],
        }
        if pname not in grouped:
            grouped[pname] = []
        grouped[pname].append(item)
    return [{"provider_name": p, "files": grouped[p]} for p in sorted(grouped.keys())]


@app.get("/api/admin/invitation-codes")
async def admin_list_codes(user: dict = Depends(require_auth)):
    """受邀码列表"""
    conn = get_db()
    cursor = conn.cursor()
    # 先查每 provider_name 注册了多少人和用户名列表
    cursor.execute("SELECT provider_name, COUNT(*) as cnt, GROUP_CONCAT(username) as usernames FROM users GROUP BY provider_name")
    provider_info = {row["provider_name"]: {"cnt": row["cnt"], "usernames": row["usernames"] or ""} for row in cursor.fetchall()}

    # 修复历史 NULL created_at：用该 provider 下最早用户创建时间
    cursor.execute("UPDATE invitation_codes SET created_at = (SELECT MIN(u.created_at) FROM users u WHERE u.provider_name = invitation_codes.provider_name) WHERE created_at IS NULL")

    cursor.execute("""
        SELECT ic.id, ic.code, ic.provider_name, ic.used, ic.max_users, ic.created_at,
               ic.used_by,
               u.username as used_by_username
        FROM invitation_codes ic
        LEFT JOIN users u ON ic.used_by = u.id
        ORDER BY ic.id DESC
    """)
    rows = cursor.fetchall()
    conn.close()
    return [{"id": r["id"], "code": r["code"], "provider_name": r["provider_name"],
             "used": r["used"],
             "registered_count": provider_info.get(r["provider_name"], {}).get("cnt", 0),
             "registered_users": provider_info.get(r["provider_name"], {}).get("usernames", "") or "-",
             "max_users": r["max_users"] or 1,
             "created_at": r["created_at"],
             "used_by_username": r["used_by_username"] or "-"} for r in rows]

@app.post("/api/admin/invitation-codes")
async def admin_create_code(body: dict, user: dict = Depends(require_auth)):
    """创建受邀码"""
    provider_name = body.get("provider_name", "").strip()
    max_users = int(body.get("max_users", 1))

    if not provider_name:
        raise HTTPException(status_code=400, detail="服务商名称不能为空")

    # 生成随机码
    chars = string.ascii_uppercase + string.digits
    code = ''.join(secrets.choice(chars) for _ in range(8))
    # 确保不重复
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM invitation_codes WHERE code = ?", (code,))
    while cursor.fetchone():
        code = ''.join(secrets.choice(chars) for _ in range(8))
        cursor.execute("SELECT id FROM invitation_codes WHERE code = ?", (code,))

    cursor.execute(
        "INSERT INTO invitation_codes (code, provider_name, max_users) VALUES (?, ?, ?)",
        (code, provider_name, max_users)
    )
    conn.commit()
    code_id = cursor.lastrowid
    conn.close()
    return {"id": code_id, "code": code, "provider_name": provider_name, "max_users": max_users}

@app.delete("/api/admin/invitation-codes/{code_id}")
async def admin_delete_code(code_id: int, user: dict = Depends(require_auth)):
    """删除受邀码"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM invitation_codes WHERE id = ?", (code_id,))
    conn.commit()
    conn.close()
    return {"success": True}

@app.post("/api/admin/invitation-codes/{code_id}/reset")
async def admin_reset_code(code_id: int, user: dict = Depends(require_auth)):
    """重置受邀码（生成新码）"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT provider_name, max_users FROM invitation_codes WHERE id = ?", (code_id,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="受邀码不存在")

    chars = string.ascii_uppercase + string.digits
    new_code = ''.join(secrets.choice(chars) for _ in range(8))
    cursor.execute(
        "UPDATE invitation_codes SET code = ?, used = 0, used_by = NULL WHERE id = ?",
        (new_code, code_id)
    )
    conn.commit()
    conn.close()
    return {"code": new_code}

# ==================== 知识库文件管理 ====================

@app.post("/api/kb/upload")
async def upload_kb_file(
    file: UploadFile = File(...),
    display_name: str = Form(...),
    category: str = Form(...),
    industry: str = Form(""),
    user: dict = Depends(require_auth)
):
    """上传知识库文件"""
    import uuid
    import shutil

    file_id = str(uuid.uuid4())
    user_kb_dir = KB_DIR / str(user["user_id"])
    user_kb_dir.mkdir(parents=True, exist_ok=True)

    # 保存文件
    ext = Path(file.filename).suffix.lower()
    safe_filename = f"{file_id}{ext}"
    filepath = user_kb_dir / safe_filename

    with open(filepath, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # 提取文本内容
    content = ""
    try:
        if ext == ".txt" or ext == ".md":
            content = filepath.read_text(encoding="utf-8")
        elif ext == ".docx":
            from docx import Document
            doc = Document(filepath)
            content = "\n".join([p.text for p in doc.paragraphs])
        elif ext == ".pdf":
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                content = "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif ext in [".xls", ".xlsx"]:
            import openpyxl
            wb = openpyxl.load_workbook(filepath)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    content += " ".join([str(c) for c in row if c]) + "\n"
        elif ext == ".csv":
            import csv
            with open(filepath, encoding="utf-8") as cf:
                reader = csv.reader(cf)
                for row in reader:
                    content += " ".join([str(c) for c in row if c]) + "\n"
        elif ext == ".doc":
            # 先尝试用 antiword 提取文本
            import subprocess
            result = subprocess.run(["antiword", str(filepath)], capture_output=True, text=True)
            if result.returncode == 0:
                content = result.stdout
    except Exception as e:
        content = f"[解析失败: {str(e)}]"

    char_count = len(content)

    # 保存到数据库
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO kb_files (id, user_id, original_filename, display_name, category, industry, filepath, status, progress, char_count)
        VALUES (?, ?, ?, ?, ?, ?, ?, 'completed', 100, ?)
    """, (file_id, user["user_id"], file.filename, display_name, category, industry, str(filepath), char_count))
    conn.commit()
    conn.close()

    return {"id": file_id, "status": "completed", "progress": 100, "char_count": char_count}


@app.get("/api/kb/files")
async def list_kb_files(category: str = "", user: dict = Depends(require_auth)):
    """获取知识库文件列表"""
    conn = get_db()
    cursor = conn.cursor()

    if category:
        cursor.execute("""
            SELECT * FROM kb_files WHERE user_id = ? AND category = ? ORDER BY created_at DESC
        """, (user["user_id"], category))
    else:
        cursor.execute("""
            SELECT * FROM kb_files WHERE user_id = ? ORDER BY created_at DESC
        """, (user["user_id"],))

    files = [dict(row) for row in cursor.fetchall()]

    # 计算完善度
    total = len(files)
    completion = min(int(total / 10 * 100), 100) if total > 0 else 0

    # 计算分类统计
    stats = {"case": 0, "template": 0, "knowledge": 0, "qa": 0, "sales": 0}
    cursor.execute("SELECT category, COUNT(*) as cnt FROM kb_files WHERE user_id = ? GROUP BY category", (user["user_id"],))
    for row in cursor.fetchall():
        if row["category"] in stats:
            stats[row["category"]] = row["cnt"]

    conn.close()
    return {"files": files, "total": total, "completion": completion, "stats": stats}


@app.get("/api/kb/enhancement")
async def get_kb_enhancement(user: dict = Depends(require_auth)):
    """获取知识库增强效果"""
    conn = get_db()
    cursor = conn.cursor()

    # 获取所有文件内容
    cursor.execute("SELECT char_count FROM kb_files WHERE user_id = ? AND status = 'completed'", (user["user_id"],))
    rows = cursor.fetchall()
    total_chars = sum(row["char_count"] for row in rows)

    if total_chars == 0:
        conn.close()
        return {"examples": []}

    # 根据内容生成3个示例
    cursor.execute("""
        SELECT original_filename, category, char_count FROM kb_files
        WHERE user_id = ? AND status = 'completed' ORDER BY char_count DESC LIMIT 5
    """, (user["user_id"],))
    kb_files = [dict(row) for row in cursor.fetchall()]
    conn.close()

    if not kb_files:
        return {"examples": []}

    # 调用AI生成示例
    kb_summary = "\n".join([
        f"- [{f['category']}] {f['original_filename']} ({f['char_count']}字)"
        for f in kb_files
    ])

    system_prompt = """你是一个售前知识库助手。根据用户上传的知识库内容，生成3个"客户问题-通用回答-增强回答"的示例。

要求：
1. 客户问题：模拟真实客户会问的问题
2. 通用回答：没有知识库时的泛泛回答
3. 增强回答：基于知识库内容的有说服力的回答，不要提到具体文件名，用"基于我们的服务经验"等表述
4. 三个示例要覆盖不同方面：行业案例、项目经验、报价方法等

输出JSON格式：
{
  "examples": [
    {"question": "问题1", "default_answer": "通用回答", "enhanced_answer": "增强回答"},
    {"question": "问题2", "default_answer": "通用回答", "enhanced_answer": "增强回答"},
    {"question": "问题3", "default_answer": "通用回答", "enhanced_answer": "增强回答"}
  ]
}"""

    user_prompt = f"知识库内容概览：\n{kb_summary}\n\n根据以上内容，生成3个有说服力的售前示例。"

    ai_result = call_minimax(system_prompt, user_prompt)
    result = ai_result["content"]

    # 解析JSON结果
    import json
    import re
    try:
        # 尝试提取JSON
        match = re.search(r'\{[\s\S]*\}', result)
        if match:
            data = json.loads(match.group())
            return data
    except:
        pass

    return {"examples": [
        {"question": "做过我们行业案例吗？", "default_answer": "有过一些相关案例", "enhanced_answer": "已服务过20+同行业客户，涵盖制造、零售等多个领域"},
        {"question": "项目一般多久完成？", "default_answer": "通常1-3个月", "enhanced_answer": "标准项目45天完成，最快可压缩至30天"},
        {"question": "怎么收费？", "default_answer": "根据项目复杂度定价", "enhanced_answer": "采用基础服务费+模块费+实施费的透明定价模式"}
    ]}


@app.delete("/api/kb/files/{file_id}")
async def delete_kb_file(file_id: str, user: dict = Depends(require_auth)):
    """删除知识库文件"""
    conn = get_db()
    cursor = conn.cursor()

    # 获取文件路径
    cursor.execute("SELECT filepath FROM kb_files WHERE id = ? AND user_id = ?", (file_id, user["user_id"]))
    row = cursor.fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="文件不存在")

    filepath = row["filepath"]

    # 删除文件
    import os
    if os.path.exists(filepath):
        os.remove(filepath)

    # 删除数据库记录
    cursor.execute("DELETE FROM kb_files WHERE id = ? AND user_id = ?", (file_id, user["user_id"]))
    conn.commit()

    # 重新计算完善度
    cursor.execute("SELECT COUNT(*) FROM kb_files WHERE user_id = ?", (user["user_id"],))
    total = cursor.fetchone()[0]
    completion = min(int(total / 10 * 100), 100) if total > 0 else 0

    conn.close()
    return {"success": True, "completion": completion}


# ==================== 客户管理 ====================

@app.get("/api/clients")
async def list_clients(user: dict = Depends(require_auth)):
    """获取客户列表（不含大字段）"""
    conn = get_db()
    cursor = conn.cursor()
    # 只查列表页需要的字段，避免返回巨大的 uploaded_files / transcript / step4_report
    # 测试账号 devuser 可以查看所有客户
    if user.get("is_test_user"):
        cursor.execute("""
            SELECT id, user_id, name, industry, initial_demand, status,
                   step1_result, step2_report, step2_todo, step2_schema,
                   step3_summary, uploaded_files, transcript,
                   step4_report, step4_presales, step4_technical,
                   step4_presales_versions, step4_technical_versions,
                   step4_input_draft, step5_schema, step5_agent_suggestions,
                   created_at, updated_at, demo_url, _completed, _saved,
                   0 AS note_count
            FROM clients ORDER BY updated_at DESC
        """)
    else:
        cursor.execute("""
            SELECT id, user_id, name, industry, initial_demand, status,
                   step1_result, step2_report, step2_todo, step2_schema,
                   step3_summary, uploaded_files, transcript,
                   step4_report, step4_presales, step4_technical,
                   step4_presales_versions, step4_technical_versions,
                   step4_input_draft, step5_schema, step5_agent_suggestions,
                   created_at, updated_at, demo_url, _completed, _saved,
                   0 AS note_count
            FROM clients WHERE user_id = ? ORDER BY updated_at DESC
        """, (user["user_id"],))
    cols = ["id", "user_id", "name", "industry", "initial_demand", "status",
            "step1_result", "step2_report", "step2_todo", "step2_schema",
            "step3_summary", "uploaded_files", "transcript",
            "step4_report", "step4_presales", "step4_technical",
            "step4_presales_versions", "step4_technical_versions",
            "step4_input_draft", "step5_schema", "step5_agent_suggestions",
            "created_at", "updated_at", "demo_url", "_completed", "_saved", "note_count"]
    clients = []
    for row in cursor.fetchall():
        d = dict(zip(cols, row))
        d["is_completed"] = int(d.pop("_completed", 0) or 0)
        d["is_saved"] = int(d.pop("_saved", 0) or 0)
        clients.append(d)
    conn.close()
    return clients

@app.post("/api/clients")
async def create_client(body: dict, user: dict = Depends(require_auth)):
    """创建客户"""
    name = body.get("name", "")
    industry = body.get("industry", "")
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO clients (user_id, name, industry) VALUES (?, ?, ?)",
        (user["user_id"], name, industry)
    )
    conn.commit()
    client_id = str(cursor.lastrowid)  # 统一返回字符串格式，与前端 ID 格式一致
    conn.close()
    return {"success": True, "id": client_id}

@app.get("/api/clients/{client_id}")
async def get_client(client_id: str, user: dict = Depends(require_auth)):
    """获取客户详情"""
    conn = get_db()
    cursor = conn.cursor()
    # 测试账号 devuser 可以查看所有客户
    if user.get("is_test_user"):
        cursor.execute("SELECT * FROM clients WHERE id = ?", (client_id,))
    else:
        cursor.execute("SELECT * FROM clients WHERE id = ? AND user_id = ?", (client_id, user["user_id"]))
    row = cursor.fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="客户不存在")

    result = dict(row)
    # FastAPI JSON 序列化会丢弃下划线开头的 key，强制 pop 后重命名
    _saved_v = result.pop("_saved", 0)
    result["is_saved"] = int(_saved_v) if _saved_v else 0
    _completed_v = result.pop("_completed", 0)
    result["is_completed"] = int(_completed_v) if _completed_v else 0
    # Parse JSON fields back to objects
    for field in ("step1_result", "step2_report", "step2_todo", "step2_schema", "step3_summary", "uploaded_files", "transcript", "step4_report", "step4_presales", "step4_technical", "step4_presales_versions", "step4_technical_versions", "step5_schema", "step5_agent_suggestions"):
        if result.get(field) and isinstance(result[field], str):
            try:
                result[field] = json.loads(result[field])
            except:
                pass
    return result

@app.put("/api/clients/{client_id}")
async def update_client(client_id: str, data: dict, user: dict = Depends(require_auth)):
    """更新客户"""
    conn = get_db()
    cursor = conn.cursor()

    # 检查所有权（测试账号 devuser 可以更新任何客户）
    if user.get("is_test_user"):
        cursor.execute("SELECT id FROM clients WHERE id = ?", (client_id,))
    else:
        cursor.execute("SELECT id FROM clients WHERE id = ? AND user_id = ?", (client_id, user["user_id"]))
    if not cursor.fetchone():
        raise HTTPException(status_code=404, detail="客户不存在")

    # 更新字段（前端发 is_completed/is_saved，数据库列名是 _completed/_saved）
    FIELD_MAP = {"is_completed": "_completed", "is_saved": "_saved"}
    allowed_fields = ["name", "industry", "scale", "budget", "tags", "initial_demand", "status", "step1_result", "step2_report", "step2_todo", "step2_schema", "step3_summary", "uploaded_files", "transcript", "step4_report", "step4_presales", "step4_technical", "step4_presales_versions", "step4_technical_versions", "step5_schema", "step5_agent_suggestions", "step4_input_draft", "demo_url", "_wecom_docid", "_wecom_url", "_step1_wecom_docid", "_step1_wecom_url", "_step4_publish_url", "_step4_technical_publish_url", "_notes_wecom_docid", "_notes_wecom_url", "is_completed", "is_saved", "company_type", "main_customers", "possible_focus", "company_intro", "admin_note_step1", "admin_note_step2", "admin_note_step3", "admin_note_step4", "admin_note_step5"]
    updates = []
    values = []
    for field in allowed_fields:
        if field in data:
            db_field = FIELD_MAP.get(field, field)  # 前端名→数据库列名
            updates.append(f"{db_field} = ?")
            val = data[field]
            # JSON fields must be serialized to string for SQLite
            if field in ("step1_result", "step2_report", "step2_todo", "step2_schema", "step3_summary", "uploaded_files", "transcript", "step4_report", "step4_presales", "step4_technical", "step4_presales_versions", "step4_technical_versions", "step5_schema", "step5_agent_suggestions", "step4_input_draft"):
                val = json.dumps(val) if val is not None else ""
            values.append(val)

    if updates:
        updates.append("updated_at = CURRENT_TIMESTAMP")
        values.append(client_id)
        cursor.execute(f"UPDATE clients SET {', '.join(updates)} WHERE id = ?", values)
        conn.commit()

    conn.close()
    return {"success": True}

@app.delete("/api/clients/{client_id}")
async def delete_client(client_id: str, user: dict = Depends(require_auth)):
    """删除客户"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM clients WHERE id = ? AND user_id = ?", (client_id, user["user_id"]))
    conn.commit()
    conn.close()
    return {"success": True}

# ==================== 报告生成 ====================

REPORTS_SYSTEM_PROMPT = """你是一个专业的企业微信智能表格售前方案顾问。根据服务商与客户的沟通记录，生成结构化的需求洞察报告。

## 🔴 输入清洗（先做，不输出过程）
沟通记录常混入三类内容，先剥离，只对①层负责：
① 客户真实表达 → 唯一可进 evidence（原话逐字引用）、唯一可当"客户明确说过"。
② 系统/AI/服务商已回复内容（"已收到你的需求""我已经帮你整理全套表格""SPECIFICATION""表1|字段|…"）→ 只能进 phaseOneScope/phaseTwoScope 的参考，绝不能进 painPoints/confirmedNeeds 的 evidence，绝不能写成客户原话。
③ 噪音（手机号、"方便电话咨询""这两天在考试"）→ 忽略。
另：salesProfile.scale 若沟通记录正文提到人数（如"员工150人"），据此填写。

严格输出以下 JSON（直接输出，不要开场白，不要 markdown 代码块，不要注释）：

{
  "customerCurrentState": "客户当前业务状态、规模、已有系统、现状做法（50字以内）",
  "salesProfile": {"scale":"人数规模，没有填待确认","budgetSignal":"预算信号/报价敏感度/紧迫性，没有填待确认","currentSystems":"目前在用的系统/工具，没有填待确认"},
  "painPoints": [{"title":"痛点标题","description":"详细描述","evidence":"客户原话引用（逐字，不得转述）"}],
  "confirmedNeeds": [{"title":"需求名称","description":"客户明确表达的具体需求","evidence":"客户原话引用"}],
  "involvedRoles": [{"role":"角色名称","responsibility":"该角色在项目中的职责"}],
  "currentProcess": "当前业务流程怎么跑（30字以内）",
  "expectedOutcome": "客户期望效果（30字以内）",
  "coreTablesExpected": ["按客户行业+需求列出本场景【必然需要】的核心表名（只列表名，如 食品生产→库存管理表/出入库记录表/生产批次表）。这是给下游建表环节的清单锚点，防止漏表。"],
  "phaseOneScope": [{"item":"一期交付项","description":"具体做什么"}],
  "phaseTwoScope": [{"item":"二期评估项","description":"为什么放二期"}],
  "pendingQuestions": [{"question":"待确认问题","whyAsk":"为什么要确认","impactIfUnknown":"不知道会影响什么"}]
}

## 填写规则
1. 只填沟通记录中客户**明确说过**的内容；没说的字段填空数组或空字符串，**不要硬编**。
2. painPoints.evidence 必须是客户**原话逐字引用**，不能是你的转述。
3. 一期只放客户明确要做、且企业微信可轻量实现的。
4. ERP对接 / AI自动判断 / 历史数据清洗 → 默认二期。
5. pendingQuestions 要具体，针对客户话语中的模糊点和遗漏，不要泛泛而问。
6. 【无重复】各数组内条目互不相同，禁止复制凑数。
7. 🆕 salesProfile 三项尽力从沟通记录提取，没有就填"待确认"，绝不臆造。
8. 🆕 coreTablesExpected 必须齐全：按客户行业列出该场景公认的核心表，即便客户没逐个点名。宁可多列，漏掉核心表视为不合格。
9. 🔴【禁止收敛性损耗·通用】客户在①层里分别表达过的每个独立业务动作/诉求，必须在 painPoints 或 confirmedNeeds 中各自成条，禁止因"业务上相关""看起来是一类""下游会合流"就在报告层揉成一条。判定：两个诉求哪怕有先后/关联关系，只要客户分别表达过，就分别成条。宁可多列，不许吞掉。（此为通用保真原则，不针对任何具体维度名。）

## 输出前自检
- 合法 JSON、无注释、无代码块；各数组无重复条目；evidence 均为原话。
- 🆕 salesProfile 三项已填或标"待确认"；coreTablesExpected 覆盖该行业公认核心表。
- 🔴 客户分别表达过的独立诉求均各自成条，无合并吞并。

"""

DEMO_SYSTEM_PROMPT = """你是一个专业的企业微信智能表格架构师。根据客户需求设计智能表格 Demo 结构。

## 输出格式（严格JSON，不要markdown代码块包裹，直接输出JSON）

{"doc_name":"表格名称","sheets":[{"sheet_name":"子表名称","fields":[{"field_title":"字段名","field_type":"文本"}],"sample_records":[{"字段名":"示例值"}]}]}

## field_type 只能取以下中文名
文本 / 多行文本 / 数字 / 单选 / 多选 / 日期时间 / 金额 / 百分比 / 进度 / 手机 / 邮箱 / 链接 / 勾选 / 人员 / 附件 / 图片 / 关联记录 / 公式 / 自动编号

## 设计原则
1. 有字段经验池时，按客户实际需求挑选组合，不照搬全部。
2. 客户没提到的需求对应的表可以不给。
3. 客户提了经验池里没有的需求，自行补合理字段。
4. 子表数量按客户实际复杂度定。
5. 🆕 **字段数量贴合真实业务，不做"轻量"删减**：核心业务表至少 12-20 个字段，覆盖该业务对象的完整信息（基础信息+状态+时间+责任人+金额/数量+备注等）；核心表字段少于 10 个视为不合格。仅字典表/辅助表可少（6-10 个）。宁可字段全，不可为简洁砍字段。
6. 🆕 **覆盖核心表清单**：需求报告 coreTablesExpected 里列出的每一张核心表都必须生成，缺表视为不合格。
7. 每个子表 3-5 条示例数据，**每条内容不同**、真实可信、贴合行业，禁止复制同一条。
8. 字段命名专业、贴合行业术语。
9. 一张表聚焦一个业务对象。

## 输出前自检
- 合法 JSON、无注释；field_type 全部来自上述清单；示例数据无重复行。
- 🆕 核心业务表字段数 ≥12；coreTablesExpected 中的表张张都在，无遗漏。"""


@app.post("/api/reports/generate")
async def generate_report(body: dict, user: dict = Depends(require_auth)):
    """生成需求分析报告或Demo方案"""
    transcript = body.get("transcript", "")
    industry = body.get("industry", "")
    output_type = body.get("output_type", "report")  # report or schema

    if not industry:
        raise HTTPException(status_code=400, detail="industry is required")

    # 构建知识库上下文
    kb = load_global_knowledge()
    query = f"{industry} {transcript[:200]}".strip()

    # 匹配行业知识
    industry_lower = industry.lower()
    industry_text = ""
    for key, data in kb["industries"].items():
        tags = [t.lower() for t in data.get("tags", [])]
        name = data.get("industry_name", "").lower()
        if industry_lower in name or any(industry_lower in t or t in industry_lower for t in tags):
            content = data.get("content", "")
            industry_text = content[:2000] if len(content) > 2000 else content
            break

    # 匹配案例
    query_lower = query.lower()
    scored_cases = []
    for case in kb["cases"]:
        score = 0
        meta = case.get("meta", {})
        case_industry = meta.get("industry", "").lower()
        case_scene = meta.get("scene", "").lower()
        if case_industry in query_lower or query_lower in case_industry:
            score += 5
        if case_scene in query_lower:
            score += 3
        if score > 0:
            scored_cases.append((score, case))
    scored_cases.sort(key=lambda x: x[0], reverse=True)
    matched_cases = scored_cases[:2]

    case_context = ""
    for score, case in matched_cases:
        meta = case.get("meta", {})
        solution = case.get("solution", {})
        case_context += f"【案例：{meta.get('industry', '')} - {meta.get('scene', '')}】\n"
        case_context += f"  架构：{solution.get('architecture', '')}\n"
        tables = solution.get("tables", [])
        if tables:
            for t in tables[:5]:
                case_context += f"  - {t.get('table_name', '')}\n"
        case_context += "\n"

    kb_context = ""
    if industry_text:
        kb_context += f"## 行业知识\n{industry_text}\n\n"
    if case_context:
        kb_context += f"## 相关交付案例\n{case_context}\n\n"

    if output_type == "schema":
        # 生成Demo方案
        system_prompt = DEMO_SYSTEM_PROMPT
        user_prompt = f"{kb_context}## 客户沟通记录\n\n{transcript}\n\n请基于以上沟通记录设计智能表格的表和字段结构，输出JSON。"
        max_tokens = 2000
    else:
        # 生成需求报告
        system_prompt = REPORTS_SYSTEM_PROMPT
        user_prompt = f"{kb_context}## 客户沟通记录\n\n{transcript}\n\n请基于以上沟通记录，生成结构化的需求分析报告。"
        max_tokens = 4000

    ai_result = call_deepseek(system_prompt, user_prompt, max_tokens=max_tokens)
    result = ai_result["content"]

    # 解析JSON（如果是schema）
    demo_json = None
    if output_type == "schema":
        import re, json
        try:
            json_match = re.search(r'```json\s*(.*?)\s*```', result, re.DOTALL)
            if json_match:
                demo_json = json.loads(json_match.group(1))
            else:
                demo_json = json.loads(result)
        except:
            demo_json = None

    # 解析报告 JSON（report类型，输出格式已经是JSON）
    parsed_report = None
    if output_type != "schema":
        try:
            parsed_report = json.loads(result)
        except Exception:
            # 降级：尝试从 markdown 中提取
            try:
                import re as re2
                m = re2.search(r'\{.*\}', result, re2.DOTALL)
                if m:
                    parsed_report = json.loads(m.group())
            except Exception:
                parsed_report = None
        if not parsed_report:
            parsed_report = {"_raw": result}  # 降级为含原始文本的对象

    return {
        "summary": parsed_report,  # 前端读取 data.summary
        "demo_json": demo_json,
        "context_used": {
            "industry_matched": bool(industry_text),
            "cases_matched": bool(case_context)
        }
    }

# ==================== Step4 方案生成（Prompt 3+4+5）====================

# Prompt 3: 需求结构化 → 生成 requirementSolutionData
STEP4_REQUIREMENT_PROMPT = """你是「企业微信服务商需求调研助手」中的需求结构化引擎。

你的任务不是写 Word、不是写 HTML、不是写建表 Prompt。

你的任务是把 Step1/2/3、用户编辑后的 Step4 输入，知识库、xlsx 交付物，整理成一个稳定的中间结构 requirementSolutionData。后续 Word / HTML / 建表 三个产物都基于它生成——所以你的准确性决定后面三个产物的成败。

====================
【输入信息】
====================

客户基础信息：
- 客户名称：{customer_name}
- 行业：{industry}
- 规模：{scale}
- 初始需求：{initial_demand}

Step1 客户画像：
{company_background}

Step1 行业痛点 / AI智搜辅助信息：
{pain_points}

Step2 信息缺口：
{gaps}

Step2 调研问题清单：
{must_ask}

Step3 原始沟通记录全文：
{transcript}

Step3 AI 摘要：
{step3_summary}

用户在 Step4「方案输入确认」中编辑并保存后的内容：
{step4_input_draft}

知识库匹配结果：
{kb_match_result}

智能表格交付物 xlsx 的 sheet 名和字段摘要：
{xlsx_sheet_summary}

服务商对客户需求总结：
{service_provider_summary}

====================
【材料使用优先级】（高→低）
====================
1. step4_input_draft（用户编辑）
2. step3_summary（Step3 AI 摘要）
3. transcript 中客户明确表达的内容
4. service_provider_summary
5. xlsx_sheet_summary
6. Step1 / Step2 背景
7. 知识库 / 行业模板 / 历史方案

硬规则：
- mainScenario 必须优先取自 step4_input_draft 或 step3_summary，**不得被 step1 industry 覆盖**。例："设计/景观建筑-跨国多区域项目管理"不得识别成"家居定制装修"。
- xlsx 的每个 sheet 与字段摘要必须完整进入 smartTableSpec.confirmedTables，**不得遗漏**。
- 🆕 **字段饱满**：smartTableSpec.fieldsByTable 中，核心业务表字段数 12-20 个，覆盖该对象完整信息（基础信息+状态+时间+责任人+金额/数量+备注等）；核心表字段 <10 视为过薄，需补齐。仅字典/辅助表可少。
- 🆕 **覆盖核心表清单**：若上游 Step3 报告有 coreTablesExpected，其列出的核心表必须张张出现在 confirmedTables 或 suggestedTables，缺表需在 warnings 说明原因。
- step3_summary 字段为空时，从 transcript 推理提取补全（并按下方三态标注为"推断"）。
- 知识库只能补充，不得覆盖客户事实。客户没说过的，不得写成"客户已确认"。
- 多轮沟通后收敛了范围，以最后范围为准。
- 客户只是"提到" AI / ERP / 系统对接 / 复杂财务核算 / 机器人自动填报 → 默认二期评估，不得写入一期，更不能标 P0。
- 沟通记录出现、但没进 xlsx 或服务商总结的扩展需求，不默认进一期。
- **严禁写入一期 P0**：AI 智能填报、机器人自动写表、OA/API 对接、复杂多区域独立阶段/财务追踪。

====================
【第 0 步 · 输入三层剥离】（最先做，不输出过程）
====================
transcript / initial_demand 常混入三类内容（实测 16.3% 混AI回复、40.2% 夹手机号），先剥离：
① 客户真实表达 → 可进 confirmedByCustomer、requirements(confirmedStatus="客户已确认")、painPoints.evidence。
② 系统/服务商/AI 已回复内容（"已收到你的需求""我已经帮你整理全套表格""SPECIFICATION""表1|字段|字段…"）
   → 进 inferredByAI / suggestedTables，source 标"AI预生成待确认"；严禁标"客户已确认"，严禁进 confirmedTables。
③ 噪音（手机号、"方便电话咨询""这两天在考试"）→ 忽略，不进业务字段。

====================
【第 0.5 步 · 数据充分性前置检查】（硬规则）
====================
检查以下三个核心字段能否从材料中确定：客户名称，行业，主场景（mainScenario）。
- 若三者中任意一项在所有材料里都找不到明确依据 → 判定材料不足。
  此时**不要瞎猜、不要硬编**，直接只输出如下 JSON 并结束：
  {"meta":{"dataInsufficient":true},"message":"客户核心信息缺失（客户名/行业/主场景至少一项无依据），请补齐 Step3 沟通记录或 Step4 输入后重跑。缺失项：<列出具体缺哪项>"}
- 若三者齐全 → 继续正常生成。

====================
【第 0.6 步 · 信息保真审计】🔴（决定方案完整性，最容易翻车的一步；抽象原则非样本规则）
====================
上游 Step1 已产出 infoUnits（信息单元清单，每条含 uid / label / kind / sourceQuote / raisedByCustomer / feasibility）。若上游未提供，你必须先从①层客户真实表达里重建一份（规则同 Step1：可独立描述即立条，禁止合并）。

⚡ **先判档位**：raisedByCustomer=true 的单元数 ≤2 或走澄清模式 → 只需保证不发生 L4 来源污染，跳过下面的全量审计，直接生成即可（不啰嗦）。≥3 → 执行下面完整的 4 类损耗审计。

**逐单元执行保真审计**，把每个 infoUnit 落到本 JSON 的产物字段，并防住 4 类损耗：
- **防 L1 收敛性损耗**：客户**分别**提到的两个单元，必须各自有独立产物（requirement/painPoint/module/processNode 之一），不许因"业务相关""下游会合流"就合并成一条。
- **防 L2 降格性损耗**：raisedByCustomer=true 的单元，priority 不得无理由降到 P2，sourceType 不得从 explicit 改写成 derived。要降必须在 note 写降级理由。
- **防 L3 静默丢弃**：企微做不了的单元，进 scope.phaseTwo 或 scope.notRecommended（标 phase + reasonForPhase），**绝不能因为做不了就从产物里消失**。
- **防 L4 来源污染**：AI/服务商已回复内容不得冒充客户单元进 confirmedByCustomer。

**回填 fidelity（保真审计结果）**：在输出的 fidelityAudit 里，逐个 uid 记录最终命运：
- kept（保留为独立产物，正常）；merged（被合并——需写并进了谁+理由，且仅当客户本就表述为一体时才允许）；downgraded（降级——需理由）；dropped（丢弃——**直接判不合格，回去补产物**）。
- 任何 raisedByCustomer=true 的单元出现 dropped，或无理由的 merged/downgraded → 不许输出，回去修。

====================
【第 0.7 步 · 售前落点承接】🆕（治"场景错位：demo好看但上门要大改"）
====================
Step1 可能产出 onsiteChecklist（售前阶段答不准、留到签约后现场逐字段核对的执行细节）。承接规则：
- onsiteChecklist 里的每一条，进 openQuestions，并标 stage="落地阶段确认"，priority 视情"中/低"。
- 🔴 **绝不把这些没问准的执行细节，在方案里硬编成"客户已确认"的具体值/字段规则**。这类细节在 requirements/fieldsByTable 里若要体现，confirmedStatus 一律标 "待确认"，rule/取值标 "⚠️ 待落地确认"。
- 原因：售前视频会议没有一线操作岗在场，字段级/流程步骤级细节此刻问不准。宁可显式留白、标清"落地阶段确认"，也不假装已确认。
- 这与信息保真不冲突：客户**说过**的诉求（infoUnit）必须全留（保真）；客户**没说准**的执行细节留白待确认（落点）。两者都反对"AI 凭空替客户做主"。

====================
【输出要求】
====================

输出严格 JSON（能被 JSON.parse 解析）：无注释、无尾逗号、无 markdown 代码块、无解释文字。结构如下：

{
  "meta": {"dataInsufficient": false, "customerName":"","industry":"","companyScale":"","mainScenario":"","secondaryScenarios":[],"serviceProvider":"","outputDate":"","version":"v1"},
  "infoUnits": [{"uid":"U1","label":"（用客户自己的话，不贴标签名）","kind":"维度/痛点/明确诉求/硬约束","raisedByCustomer":true,"sourceQuote":"（①层客户原话≤30字）","status":"covered/pending/unclear/derived","priority":"高优先级(P0)/中优先级(P1)/低优先级(P2)","feasibility":"可原生实现/二期/需外部产品","note":""}],
  "fidelityAudit": [{"uid":"U1","fidelity":"kept/merged/downgraded/dropped","landedOn":["requirements[?]","moduleRecommendation[?]"],"phase":"一期/二期评估/暂不建议","reason":"（非kept必填理由）"}],
  "sourceTrace": {"confirmedByCustomer":[],"fromStep3Summary":[],"fromUserEditedInput":[],"fromServiceProviderSummary":[],"fromXlsxOrDeliveryFile":[],"fromKnowledgeBase":[],"inferredByAI":[],"pendingConfirmation":[]},
  "customerFacts": {
    "customerCurrentState":"",
    "existingTools":[],
    "currentProcess":[{"stepName":"","role":"","currentMethod":"","problem":"","evidenceQuote":""}],
    "involvedRoles":[],
    "explicitNeeds":[]
  },
  "painPoints": [{"title":"","description":"","businessImpact":"","evidence":"","priority":"P0/P1/P2","uid":""}],
  "requirements": [{"requirementName":"","customerExpression":"","businessTranslation":"","priority":"P0/P1/P2","phase":"一期/二期评估/暂不建议","reasonForPhase":"","confirmedStatus":"客户已确认/用户编辑确认/AI推断/待确认","uid":""}],
  "scope": {
    "phaseOne":[{"item":"","reason":"","deliveryForm":""}],
    "phaseTwo":[{"item":"","reason":"","prerequisites":[]}],
    "notRecommended":[{"item":"","reason":""}]
  },
  "businessProcess": {
    "currentFlow":[],
    "targetFlow":[],
    "processNodes":[{"nodeName":"","responsibleRole":"","input":"","output":"","systemAction":"","reminderNeeded":true}]
  },
  "moduleRecommendation": [{"moduleName":"","moduleType":"智能表格/审批/自动化/权限/看板/机器人AI/系统对接","solvedProblem":"","phase":"一期/二期评估/暂不建议","notes":"","uid":""}],
  "smartTableSpec": {
    "scenarioComplexity":"简单流程型/跨部门协同型/多表主数据型/看板同步型/系统对接型",
    "confirmedTables":[{"tableName":"","tablePurpose":"","source":"xlsx/客户确认/服务商总结/知识库建议","roles":[]}],
    "suggestedTables":[],
    "phaseTwoTables":[],
    "fieldsByTable":[{"tableName":"","fields":[{"fieldName":"","fieldType":"（中文名，见字段类型对照表）","required":true,"rule":"","source":"xlsx/客户确认/知识库建议/AI推断"}]}],
    "relations":[],
    "views":[],
    "automations":[],
    "permissions":[],
    "dashboards":[],
    "warnings":[]
  },
  "openQuestions": [{"question":"","whyAsk":"","impactIfUnknown":"","priority":"高/中/低","stage":"售前待补/落地阶段确认"}]
}

====================
【JSON 完整性强制要求】
====================

**所有字段（除 explicitly optional 标注外）必须完整填写，不得留空字符串、不得留空数组、不得留空对象。**

字段缺失时的处理规则（按优先级）：
1. 有值 → 直接使用
2. 无值但有原材料 → 从原材料推理填入（引用证据到 sourceTrace）
3. 无值且无原材料 → 填入 "待确认" 或合理的占位描述

**严禁以下行为**：
- 字段留空（""）、留空数组（[]）、留空对象（{}）
- 把对象或数组当成字符串写入（如写成 "[object Object]"）
- 仅复制输入材料而不做结构化整合
- 用省略号 "..." 代替完整内容

**JSON 结构完整性自检（输出前必查）**：
- meta 全部字段非空
- confirmedTables 每个 table 有 tableName + tablePurpose + roles（非空数组）
- fieldsByTable 每个 field 有 fieldName + fieldType + required
- requirements 每个 item 有 requirementName + priority + phase + reasonForPhase
- openQuestions 每个 item 的 question + whyAsk + impactIfUnknown 为非空字符串
- 数组类型字段（如 painPoints、requirements）最少有 1 条（无内容填 "待补充" + 说明原因）

## 输出前自检（不通过就重写，不要输出）
1. 【合法 JSON】无注释、无尾逗号、无代码块。
2. 【核心字段一致】customerName/industry/mainScenario 前后不矛盾；材料不足时已按第0步回退。
3. 【无复制行】各数组元素内容互不相同。
4. 【三态到位】缺失字段填 "⚠️ 待确认" 且已登记 openQuestions，没有硬编瞎猜。
5. 【xlsx 全进表】xlsx 的 sheet/字段已全部进入 confirmedTables，无遗漏。
6. 🔴【信息保真】fidelityAudit 中所有 raisedByCustomer=true 的单元 fidelity 无 dropped；merged/downgraded 均有理由；客户分别提的单元没被合并；二期/不可行单元已进 scope 而非消失。（此项在"轻档"可豁免，仅需保证 L4 不污染。）
7. 🆕【售前落点】onsiteChecklist 的执行细节已进 openQuestions 并标 stage="落地阶段确认"，没有被硬编成"客户已确认"的字段值/规则。
8. 🆕【字段饱满】核心业务表字段数 ≥12；coreTablesExpected 的核心表张张都在，缺表已在 warnings 说明。

====================
【范围判断规则】
====================
🔴 判定二期看【能力本身的复杂度】，不看客户把它说得多"基础"。客户说"就要个自动生成讲解视频""就要三端隔离"——只要能力本身属于 AI生成/系统集成/复杂权限，一律二期评估，不因客户口气轻松就放进一期。

一期：客户明确的核心需求 + 痛点强 + 企业微信入口/智能表格/审批/自动化/权限/看板可轻量实现 + 不依赖复杂接口/复杂AI判断/大量历史数据清洗。
二期评估：ERP/OA/CRM/财务对接、数据回写，AI自动判断，复杂财务核算，历史数据清洗，多系统权限联动，机器人自动填报，高级经营分析。
暂不建议：替代完整 ERP/CRM/财务、客户没提但模板里有的模块，强监管实时决策，超出轻量交付边界的复杂系统。

直接输出有效 JSON，不要 markdown 代码块包裹。"""

# Prompt 4: Word 内容生成
STEP4_WORD_PROMPT = """你是一个企业微信智能表格售前方案顾问。请基于以下结构化需求数据，生成《需求确认 & 方案设计表》Word 正文内容。

【requirementSolutionData】
{requirement_data}

请严格按以下 JSON 结构输出（直接输出 JSON，不要任何开场白，不要 markdown 代码块）：

{
  "docTitle": "需求确认 & 方案设计表",
  "subtitle": "企业微信定制开发",
  "version": "v1.0",
  "introNotice": "本文件用于服务商与客户共同确认需求范围、一期边界、智能表格搭建口径、权限与待确认问题。客户未确认内容不得写入一期交付承诺。字段与公式为参考设计，最终以落地阶段与客户核对为准。",

  "customerInfoTable": [
    { "field": "客户名称", "value": "" },
    { "field": "行业", "value": "" },
    { "field": "规模", "value": "" },
    { "field": "办公区域/使用范围", "value": "" },
    { "field": "主场景", "value": "" },
    { "field": "方案口径", "value": "" }
  ],

  "needsOverviewTable": {"moduleCount":0,"featureCount":0,"summary":"如 5大模块 × 23项功能"},
  "companyProfileTable": [{"field":"注册资本/成立年份/统一社会信用代码/团队规模","value":"有据才填，无工商数据则本表整体省略，不编造"}],
  "currentPainTable": [
    {"businessArea":"","currentStateOrPain":""}
  ],
  "capabilityMappingTable": [
    {"need":"","implementWay":"","verdict":"原生可实现/降级可实现/无法实现","alternative":"无法实现时填替代方案，否则留空"}
  ],
  "feasibilityDistributionTable": {"native":"X项(X%)","degraded":"X项(X%)","cannot":"X项(X%)","coverageNote":"覆盖约X%"},
  "scenarioBoundary": {
    "scenarioJudgement":"",
    "phaseOne":[""],
    "phaseTwo":[""],
    "notRecommended":[""]
  },
  "requirementPriorityTable": [
    {"requirement":"","priority":"P0/P1/P2","phase":"一期/二期评估/暂不建议","implementationApproach":""}
  ],
  "moduleDetailTable": [
    {"moduleName":"","subFeature":"","capabilities":"该子功能3条能力点，用、分隔","phase":"一期/二期评估/暂不建议"}
  ],
  "processDesignTable": [
    {"item":"","description":""}
  ],
  "wecomArchitectureTable": [
    {"layer":"企业微信入口层/智能表格数据层/自动化与提醒层/权限与看板层","designDescription":""}
  ],
  "smartTableDeliveryTable": [
    {"tableName":"","type":"主表/业务表/辅助表","purpose":"","roles":"","phaseOne":"是/否"}
  ],
  "dataBlueprintTable": [
    {"tableName":"","purpose":"这张表管什么(一句话)","forRole":"主要使用角色"}
  ],
  "dataBlueprintNote": "本表为数据蓝图，说明将搭建哪几张表；具体字段、类型、公式将在方案确认后的智能表格搭建环节(Step5)展开，并与贵司操作团队共创定准。",
  "dataFlowNote": "一句话说清数据从录入表→汇总表→看板的闭环流向",
  "automationTable": [
    {"ruleName":"","trigger":"","action":"","priority":"P0/P1/P2"}
  ],
  "permissionTable": [
    {"role":"","viewScope":"","operation":"","sensitiveFields":""}
  ],
  "dashboardTable": [
    {"dashboard":"","users":"","metrics":"","filters":""}
  ],
  "valueTable": [
    {"metric":"如 会议决策落地率提升约80%","label":"衡量什么","basis":"有据填出处/无据填预估"}
  ],
  "efficiencyComparisonTable": [
    {"dimension":"如 逾期发现","before":"实施前现状","after":"实施后(无据标预估)"}
  ],
  "landingCollaborationNote": "售前先讲清业务价值与蓝图，字段权限落地共创定准。基于同类项目经验，落地阶段30-50%现场调整正常且负责任——这正是方案能真正跑起来、而非demo好看上线难用的关键。共创步骤：需求共创→字段定准→原型试跑→定稿上线。",
  "dataBoundaryTable": [
    {"dataObject":"","phaseOneMethod":"","phaseTwoEvaluation":"","boundaryNote":""}
  ],
  "implementationPlanTable": [
    {"phase":"","workContent":"","customerCooperation":"","output":""}
  ],
  "optionalModules": {
    "_note":"以下四块默认省略，仅当 Step4 用户追加要求时才输出对应块",
    "competitorComparison": [{"dimension":"","wecom":"","competitor":"竞品侧标参考/估算，不硬编竞品精确价"}],
    "firstCoopGuarantee": [{"guarantee":"","desc":""}],
    "quotationSummary": {"oneTime":"参考区间","annual":"参考区间","paymentPlan":[{"milestone":"","ratio":""}],"disclaimer":"参考区间，以正式报价为准"},
    "wecomCapabilityList": [{"capability":"","note":""}]
  },
  "pendingQuestions": [""],
  "confirmationItems": [
    "一期范围是否按本文件定义执行 □ 确认 □ 调整",
    "字段与权限是否允许按试运行反馈微调 □ 确认 □ 调整",
    "OA/机器人/AI 能力是否作为二期评估 □ 确认 □ 调整",
    "客户签字/盖章：________________"
  ]
}

## 填写规则（v11 · 内容饱满度硬指引）
1. customerInfoTable：value 全部从 requirementSolutionData.meta 与 customerFacts 读取；不得自行发明客户名/行业/场景，材料没有的填 "⚠️ 待确认"。
2. currentPainTable：≥4 行（覆盖客户提到的各业务面），每条痛点描述 ≥20 字，讲清"现状怎么做+卡在哪+什么后果"。
3. capabilityMappingTable：承接 requirementSolutionData，逐项需求判 verdict；"无法实现"必须填 alternative 替代方案，绝不把做不到的写成可实现。
4. scenarioBoundary.scenarioJudgement：必填，写明场景判断。
5. requirementPriorityTable：**严禁**将 AI 能力、机器人自动写表、OA 对接写成 P0；为空时填"[{"requirement": "【待补充】需求列表为空，请返回 Step3 补充沟通记录", "priority": "P0", "phase": "待确认", "implementationApproach": ""}]"
6. moduleDetailTable（做厚的核心，不能为空）：每个一期模块 2-4 个子功能行，capabilities 每格 3 条具体能力点（"员工能做什么+系统会做什么"），不写空话。
7. smartTableDeliveryTable：phaseOne="是"的表必须全部来自 smartTableSpec.confirmedTables 且 phase="一期"；每条必须有 tableName + type + purpose + roles（非空）。
8. dataBlueprintTable（替代原字段表）：只列"要建哪几张表 + 一句话用途 + 主要使用角色"，**禁止列字段名、字段类型、公式**。字段设计属于 Step5 建表环节，售前 Word 不出现。dataBlueprintNote 固定说明字段在 Step5 展开。
9. valueTable：3-5 条量化价值，无客户材料佐证的数字 basis 一律填"预估"，绝不硬编精确承诺值。
10. landingCollaborationNote：固定写清"字段权限在搭建环节共创、30-50% 调整正常且负责任，共创四步"，作为信任状。
11. needsOverviewTable：统计 requirements/modules，给"X大模块×Y项功能"。
12. companyProfileTable：**仅当材料有工商/公司信息时填**；无据则整表省略，绝不编造。
13. feasibilityDistributionTable：把 capabilityMappingTable 汇总成 原生/降级/不可 三档占比 + 覆盖率。
14. efficiencyComparisonTable：3-5 行"实施前 vs 实施后"，实施后无据数字标"预估"。
15. optionalModules：四块默认全部省略；**仅当 requirementSolutionData.optionalModules.enabled 列出时才填**。竞品价格/费用一律标"参考/估算，以正式报价为准"，不硬编。
16. pendingQuestions：≥4 条，每项讲清"问什么+为什么影响交付"。
17. **所有单元格都必须是字符串**，确实无值填 "⚠️ 待确认"，不得出现 [object Object]。

## 三条铁律
A.【每行必须不同】任何一张表里，**不允许出现两行内容完全相同**。
B.【待确认不铺满】"⚠️ 待确认"只能用于**个别真正缺数据的格子**。如果整张表大面积待确认，说明上游数据不足——该表只输出一行"【数据不足】该模块缺少客户材料支撑，请补充后重生成"。
C.【不许自相矛盾】客户名，行业、主场景在全文必须一致。

直接输出有效 JSON，不要 markdown 代码块包裹。"""




# Prompt 5: 可视化 HTML 内容生成
# Prompt 6: 技术路线及报价方案（文档风·11章26表·可下载Word）— V11
STEP4_TECHDOC_PROMPT = """你是一个企业微信定制开发技术方案顾问。请基于结构化需求数据，生成《技术路线及报价方案》完整 HTML 文档（文档风，可下载 Word）。

【requirementSolutionData】
{requirement_data}

请直接输出完整 HTML 文档（不要 JSON，不要 markdown 代码块包裹）。

## 文档视觉基因（锁死，禁止修改）
- 正文字体：font-family:"STKaiti","华文楷体","KaiTi","楷体",serif（正文华文楷体）
- 背景：白色 #ffffff
- 排版：白底、A4 文档感、双线封面、编号章节（一、二…十一）、小节（1.1/1.2）、密集表格（表头深蓝底白字 #1f4e79）
- 封面：双线边框，深蓝顶标 #1f4e79
- 章节标题：左 left-border 5px solid #1f4e79 + 浅蓝背景 #f2f5f9
- 表头：#1f4e79 底 + 白字加粗，偶数行：#f7f9fb
- 右下角固定「⬇ 下载 Word 版」按钮（Blob msword 导出，Word 样式里字体同样设为华文楷体）
- 无营销渐变卡片、无 KPI 大数字、无效果图——这是严肃的技术确认文档

## 文档骨架（11 章 26 表，锁死结构，照此生成）
封面：企业微信定制开发 / 技术路线及报价方案 / 需求确认&方案设计表 / "请勿将未确认需求写入一期交付承诺"
元信息表×2：客户名称|所属行业|服务商|项目场景；方案版本|输出日期|沟通阶段|客户联系人

一、客户基础信息与当前现状
  1.1 客户基础信息确认表（项目|填写内容|备注来源，8行）
  1.2 当前业务运转方式（环节|当前做法|主要问题|会议依据客户原话，≥4行）
  1.3 核心痛点与优先级（痛点编号|描述|业务影响|优先级|对应企微方案，≥4行）
二、场景类型判断与方案边界
  2.1 场景类型判断（判断项|填写内容：场景大类/是否已有源系统/本期定位|交付边界，4行）
  2.2 一期/二期/不建议范围（类型|范围说明|原因，3行）
三、需求理解与优先级确认
  3.1 需求清单（需求项|客户描述|业务影响|优先级|一期二期|企微实现方式，≥5行）
  3.2 客户原话与业务翻译（客户原话|业务翻译|是否已确认，≥4行）
四、业务流程设计
  4.1 当前流程与目标流程（阶段|当前流程|优化后流程|企微动作，≥4行）
  4.2 流程节点确认表（节点序号|名称|操作角色|输入|输出|是否提醒，≥5行）
五、企业微信方案总览
  5.1 能力架构（层级|能力|本项目使用方式，4行）
六、智能表格交付设计【技术核心，字段在此展开】
  6.1 智能表格总览（表名|类型|用途|使用对象|一期必做，≥5行）
  6.2 核心字段设计表（所属表|字段分组|字段名|类型|必填|填写角色|规则说明，逐字段展开，用rowspan按表分组，≥16行）
  6.3 表间关联关系（主表|关联表|关联字段|自动带出汇总|注意事项，≥3行）
七、审批与自动化设计
  7.1 自动化规则表（规则名|触发条件|执行动作|通知对象|优先级，≥4行）
  7.2 审批流程设计（审批名|发起角色|审批人|通过后动作|同步表格，≥3行）
八、权限与数据看板设计
  8.1 权限矩阵（角色|新增|查看范围|可编辑字段|看板权限|敏感字段，覆盖全部角色）
  8.2 数据看板设计（看板名|使用对象|核心指标|筛选维度|是否下钻，≥3行）
九、数据来源、系统对接与交付边界
  9.1 数据来源与接入方式（数据对象|来源方式|一期接入|二期评估，≥3行）
  9.2 交付清单（类别|交付内容|是否包含|备注，≥6行）
十、实施计划、报价口径与变更机制
  10.1 实施计划（阶段|工作内容|客户配合|输出物，≥4行）
  10.2 报价口径建议（费用模块|范围说明|是否本次包含|备注参考区间，≥6行；价格一律标"参考区间，以正式报价为准"）
  10.3 范围变更机制（事项|是否范围内|处理方式，3行）
十一、待客户确认问题与签署
  11.1 待确认问题清单（问题编号|待确认问题|负责人|截止时间|确认结果，≥5行）
  11.2 客户确认（确认事项|确认说明，含盖章签字栏）

## 填写规则（治"内容不对、满屏待确认、残留占位符"）
1. 内容来自 requirementSolutionData（06 产出），客户名/行业/痛点/需求/场景/字段全部取真实数据，不残留 `{{来源}}`/`{{}}` 占位符。
2. ⚠️待确认只用于真缺失（只有客户确实没定的才标）；能从 Step1/Step3 推出的一律填实，禁止满屏待确认。
3. 字段设计要厚：6.2 逐字段展开，核心表 8-13 字段/表，用 rowspan 按表分组，带类型/必填/角色/规则/公式。这是技术方案的价值所在。
4. 报价只给参考区间：10.2 每项标"参考 X-Y 元"，统一注"参考区间，以正式报价为准"，不硬编精确价。
5. 一期/二期边界沿用 06 的 confidence 与 scope，ERP对接等一律二期，不写进一期承诺。
6. 内容饱满度：每个表按行数下限填满，每格写实不空泛。

## 输出前自检（不通过就重写）
1. 合法完整 HTML，华文楷体，含下载 Word 按钮。
2. 11 章 26 表齐全，字段设计（6.2）逐字段展开 ≥16 行。
3. 无 `{{}}` 占位符残留；⚠️待确认只在真缺失处。
4. 报价均为参考区间且带 disclaimer。
5. 视觉为文档风（白底/编号章节/密集表），无营销渐变卡片。

直接输出完整 HTML，不要任何前缀说明。"""






STEP4_HTML_PROMPT = """你是一个企业微信智能表格可视化方案顾问。请基于结构化需求数据，生成客户友好的可视化方案 HTML 内容。

【requirementSolutionData】
{requirement_data}

请按以下 JSON 结构输出（直接输出 JSON，不要任何开场白）：

{
  "pageTitle": "",
  "hero": {
    "title": "",
    "subtitle": "",
    "tags": [],
    "summary": ""
  },
  "customerStageJudgement": {
    "title": "",
    "description": "",
    "keyFacts": []
  },
  "insightSection": {
    "mainInsight": "",
    "painCards": [
      {
        "title": "",
        "description": "",
        "impact": ""
      }
    ]
  },
  "scenarioBreakdown": [
    {
      "scenarioName": "",
      "currentProblem": "",
      "targetState": "",
      "wecomSolution": "",
      "value": ""
    }
  ],
  "architecture": {
    "positioning": "",
    "layers": [
      {
        "layerName": "",
        "capability": "",
        "usage": ""
      }
    ]
  },
  "recommendedModules": [
    {
      "moduleName": "",
      "description": "",
      "phase": "一期/二期评估/暂不建议",
      "value": ""
    }
  ],
  "roadmap": [
    {
      "phaseName": "",
      "workContent": "",
      "customerCooperation": "",
      "output": ""
    }
  ],
  "valuePoints": [
    {
      "title": "",
      "description": ""
    }
  ],
  "pendingQuestions": []
}

内容要求：少字高信息密度、不堆砌字段、突出客户现状和核心痛点、不确定内容写"待确认"或"二期评估"。直接输出 JSON。

## 强化要求
1. **scenarioBreakdown 必须生成 3-5 个具体业务场景**，每个场景必须同时包含：
   - scenarioName：具体业务场景名（如"项目立项审批"/"供应商报价管理"）
   - currentProblem：客户当前面临的具体问题（引用客户原话或痛点描述）
   - targetState：使用企业微信智能表格后的目标状态（具体可衡量）
   - wecomSolution：具体解决方案（智能表格+审批+自动化+机器人AI的组合）
   - value：带给客户的量化或质化价值
2. **architecture.layers 必须生成 4 层**，对应企业微信智能表格标准架构：
   - 企业微信入口层（消息通知/工作台/分享）
   - 智能表格数据层（多表关联/数据录入）
   - 自动化与提醒层（自动化规则/机器人推送）
   - 权限与看板层（角色权限/数据看板）
3. **recommendedModules 最少 2 条**，区分一期/二期
4. **roadmap 最少 2 条**（一期 + 二期各 1 条）
5. **painCards 最少 2 条**，每条 impact 必须量化或具体描述
6. 所有字符串字段不得为空（用"待确认"填充），数组不得为空（至少 1 条占位）

**【强制禁止】**：
- 输出内容必须以 `{` 开头，以 `}` 结尾，不允许在 JSON 前后出现任何其他文字
- 禁止输出任何解释、说明、注释、markdown 代码块包裹
- 如果输出内容被截断导致 JSON 不完整，必须在截断处停止，不要补全或续写

直接输出有效 JSON，不要 markdown 代码块包裹，不要有任何开场白或结尾语。"""

def parse_json_response(raw):
    """解析 MiniMax 返回的 JSON，处理各种异常情况"""
    import re, json
    if not raw:
        return None
    try:
        return json.loads(raw)
    except:
        pass
    # 尝试从 markdown 代码块中提取
    m = re.search(r'```(?:json)?\s*([\s\S]*?)```', raw)
    if m:
        try:
            return json.loads(m.group(1).strip())
        except:
            pass
    # 尝试找 JSON 对象
    m = re.search(r'\{[\s\S]*\}', raw)
    if m:
        try:
            return json.loads(m.group())
        except:
            pass
    return None

def _safe(v):
    """将任意值转为安全字符串，用于质检检测 [object Object]"""
    if v is None or v == "":
        return ""
    if isinstance(v, str):
        return v
    if isinstance(v, (int, float, bool)):
        return str(v)
    if isinstance(v, list):
        return "|||".join(_safe(x) for x in v)
    if isinstance(v, dict):
        for k in ("question", "whyAsk", "impactIfUnknown", "title", "description",
                  "item", "reason", "name", "requirement", "priority",
                  "phase", "tableName", "fieldName"):
            if k in v and v[k]:
                return _safe(v[k])
        return json.dumps(v, ensure_ascii=False)
    return str(v)

def db_get_client(client_id):
    """同步获取客户字典（用于非 async 函数）"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM clients WHERE id = ?", (client_id,))
    row = cursor.fetchone()
    conn.close()
    if not row:
        return None
    d = dict(row)
    # JSON 字段反序列化
    for k in JSON_FIELDS:
        if k in d and d[k]:
            try:
                d[k] = json.loads(d[k])
            except Exception:
                pass
    return d


JSON_FIELDS = {"step1_result", "step2_report", "step2_todo", "step2_schema", "step3_summary",
                "uploaded_files", "transcript", "step4_report", "step4_presales", "step4_technical",
                "step4_presales_versions", "step4_technical_versions", "step5_schema",
                "step5_agent_suggestions", "step4_input_draft", "_completed", "_saved"}

def db_update_client(client_id, updates):
    """同步更新客户字段"""
    if not updates:
        return
    conn = get_db()
    cursor = conn.cursor()
    sets = []
    vals = []
    for k, v in updates.items():
        if k in JSON_FIELDS and v is not None:
            v = json.dumps(v, ensure_ascii=False)
        sets.append(f"{k} = ?")
        vals.append(v if v is not None else "")
    vals.append(client_id)
    cursor.execute(f"UPDATE clients SET {', '.join(sets)}, updated_at = CURRENT_TIMESTAMP WHERE id = ?", vals)
    conn.commit()
    conn.close()


def validate_requirement_doc(word_content, requirement_data=None):
    """质检生成的 Word 内容 JSON，不通过返回 errors"""
    errors = []
    warnings = []

    if not word_content or not isinstance(word_content, dict):
        errors.append("wordContent 为空或非对象")
        return {"pass": False, "errors": errors, "warnings": warnings}

    raw_str = json.dumps(word_content, ensure_ascii=False)

    # 1. 检查 [object Object]
    if "[object Object]" in raw_str:
        errors.append("出现 [object Object]，字段映射错误（对象未展开）")

    # 2. 检查 docTitle
    if not _safe(word_content.get("docTitle", "")):
        errors.append("docTitle 为空")

    # 3. 检查 11 个章节是否存在（容许某些章节为空列表但不允许多段落字符串空值）
    required_sections = [
        "customerInfoTable", "currentPainTable", "scenarioBoundary",
        "requirementPriorityTable", "processDesignTable", "wecomArchitectureTable",
        "smartTableDeliveryTable", "keyFieldsByTable", "automationTable",
        "permissionTable", "dashboardTable",
        "dataBoundaryTable", "implementationPlanTable",
        "pendingQuestions", "confirmationItems"
    ]
    for sec in required_sections:
        if sec not in word_content:
            errors.append(f"缺少章节：{sec}")

    # 4. 检查 requirements 是否为空（来自 requirementData）
    if requirement_data:
        reqs = requirement_data.get("requirements") or []
        if not reqs:
            errors.append("requirements 为空，请补充 Step3 沟通记录后重新生成")
        smart = requirement_data.get("smartTableSpec") or {}
        tables = smart.get("confirmedTables") or []
        if not tables:
            errors.append("confirmedTables 为空，请补充 xlsx 或 Step3 后重新生成")

    # 5. 检查空 table（数组为空，或数组每项全空字符串）
    def check_table(name, arr):
        if not isinstance(arr, list):
            errors.append(f"{name} 不是数组")
            return
        if len(arr) == 0:
            errors.append(f"{name} 为空表格")
            return
        for i, row in enumerate(arr):
            row_str = _safe(row)
            if not row_str.strip(" |"):
                errors.append(f"{name}[{i}] 所有字段为空")

    check_table("customerInfoTable", word_content.get("customerInfoTable"))
    check_table("requirementPriorityTable", word_content.get("requirementPriorityTable"))
    check_table("smartTableDeliveryTable", word_content.get("smartTableDeliveryTable"))

    # 6. 检查行业/场景识别错误
    # 精确匹配错误行业（仅针对"家居装修误识别"场景，不拒正常含"装修"词根的行业如"建筑装饰"）
    industry_wrong_exact = ["家居定制装修", "家居装修", "全屋定制"]
    scenario = _safe(word_content.get("scenarioBoundary", {}).get("scenarioJudgement", ""))
    cust_info = word_content.get("customerInfoTable", [])
    industry_val = ""
    for row in cust_info:
        if _safe(row.get("field", "")) == "行业":
            industry_val = _safe(row.get("value", ""))
            break
    if industry_val in industry_wrong_exact:
        errors.append(f"行业识别错误：'{industry_val}'（应为设计/景观建筑-跨国多区域项目管理）")
    if "家居" in scenario or "家居定制" in scenario:
        errors.append(f"场景判断错误：'{scenario}'")

    # 7. 检查把 AI/机器人/OA/ERP 写成 P0 一期
    p0_phase_one = []
    for row in (word_content.get("requirementPriorityTable") or []):
        priority = _safe(row.get("priority", ""))
        phase = _safe(row.get("phase", ""))
        req_text = _safe(row.get("requirement", ""))
        if priority == "P0" and phase == "一期":
            for keyword in ["AI", "机器人", "OA", "ERP", "系统对接", "智能填报", "自动写表"]:
                if keyword in req_text:
                    p0_phase_one.append(f"P0 一期不能包含：{req_text}")
    p0_phase_one and errors.extend(p0_phase_one)

    # 8. 检查章节内容是否为泛泛段落而非表格
    for sec, has_table in [
        ("requirementPriorityTable", True),
        ("smartTableDeliveryTable", True),
        ("automationTable", True),
        ("permissionTable", True),
    ]:
        arr = word_content.get(sec, [])
        if has_table and isinstance(arr, list) and len(arr) > 0:
            first_row_str = _safe(arr[0])
            if first_row_str.count("|||") < 2 and len(first_row_str) > 200:
                warnings.append(f"{sec} 可能只是泛泛段落，建议改为结构化表格")

    return {"pass": len(errors) == 0, "errors": errors, "warnings": warnings}


# ==================== Step1 调研问题生成 ====================

STEP1_SYSTEM_PROMPT = """你是一名资深的定制开发售前调研顾问，服务于企业微信智能表格 / 低代码定制开发场景。
你的职责：把客户进线时的原始表达，转化为一份"能直接带去和客户开会、并推动成交"的调研材料。

六条铁律：
1.【忠于原文】客户明确说过的，是付费锚点，必须被识别、保真、优先；不得被你的主观判断降级或忽略。
2.【边界优先】必问问题只围绕"客户真实场景"，不堆砌行业泛问题；宁可少问，不可越界。
3.【诚实标注来源】每条痛点、缺口、问题都要标明是"客户明确提的"还是"你推导补的"，绝不把推测伪装成客户原意。
4.【缺料不编造】客户原始表达不足以支撑标准调研时，走澄清模式，只提开放式澄清问题，绝不凭空编造需求或场景。
5.【看人下问】售前进线阶段坐在对面的通常是老板/进线人，一线操作岗多半不在场。同一需求，问管理层和问执行层是两套问法：管理层只讲目标与痛，答不出字段细节；执行层才谈操作与数据。必须先判断本次访谈对象，再决定问题的抽象层级——绝不拿字段级、流程步骤级的细节去问管理层。
6.【售前落点】这是"售前视频会议"，不是"签约后现场逐部门实装调研"。你出的问题必须满足两个条件才放进必问清单：①对面此刻的人（多为决策者）当场答得上；②答案能帮销售判断方向，推动成交。凡是"必须叫上一线操作岗、对着现有表格逐字段逐流程核对"才能答准的执行细节，一律不在售前阶段追问——把它标记为"落地阶段确认"，留到签约后现场再问。宁可框架清楚、细节留白，不可细节堆满、客户当场答不上还答不准。

你只输出 JSON，不输出任何解释，开场白或 markdown 代码块。"""

STEP1_USER_PROMPT = Template("""## 客户基本信息
- 客户名称：${company_name}
- 行业：${industry}
- 规模：${scale}
- 需求标签：${tags}
- 原始需求：${initial_demand}
- AI 补充简介：${company_intro}

## 🔴 输入清洗（最先做，不输出清洗过程）
把【原始需求】在心里剥成三层，只对第①层负责：
① 客户真实表达：客户自己说的诉求/现状/痛点。← 唯一可标 sourceType="explicit" 的来源。
② 系统/服务商/AI 已回复内容：形如"已收到你的需求""我已经帮你整理""SPECIFICATION""全套可复制表格""表1|字段|字段…"等。→ 只能当参考线索，标 "derived"，严禁标 "explicit"，严禁写成"客户已确认"。
③ 噪音：手机号（1xxxxxxxxxx）、"方便电话咨询""这两天在考试""与客户语音沟通情况"等寒暄。→ 忽略，不进任何业务字段。
清洗后：若①层有效业务内容 < 15 字（例：只剩"方便电话咨询？"），直接判为"需求未明确"，走【分支B 澄清模式】。

## 行业/规模兜底
- 若 ${industry} 为空，从①层正文或公司名推断行业，标 sourceType=derived。
- 若 ${scale} 为空，从①层正文提取人数（如"员工150人""项目团队100人以内"）作为规模，标 derived。

## 多场景处理
- 若需求标签/正文体现多个场景（如"项目管理,财务成本"），拆分识别；主场景取①层中着墨最多、痛点最强者，其余作为次要场景在 note 中说明。

## 🆕 part0（访谈对象判定）—— 最先做，决定所有问题的抽象层级
根据 ①层正文 / tags / decision_role 线索，判断本轮访谈最可能的对象，输出 interviewee：
- "管理层"（老板/总经理/合伙人）：关心经营目标、效率，成本，能不能落地，答不出字段/流程步骤细节。
- "业务执行层"（一线员工/主管）：关心具体操作、数据怎么填、流程卡在哪。
- "职能技术层"（IT/流程负责人）：关心系统对接、权限、数据结构。
无法判断时默认"管理层"（进线阶段多为老板拍板）。后续 must_ask / deep_dive 的问法，必须匹配 interviewee 的层级——绝不拿字段级细节问管理层。

## 🆕 问题落点分层（贯穿 must_ask / deep_dive）—— 治"场景错位"
这是"售前视频会议"阶段，不是"签约后现场逐部门实装调研"。给每条问题判一个 askStage：
- "presale"（售前层）：对面此刻的人（多为决策者）当场答得上，且答案能帮销售判断方向，推动成交。例："这件事目前从头到尾大概怎么流转？""现在最头疼的卡点在哪？""大概涉及多少人？"
- "onsite"（落地层）：必须叫上一线操作岗、对着现有表格/系统逐字段逐流程核对才能答准的执行细节。例："这张表具体要哪些字段、字段什么类型""每个审批节点的具体条件阈值""每个岗位每天几点录哪条数据"。
规则：**售前问卷（must_ask / deep_dive）只放 askStage="presale" 的问题**；判为 "onsite" 的，不拿去问客户，改收进 onsiteChecklist（留到签约后现场再确认）。宁可框架清楚、细节留白，不可细节堆满、客户当场答不上还答不准。

## 🔴 信息单元清单（infoUnits）—— 客户信息保真的通用载体，下游 06/08 据此审计
目的：客户在①层里明确表达过的每一个**信息单元**，都要单独立条、带身份证，一路带到最终方案。防止它在流水线里被合并、降格或悄悄删掉。
> 注意：保真（不丢信息）与售前落点（不越界多问）是两件事。infoUnits 记录客户"说过什么"，必须全留；askStage 决定这些信息里"哪些细节现在追问、哪些留到落地"。留到落地 ≠ 从 infoUnits 删掉。
**什么是"信息单元"**：任何"客户说过，下游不该弄丢"的原子信息——可以是一个业务维度、一个痛点、一个明确诉求、一条硬约束。
【kind 参考池】维度类：知识管理/任务管理/会议管理/客户沟通/话术管理/信息更新同步/数据看板/审批流程/培训考核/财务成本/项目管理/营销获客/权限安全/AI内容生成/外部系统对接；其它 kind：痛点/明确诉求/硬约束（如"必须手机端能用""数据不能上公有云"）。
【立条规则】
1. 逐句扫描①层，凡出现一个可独立描述的诉求/现状/约束，就立一条 infoUnit。
2. 🔴 **禁止收敛性损耗（L1）**：客户**分别**提到的两个单元，哪怕业务上相关、有先后关系，也必须是两条，不许因"下游会合流""看起来是一类"就合并。
3. 每条必带 `sourceQuote`（从①层摘一句最能证明它存在的客户原话，≤30字）。
4. `raisedByCustomer`：true=①层明确说过；false=你基于场景补的（此时 status="derived"）。
5. `status`：covered（信息已够设计）/ pending（提了但细节不足待追问）/ unclear（表述模糊）/ derived（AI补的）。
6. 🔴 **禁止静默丢弃（L3）**：企微做不了的单元（如"1分钟生成讲解视频""三端数据隔离"）**不许删**，照样立条，feasibility 标 "二期"/"需外部产品"，理由写 note。
⚡ **复杂度分级**：若①层信息单元 ≤ 2 条，或走澄清模式 → infoUnits 简单列出即可，不必逐条填全字段，不啰嗦。若 ≥ 3 条（尤其多维度进线）→ 逐条填全 uid/sourceQuote/status，供下游全量审计。

## 概念定义（全程严格遵守）
【sourceType 来源类型】对每条 gap / 问题判定其一：
- "explicit"（明确提及）：能在【原始需求】原文找到对应文字。
- "implicit"（隐含暗示）：客户没直说，可由上下文合理推断。
- "derived"（推导补全）：客户完全没提，由行业+标签补出。

【scopeBoundary 场景边界】= 客户已明确表达的需求点，映射到该行业/场景标准能力地图后所覆盖的范围。它是必问问题不可越过的红线。

## 前置判断（先做，决定分支）
属于"需求未明确"（满足任一即是）：
- industry 为"无明确场景"或空；
- initial_demand 为空，或含"无文字描述/未清晰描述/客户没有描述/待沟通"等；
- 通篇只有联系方式、无任何业务诉求。
→ 是：走【分支B 澄清模式】；否：走【分支A 标准模式】。

================== 分支A：标准模式 ==================
### part1（客户画像）
- company_background：≤100字。
- pain_points：精确5条，每条≤25字，每条标 sourceType。
- customer_type：如"xx行业中型民营企业"。
- main_customers：主要客户群体。
🆕 销售维度（有则填，无则填"待确认"，绝不臆造）：
  - company_scale_guess：人数规模（从①层正文或公司名推断，如"约150人"）。
  - budget_signal：预算信号（客户有无提到预算/报价敏感度/紧迫性）。
  - current_systems：目前在用的系统/工具（钉钉/Excel/某ERP等）。
  - decision_role：进线人角色（老板/中层/执行，呼应 interviewee）。

### part2（待确认信息清单）
- gaps[]：条数随复杂度浮动，每条 { gap, priority, whyNeed, sourceType, uid }。
【优先级判定规则——强制，先判 sourceType 再定 priority】
- priority 取值："高优先级(P0)" / "中优先级(P1)" / "低优先级(P2)"。
- 硬规则一：sourceType="explicit" 的条目，priority 不得低于"高优先级(P0)"。仅当明显是边角诉求（如随口一提的外观偏好）方可降为"中优先级(P1)"，并在 whyNeed 末尾写"（降级原因：…）"。
- 硬规则二：sourceType="derived" 的条目，默认"中优先级(P1)"。仅当它是场景成败关键项（缺了方案无法落地）方可升为"高优先级(P0)"。

### part3（访谈提纲）—— 按 A→E 顺序推导
A. 抽取【原始需求】中客户明确提及的需求点，记为"已知点"。
B. 基于 industry+tags 推导该场景标准能力地图，与"已知点"取交集，得【场景边界】，写入 scope_boundary。
🆕 C. must_ask[]：**先框架后细节 + 只放售前层**，按此规则生成——
   【排序·强制】前 2-3 条固定为框架性问题，顺序：
     ① 业务场景/流程类（这件事从头到尾怎么流转、涉及哪些环节和角色）——至少 2 条，永远排最前；
     ② 现状与痛点类（现在靠什么做、最大的卡点是什么）；
     ③ 才是执行细节类（但仅限 askStage="presale" 的，见下）。
   【落点·强制】每条标 askStage。**must_ask 只收 askStage="presale" 的问题**；凡判为 "onsite"（要操作岗在场逐字段核对）的执行细节，不放这里，改写进 onsiteChecklist。
   【数量·随复杂度】简单进线（原始需求≤2个诉求）给 3-5 条；复杂进线给 6-10 条。宁缺毋滥，严禁把简单需求问复杂，严禁生成边界外行业泛问题。
   【层级·匹配对象】问法匹配 interviewee：对管理层问目标与痛，不问字段。
   若 budget_signal / current_systems 为"待确认"，可在 must_ask 末尾各加 1 条温和探测，语气克制放最后，askStage="presale"。
   每条 { question, dimension, note, needRole, whyAsk, impactIfUnknown, sourceType, askStage, uid }，且必须落在 scope_boundary 内。
D. deep_dive[]：3-6条深挖问题，针对"已知点"的根因/量化，须在场景边界内，且**同样只放 askStage="presale"**。每条 { question, dimension, note, askStage, uid }。
🆕 onsiteChecklist[]：把判为 "onsite" 的执行细节问题收在这里（留到签约后现场、操作岗在场时确认，不在本次售前问客户）。每条 { item, whyOnsite }。可为空数组。
E. industry_experience[]：2-3条行业经验问题，建立专业信任，允许超边界但仅作行业共性探讨，不得与 must_ask 重复。每条 { question, note }。

================== 分支B：澄清模式 ==================
（目的：把模糊进线变成可澄清的开放问题，绝不凭空编造需求）
### part1
- company_background：仅依据已知行业/规模克制描述，不臆造业务细节，≤50字。
- pain_points：最多3条，sourceType 全 "derived"，whyNeed 注"基于行业推测，待客户确认"。
- customer_type / main_customers：基于行业常识保守填。
- 销售维度四项一律填"待确认"。
### part2
- gaps[]：3-5条，priority 一律 "待澄清(pending)"，sourceType 一律 "derived"。
### part3
- scope_boundary："客户需求尚未明确，本轮以澄清为主"。
- must_ask[]：5-7条全开放式澄清问题，sourceType 一律 "explicit_clarify"，dimension 标 "需求澄清"，askStage 一律 "presale"。方向（按行业改写，勿照抄）：
    · 您目前最想优先解决的具体问题是什么？
    · 现在主要靠什么工具/由谁完成？流程如何？
    · 理想状态希望它变成什么样？
    · 大概涉及多少人、多少数据量？
    · 有没有现成表格/系统/截图可参考？
  每条 { question, dimension, note, needRole, whyAsk, impactIfUnknown, sourceType, askStage }。
- deep_dive[]：[]（需求未明确不深挖）。
- onsiteChecklist[]：[]（需求未明确，无落地细节可收）。
- industry_experience[]：1-2条行业共性问题帮打开话题。

## 输出格式（两分支通用，严格 JSON，直接输出，无 markdown 代码块，无注释，无解释）
{
  "mode": "standard | clarify",
  "interviewee": "管理层 | 业务执行层 | 职能技术层",
  "part1": {"company_background":"","pain_points":[{"text":"","sourceType":""}],"customer_type":"","main_customers":"","company_scale_guess":"","budget_signal":"","current_systems":"","decision_role":""},
  "infoUnits": [
    {"uid":"U1","label":"（用客户自己的话，不贴标签名）","kind":"维度/痛点/明确诉求/硬约束","raisedByCustomer":true,"sourceQuote":"（①层客户原话≤30字）","status":"covered/pending/unclear/derived","priority":"高优先级(P0)/中优先级(P1)/低优先级(P2)","feasibility":"可原生实现/二期/需外部产品","note":""}
  ],
  "part2": {"gaps":[{"gap":"","priority":"","whyNeed":"","sourceType":"","uid":""}]},
  "part3": {"scope_boundary":"","must_ask":[{"question":"","dimension":"","note":"","needRole":"","whyAsk":"","impactIfUnknown":"","sourceType":"","askStage":"presale","uid":""}],"deep_dive":[{"question":"","dimension":"","note":"","askStage":"presale","uid":""}],"onsiteChecklist":[{"item":"","whyOnsite":""}],"industry_experience":[{"question":"","note":""}]}
}

## 输出前自检
- 合法 JSON、无注释；各数组无重复条目；priority 带 P 级标注。
- 🆕 **售前落点**：must_ask 与 deep_dive 里没有任何 askStage="onsite" 的条目；需要操作岗在场逐字段核对的细节都进了 onsiteChecklist，没拿去问客户。
- 🆕 **先框架后细节**：must_ask 前 2-3 条是业务流程/场景类框架问题；简单进线（≤2诉求）未被撑成 6-10 条。
- 🆕 **看人下问**：问题层级匹配 interviewee；没有拿字段级细节去问管理层。
- 🆕 **销售维度**：part1 四项销售字段已填或标"待确认"，无臆造。
- 🔴 **保真-无合并**：①层客户分别提到的每个可独立单元，在 infoUnits 里都能找到独立一条，且每条有 sourceQuote 佐证；没有把两个单元揉成一条。
- 🔴 **保真-无丢弃**：企微做不了的单元也在 infoUnits 里（标 feasibility），没被删掉；"留到落地"的细节只影响 askStage，绝不导致 infoUnits 删条。
- 🔴 **回指**：gaps/must_ask/deep_dive 中与某单元相关的条目，uid 回指到 infoUnits（无法归属填 ""）。
- ⚡ **分级**：若信息单元 ≤2 或澄清模式，infoUnits 可精简，不强求填满全字段。
""")


@app.post("/api/question_list")
async def question_list(body: dict, user: dict = Depends(require_auth)):
    """生成 Step1 调研问题（客户画像 + 信息缺口 + 调研清单）"""
    company_name = body.get("company_name", "")
    industry = body.get("industry", "")
    scale = body.get("scale", "")
    tags = body.get("tags", "")
    initial_demand = body.get("initial_demand", "")
    company_intro = body.get("company_intro", "")

    user_prompt = STEP1_USER_PROMPT.substitute(
        company_name=company_name,
        industry=industry,
        scale=scale or "未填写",
        tags=tags or "未填写",
        initial_demand=initial_demand or "未填写",
        company_intro=company_intro or "暂无"
    )

    ai_result = call_deepseek(STEP1_SYSTEM_PROMPT, user_prompt, max_tokens=6000)
    raw = ai_result["content"]
    # Note: question_list doesn't have client_id, skip token recording here
    if raw.startswith("Error:"):
        return {"success": False, "error": raw}

    # 解析 JSON
    result = parse_json_response(raw)
    if not result:
        return {"success": False, "error": "AI 返回格式异常，请重试"}

    return {
        "success": True,
        "result": result,   # 前端 updateClient({ step1_result: data.result })
        "summary": result   # 兼容前端 data.summary 判断
    }


@app.post("/api/step4/generate")
async def generate_step4_artifacts(body: dict, user: dict = Depends(require_auth)):
    """生成 Step4 售前方案和技术路线方案（Prompt 3→4→5 三步）"""
    client_id = body.get("client_id")
    artifact_type = body.get("type", "both")

    if not client_id:
        raise HTTPException(status_code=400, detail="client_id is required")

    # 获取客户数据
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM clients WHERE id = ? AND user_id = ?", (client_id, user["user_id"]))
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="客户不存在")

    client = dict(row)
    for field in ("step1_result", "step2_report", "step2_todo", "step2_schema", "step3_summary", "uploaded_files", "step4_input_draft"):
        if client.get(field) and isinstance(client[field], str):
            try:
                client[field] = json.loads(client[field])
            except:
                pass

    # 构建通用上下文
    customer_name = client.get("name", "")
    industry = client.get("industry", "")
    scale = client.get("scale", "")
    initial_demand = client.get("initial_demand", "")

    step1 = client.get("step1_result", {}) or {}
    company_background = step1.get("part1", {}).get("company_background", "") or ""
    pain_points_raw = step1.get("part1", {}).get("pain_points", []) or []
    pain_points = "\n".join(str(p) if isinstance(p, dict) else p for p in pain_points_raw)
    gaps = "\n".join([f"- {g.get('gap', '')}" for g in (step1.get("part2") or [])])
    must_ask = step1.get("part3", {}) or {}
    must_ask_text = "\n".join([f"{i+1}. {q.get('question', '')}" for i, q in enumerate(must_ask.get("must_ask", []) or [])])

    uploaded_files = client.get("uploaded_files") or []
    transcript = "\n\n".join([f"【{f.get('name', '记录')}】{f.get('content', '') or f.get('text', '')}" for f in uploaded_files if f.get('content') or f.get('text')])

    # 知识库匹配结果
    kb_match_result = ""
    # 预留字段，后续接入知识库时填充

    # xlsx sheet 名和字段摘要
    xlsx_sheet_summary = ""
    step2_schema = client.get("step2_schema") or {}
    if step2_schema:
        sheets = step2_schema.get("sheets") or []
        if sheets:
            lines = []
            for s in sheets:
                name = s.get("name", "")
                cols = [c.get("name", "") for c in (s.get("columns") or []) if c.get("name")]
                lines.append(f"表名：{name}，字段：{', '.join(cols)}")
            xlsx_sheet_summary = "\n".join(lines)

    # 服务商对客户需求总结
    service_provider_summary = ""
    step2_report = client.get("step2_report") or {}
    if isinstance(step2_report, dict):
        summary = step2_report.get("summary", "") or step2_report.get("service_summary", "") or step2_report.get("demand_summary", "")
        if summary:
            service_provider_summary = summary
        else:
            # 尝试从其他常见字段提取
            for key in ("demand_analysis", "provider_summary", "needs_summary"):
                if step2_report.get(key):
                    service_provider_summary = step2_report.get(key)
                    break

    # 用户输入摘要（方案输入确认内容）— step4_input_draft 9字段结构
    step4_input_draft = client.get("step4_input_draft") or {}
    if isinstance(step4_input_draft, str):
        try:
            step4_input_draft = json.loads(step4_input_draft)
        except:
            step4_input_draft = {}
    input_summary = ""
    if step4_input_draft:
        def field_to_text(label, val):
            if not val:
                return ""
            if isinstance(val, list):
                val = "、".join(val)
            return f"{label}：{val}"
        lines = [
            field_to_text("客户现状", step4_input_draft.get("customerCurrentState")),
            field_to_text("核心问题", step4_input_draft.get("painPoints")),
            field_to_text("已确认需求", step4_input_draft.get("confirmedNeeds")),
            field_to_text("涉及角色", step4_input_draft.get("involvedRoles")),
            field_to_text("当前流程", step4_input_draft.get("currentProcess")),
            field_to_text("期望效果", step4_input_draft.get("expectedOutcome")),
            field_to_text("一期范围", step4_input_draft.get("phaseOneScope")),
            field_to_text("二期评估", step4_input_draft.get("phaseTwoScope")),
            field_to_text("待确认问题", step4_input_draft.get("pendingQuestions")),
        ]
        input_summary = "\n".join([l for l in lines if l])

    # step4_input_draft JSON 字符串
    step4_input_draft_str = json.dumps(step4_input_draft, ensure_ascii=False, indent=2) if step4_input_draft else "暂无用户已编辑输入"

    # step4_report.requirementData（历史草稿）
    step4_report = client.get("step4_report") or {}
    if isinstance(step4_report, str):
        try:
            step4_report = json.loads(step4_report)
        except:
            step4_report = {}
    step4_report_str = ""
    if step4_report.get("requirementData"):
        step4_report_str = json.dumps(step4_report.get("requirementData"), ensure_ascii=False, indent=2)

    # step3_summary（第三优先级）
    step3_summary = client.get("step3_summary") or {}
    if isinstance(step3_summary, str):
        try:
            step3_summary = json.loads(step3_summary)
        except:
            step3_summary = {}
    # 渲染 step3_summary 为可读文本（兼容新旧格式）
    step3_summary_text = ""
    if isinstance(step3_summary, dict):
        parts = []
        # ---- 新格式优先 ----
        if step3_summary.get("customerCurrentState"):
            parts.append("客户现状：" + str(step3_summary["customerCurrentState"]))
        if step3_summary.get("painPoints"):
            pains = step3_summary["painPoints"]
            if isinstance(pains, list):
                for p in pains:
                    if isinstance(p, dict):
                        parts.append(f"核心问题：{p.get('title','')}")
                    elif isinstance(p, str):
                        parts.append(f"核心问题：{p}")
        if step3_summary.get("confirmedNeeds"):
            needs = step3_summary["confirmedNeeds"]
            if isinstance(needs, list):
                for n in needs:
                    if isinstance(n, dict):
                        parts.append(f"已确认需求：{n.get('title','')}")
                    elif isinstance(n, str):
                        parts.append(f"已确认需求：{n}")
        if step3_summary.get("involvedRoles"):
            roles = step3_summary["involvedRoles"]
            if isinstance(roles, list):
                for r in roles:
                    if isinstance(r, dict):
                        parts.append(f"涉及角色：{r.get('role','')}（{r.get('responsibility','')}）")
                    elif isinstance(r, str):
                        parts.append(f"涉及角色：{r}")
        if step3_summary.get("currentProcess"):
            parts.append("当前流程：" + str(step3_summary["currentProcess"]))
        if step3_summary.get("expectedOutcome"):
            parts.append("期望效果：" + str(step3_summary["expectedOutcome"]))
        if step3_summary.get("phaseOneScope"):
            scope = step3_summary["phaseOneScope"]
            if isinstance(scope, list):
                for s in scope:
                    if isinstance(s, dict):
                        parts.append(f"一期范围：{s.get('item','')}")
                    elif isinstance(s, str):
                        parts.append(f"一期范围：{s}")
        if step3_summary.get("pendingQuestions"):
            qs = step3_summary["pendingQuestions"]
            if isinstance(qs, list):
                for q in qs:
                    if isinstance(q, dict):
                        parts.append(f"待确认问题：{q.get('question','')}")
                    elif isinstance(q, str):
                        parts.append(f"待确认问题：{q}")
        # ---- 旧格式兼容 ----
        if not parts:
            if step3_summary.get("key_requirements"):
                parts.append("关键需求：" + "、".join(step3_summary["key_requirements"]))
            if step3_summary.get("roles_and_responsibilities"):
                roles = step3_summary["roles_and_responsibilities"]
                if isinstance(roles, list):
                    for r in roles:
                        if isinstance(r, dict):
                            parts.append(f"{r.get('role','')}：{r.get('responsibility','')}")
            if step3_summary.get("decision_chain"):
                dc = step3_summary["decision_chain"]
                if dc.get("decision_maker"):
                    parts.append(f"拍板人：{dc['decision_maker']}")
                if dc.get("influencer"):
                    parts.append(f"影响者：{dc['influencer']}")
                if dc.get("executor"):
                    parts.append(f"执行者：{dc['executor']}")
            if step3_summary.get("progress_and_stages"):
                ps = step3_summary["progress_and_stages"]
                if ps.get("current_stage"):
                    parts.append(f"当前阶段：{ps['current_stage']}")
            if step3_summary.get("risk_points"):
                risks = step3_summary["risk_points"]
                if isinstance(risks, list) and risks:
                    risk_texts = []
                    for r in risks:
                        if isinstance(r, dict) and r.get("risk"):
                            risk_texts.append(r["risk"])
                    if risk_texts:
                        parts.append("风险点：" + "、".join(risk_texts))
        step3_summary_text = "\n".join(parts)

    def build_context(prompt_template):
        def json_escape(s):
            if s is None:
                return ""
            return json.dumps(str(s), ensure_ascii=False)[1:-1]  # 去掉首尾引号

        return (prompt_template
            .replace("{customer_name}", json_escape(customer_name))
            .replace("{industry}", json_escape(industry))
            .replace("{scale}", json_escape(scale))
            .replace("{initial_demand}", json_escape(initial_demand))
            .replace("{company_background}", json_escape(company_background))
            .replace("{pain_points}", json_escape(pain_points))
            .replace("{gaps}", json_escape(gaps))
            .replace("{must_ask}", json_escape(must_ask_text))
            .replace("{transcript}", json_escape(transcript) if transcript else "暂无沟通记录")
            .replace("{input_summary}", json_escape(input_summary) if input_summary else "暂无用户已确认输入")
            .replace("{step4_input_draft}", step4_input_draft_str)
            .replace("{step4_report}", step4_report_str or "暂无Step4历史草稿")
            # 传 JSON 字符串，让 Prompt 3 的 AI 自己推理填充空字段
            .replace("{step3_summary}", json.dumps(step3_summary, ensure_ascii=False, indent=2) if step3_summary else "暂无Step3 AI摘要")
            .replace("{kb_match_result}", json_escape(kb_match_result) if kb_match_result else "暂无知识库匹配结果")
            .replace("{xlsx_sheet_summary}", json_escape(xlsx_sheet_summary) if xlsx_sheet_summary else "暂无xlsx交付物摘要")
            .replace("{service_provider_summary}", json_escape(service_provider_summary) if service_provider_summary else "暂无服务商需求总结")
        )

    result = {}

    # ====== Step 1: Prompt 3 → requirementSolutionData ======
    req_prompt = build_context(STEP4_REQUIREMENT_PROMPT)
    req_ai = call_codebuddy(STEP4_REQUIREMENT_PROMPT, req_prompt, max_tokens=15000)
    req_raw = req_ai["content"]
    record_ai_tokens(client_id, req_ai["usage"]["total_tokens"])
    requirement_data = parse_json_response(req_raw)
    if not requirement_data:
        return {"success": False, "error": "Prompt 3 生成失败（输出被截断或格式异常），请稍后重试。详情：" + (req_raw[:300] if req_raw else "空响应")}

    result["requirementData"] = requirement_data
    req_json_str = json.dumps(requirement_data, ensure_ascii=False)

    # ====== Step 2: Prompt 4 → Word 内容 ======
    if artifact_type in ("both", "presales", "word"):
        word_prompt = STEP4_WORD_PROMPT.replace("{requirement_data}", req_json_str)
        word_content = None
        for attempt in range(3):
            word_ai = call_codebuddy(STEP4_WORD_PROMPT, word_prompt, max_tokens=8000)
            word_raw = word_ai["content"]
            record_ai_tokens(client_id, word_ai["usage"]["total_tokens"])
            word_content = parse_json_response(word_raw)
            if word_content:
                v = validate_requirement_doc(word_content, requirement_data)
                if v["pass"]:
                    break
                if attempt < 2:
                    continue  # 重试
                return {
                    "success": False,
                    "error": "生成结果未通过质检（已重试 2 次）：\n" + "\n".join(v["errors"]),
                    "validation": v
                }
            elif attempt >= 2:
                return {"success": False, "error": "Word 内容生成失败（输出截断），请稍后重试"}
        if word_content:
            result["wordContent"] = word_content

    # ====== Step 3: Prompt 5 → HTML 内容 ======
    if artifact_type in ("both", "html"):
        html_prompt = STEP4_HTML_PROMPT.replace("{requirement_data}", req_json_str)
        html_ai = call_codebuddy(STEP4_HTML_PROMPT, html_prompt, max_tokens=8000)
        html_raw = html_ai["content"]
        record_ai_tokens(client_id, html_ai["usage"]["total_tokens"])
        html_content = parse_json_response(html_raw)
        if html_content:
            result["htmlContent"] = html_content

    # ====== Step 4: Prompt 6 → 技术文档 HTML（文档风，11章26表） ======
    if artifact_type in ("both", "technical"):
        techdoc_prompt = STEP4_TECHDOC_PROMPT.replace("{requirement_data}", req_json_str)
        techdoc_ai = call_codebuddy(STEP4_TECHDOC_PROMPT, techdoc_prompt, max_tokens=12000)
        techdoc_raw = techdoc_ai["content"]
        record_ai_tokens(client_id, techdoc_ai["usage"]["total_tokens"])
        if techdoc_raw and not techdoc_raw.startswith("Error:"):
            result["technicalContent"] = techdoc_raw

    # ====== 推断补全（对 requirementData 空白字段做 fallback ======
    rd = result.get("requirementData") or {}
    meta2 = rd.get("meta") or {}
    facts2 = rd.get("customerFacts") or {}
    bp2 = rd.get("businessProcess") or {}
    st2 = rd.get("smartTableSpec") or {}
    scope2 = rd.get("scope") or {}
    modules2 = rd.get("moduleRecommendation") or []
    pains2 = rd.get("painPoints") or []
    reqs2 = rd.get("requirements") or []
    oqs2 = rd.get("openQuestions") or []
    roles2 = facts2.get("involvedRoles") or []
    confirmed_tables2 = (st2.get("confirmedTables") or []) if st2 else []
    fields_by_table2 = (st2.get("fieldsByTable") or []) if st2 else []
    fields_map2 = {ft.get("tableName", ""): ft.get("fields", []) for ft in fields_by_table2} if fields_by_table2 else {}
    if isinstance(modules2, dict): modules2 = [modules2]
    if isinstance(pains2, dict): pains2 = [pains2]
    if isinstance(reqs2, dict): reqs2 = [reqs2]
    if isinstance(oqs2, dict): oqs2 = [oqs2]
    if isinstance(roles2, dict): roles2 = [roles2]
    if isinstance(confirmed_tables2, dict): confirmed_tables2 = [confirmed_tables2]

    def _l2(v): return v if isinstance(v, list) else ([v] if v else [])
    def _s2(v):
        if not v or isinstance(v, bool): return ""
        if isinstance(v, (int, float)): return str(v)
        if isinstance(v, str): return v.strip()
        if isinstance(v, list): return "、".join(_s2(x) for x in v if x)
        if isinstance(v, dict): return v.get("title") or v.get("name") or v.get("item") or str(v)
        return str(v)

    # 1. painPoints 为空 → 从 currentFlow.problem 推断
    if not pains2:
        for f in _l2(bp2.get("currentFlow", [])):
            prob = f.get("problem", "")
            if prob:
                pains2.append({"title": _s2(prob)[:30], "description": "当前流程中存在该问题。", "businessImpact": "影响整体效率。", "priority": "P1"})
        rd["painPoints"] = pains2

    # 2. requirements 为空但有 confirmed_tables → 从表推断需求
    if not reqs2 and confirmed_tables2:
        for t in confirmed_tables2:
            tn = _s2(t.get("tableName", ""))
            ph = _s2(t.get("phase", "一期"))
            reqs2.append({"requirementName": "搭建" + tn + "智能表格", "customerExpression": "需要管理" + tn + "相关数据。", "businessTranslation": "通过企业微信智能表格实现" + tn + "在线化管理。", "priority": "P1", "phase": ph, "confirmedStatus": "AI推断"})
        rd["requirements"] = reqs2

    # 3. businessProcess.processNodes 为空 → 从 currentFlow/targetFlow 推断
    nodes2 = _l2(bp2.get("processNodes", []))
    cur_flow2 = _l2(bp2.get("currentFlow", []))
    tgt_flow2 = _l2(bp2.get("targetFlow", []))
    if not nodes2 and (cur_flow2 or tgt_flow2):
        for i, cf in enumerate(cur_flow2):
            ti = tgt_flow2[i] if i < len(tgt_flow2) else {}
            nodes2.append({
                "nodeName": _s2(cf.get("stepName", "")),
                "responsibleRole": _s2(cf.get("role", "")),
                "input": _s2(cf.get("input", "")),
                "output": _s2(cf.get("output", "")),
                "systemAction": _s2(ti.get("systemAction", cf.get("currentMethod", ""))),
                "reminderNeeded": bool(ti.get("reminderNeeded", False))
            })
        bp2["processNodes"] = nodes2

    # 4. automations 为空 → 从 painPoints 推断标准自动化
    if not _l2(st2.get("automations", [])) and pains2:
        autos2 = []
        for p in pains2[:3]:
            autos2.append({"name": "「" + _s2(p.get("title", "")) + "」跟进提醒", "trigger": "当记录状态变更时", "action": "推送提醒至负责人", "notifyTarget": "、".join(roles2[:2]) if roles2 else "负责人", "priority": "中"})
        st2["automations"] = autos2

    # 5. permissions 为空 → 从 roles + confirmed_tables 推断
    if not _l2(st2.get("permissions", [])) and roles2:
        perms2 = []
        for r in roles2:
            perms2.append({"role": _s2(r), "addScope": "本职范围内", "viewScope": "、".join(_s2(t.get("tableName", "")) for t in confirmed_tables2[:3]), "editableFields": "本职相关字段", "sensitiveFields": "——"})
        st2["permissions"] = perms2

    # 6. dashboards 为空 → 从 confirmed_tables 推断
    if not _l2(st2.get("dashboards", [])) and confirmed_tables2:
        st2["dashboards"] = [{"dashboardName": _s2(meta2.get("mainScenario", "业务")) + "管理看板", "users": "、".join(roles2[:2]) if roles2 else "业务负责人", "metrics": "记录数量、跟进状态、转化率", "filters": "时间范围"}]

    # 7. openQuestions 为空 → 从 pains 推断
    if not oqs2 and pains2:
        oqs2 = [{"question": "如何量化评估「" + _s2(p.get("title", "")) + "」的改善效果？", "whyAsk": "用于设定实施目标", "owner": "服务商+客户", "priority": "中"} for p in pains2[:3]]
        rd["openQuestions"] = oqs2

    # 8. implementationPlanTable 为空 → 标准4阶段
    if not _l2(rd.get("implementationPlanTable", [])):
        rd["implementationPlanTable"] = [
            {"phase": "需求确认", "workContent": "确认需求范围、一期二期边界、智能表格结构", "customerCooperation": "提供业务需求，参与评审", "output": "需求确认文档"},
            {"phase": "智能表格搭建", "workContent": "按确认的字段设计搭建智能表格", "customerCooperation": "参与字段确认、提供基础数据", "output": "可用的智能表格 Demo"},
            {"phase": "规则配置", "workContent": "配置自动化规则、审批流、权限", "customerCooperation": "参与规则评审、测试确认", "output": "配置完成的规则"},
            {"phase": "试运行与优化", "workContent": "上线试运行，收集反馈并优化", "customerCooperation": "提供试运行数据、反馈问题", "output": "上线文档+优化建议"},
        ]

    # 9. quoteScopeTable 为空 → 从 confirmed_tables 生成口径
    if not _l2(rd.get("quoteScopeTable", [])) and confirmed_tables2:
        qt2 = []
        for t in confirmed_tables2:
            ph = _s2(t.get("phase", ""))
            if "一期" in ph and "二期" not in ph:
                qt2.append({"feeModule": _s2(t.get("tableName", "")), "scope": "智能表格搭建及基础配置", "included": "是", "note": "费用口径待与服务商确认"})
        rd["quoteScopeTable"] = qt2

    # 10. changeManagementTable 为空 → 标准变更机制
    if not _l2(rd.get("changeManagementTable", [])):
        rd["changeManagementTable"] = [
            {"item": "新增智能表格字段", "inScope": "否", "handling": "评估后进入变更单处理"},
            {"item": "跨表关联新增", "inScope": "否", "handling": "评估工作量后单独报价"},
            {"item": "自动化规则增加", "inScope": "否", "handling": "评估后进入变更单处理"},
        ]

    # 11. deliveryList 为空 → 从 confirmed_tables 生成
    if not _l2(rd.get("deliveryList", [])) and confirmed_tables2:
        dl2 = []
        for t in confirmed_tables2:
            ph = _s2(t.get("phase", ""))
            dl2.append({"category": "企业微信智能表格", "content": _s2(t.get("tableName", "")), "included": "✅ 一期" if "一期" in ph and "二期" not in ph else "🔄 二期", "note": _s2(t.get("tablePurpose", ""))})
        rd["deliveryList"] = dl2

    result["requirementData"] = rd
    result["success"] = True
    return result


# ==================== Step4 模板化新 API（按钮① HTML / 按钮② Word）====================

# 模板包路径
TEMPLATE_BASE = str(Path(__file__).parent / "templates" / "双按钮模板包")
OUTPUTS_BASE = str(Path(__file__).parent.parent / "outputs")


def _read_template(path: str) -> str:
    """读取服务器上的模板文件"""
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        return ""


def _load_client_context(client: dict, user: dict) -> dict:
    """从 client 数据中提取所有步骤的完整上下文，返回 dict"""
    import logging
    logger = logging.getLogger("uvicorn")

    def jload(val):
        if isinstance(val, str):
            try:
                return json.loads(val)
            except:
                return val
        return val

    ctx = {
        "customer_name": client.get("name", ""),
        "industry": client.get("industry", ""),
        "scale": client.get("scale", ""),
        "tags": client.get("tags", ""),
        "initial_demand": client.get("initial_demand", ""),
    }

    logger.info(f"[Word诊断] client_id={client.get('id')} name={client.get('name')} step1_result=exists:{bool(client.get('step1_result'))} step2_report=exists:{bool(client.get('step2_report'))} step2_schema=exists:{bool(client.get('step2_schema'))} step3_summary=exists:{bool(client.get('step3_summary'))} step4_input_draft=exists:{bool(client.get('step4_input_draft'))}")

    # Step1 result
    s1 = jload(client.get("step1_result")) or {}
    if isinstance(s1, dict):
        p1 = s1.get("part1", {}) or {}
        ctx["company_background"] = p1.get("company_background", "")
        pains_raw = p1.get("pain_points", []) or []
        if isinstance(pains_raw, list):
            ctx["pain_points_text"] = "\n".join(
                f"- {str(p.get('title',''))}: {str(p.get('description',''))}" if isinstance(p, dict) else f"- {p}"
                for p in pains_raw if p
            )
        else:
            ctx["pain_points_text"] = str(pains_raw)
        ctx["pain_points"] = s1.get("part1", {}).get("pain_points", [])
        gaps_raw = s1.get("part2") or []
        ctx["gaps_text"] = "\n".join(
            f"- {g.get('gap','')}" if isinstance(g, dict) else f"- {g}"
            for g in gaps_raw if g
        )
        must_ask_list = (s1.get("part3") or {}).get("must_ask") or []
        ctx["must_ask_text"] = "\n".join(
            f"{i+1}. {q.get('question','')}" if isinstance(q, dict) else f"{i+1}. {q}"
            for i, q in enumerate(must_ask_list) if q
        )
        ctx["step1_summary"] = s1.get("summary", "")
    else:
        ctx["company_background"] = ""
        ctx["pain_points_text"] = ""
        ctx["pain_points"] = []
        ctx["gaps_text"] = ""
        ctx["must_ask_text"] = ""
        ctx["step1_summary"] = ""

    # Step2 report
    s2 = jload(client.get("step2_report")) or {}
    if isinstance(s2, dict):
        ctx["step2_summary"] = s2.get("summary", "") or s2.get("service_summary", "") or s2.get("demand_summary", "") or str(s2)
    else:
        ctx["step2_summary"] = str(s2) if s2 else ""
    ctx["step2_report_full"] = json.dumps(s2, ensure_ascii=False) if isinstance(s2, dict) else (s2 or "")

    # Step2 schema (xlsx sheets)
    s2schema = jload(client.get("step2_schema")) or {}
    sheets = (s2schema.get("sheets") or []) if isinstance(s2schema, dict) else []
    ctx["xlsx_summary"] = ""
    if sheets:
        lines = []
        for sh in sheets:
            cols = [c.get("name","") for c in (sh.get("columns") or []) if c.get("name")]
            lines.append(f"表名：{sh.get('name','') if isinstance(sh,dict) else sh}，字段：{', '.join(cols)}")
        ctx["xlsx_summary"] = "\n".join(lines)

    # Step3: 上传文件内容 + 沟通纪要全文
    uploaded_files = jload(client.get("uploaded_files")) or []
    transcript_list = []
    for f in uploaded_files:
        if isinstance(f, dict):
            content = f.get("content", "") or f.get("text", "") or ""
            name = f.get("name", "记录")
            if content:
                transcript_list.append(f"【{name}】\n{content}")
    transcript_raw = jload(client.get("transcript"))
    if transcript_raw:
        if isinstance(transcript_raw, list):
            for t in transcript_raw:
                c = t.get("content","") or t.get("text","") or ""
                n = t.get("name","记录")
                if c:
                    transcript_list.append(f"【{n}】\n{c}")
        elif isinstance(transcript_raw, str) and transcript_raw:
            transcript_list.append(transcript_raw)
    ctx["step3_transcript_full"] = "\n\n".join(transcript_list) if transcript_list else "（暂无沟通记录）"

    # Step3 summary
    s3 = jload(client.get("step3_summary")) or {}
    if isinstance(s3, dict):
        parts = []
        for key in ("customerCurrentState","painPoints","confirmedNeeds","involvedRoles",
                    "currentProcess","expectedOutcome","phaseOneScope","phaseTwoScope","pendingQuestions"):
            v = s3.get(key)
            if v:
                if isinstance(v, list):
                    parts.append(f"{key}：\n" + "\n".join(
                        f"- {x.get('title','') if isinstance(x,dict) else x}" for x in v if x))
                else:
                    parts.append(f"{key}：{v}")
        ctx["step3_summary_full"] = "\n".join(parts) if parts else str(s3)
    else:
        ctx["step3_summary_full"] = str(s3) if s3 else ""

    # Step4 input draft（9字段完整 JSON）
    s4draft = jload(client.get("step4_input_draft")) or {}
    draft_parts = []
    if s4draft:
        for key, label in [
            ("customerCurrentState","客户现状"),
            ("painPoints","核心问题"),
            ("confirmedNeeds","已确认需求"),
            ("involvedRoles","涉及角色"),
            ("currentProcess","当前流程"),
            ("expectedOutcome","期望效果"),
            ("phaseOneScope","一期范围"),
            ("phaseTwoScope","二期评估"),
            ("pendingQuestions","待确认问题"),
        ]:
            v = s4draft.get(key)
            if v:
                if isinstance(v, list):
                    draft_parts.append(f"【{label}】\n" + "\n".join(f"- {x.get('title','') if isinstance(x,dict) else x}" for x in v if x))
                else:
                    draft_parts.append(f"【{label}】\n{v}")
    ctx["step4_input_draft_full"] = "\n\n".join(draft_parts) if draft_parts else "（暂无 Step4 方案输入确认内容）"
    ctx["step4_input_draft_json"] = json.dumps(s4draft, ensure_ascii=False, indent=2) if s4draft else "{}"

    ctx["provider_name"] = user.get("provider_name", "")

    logger.info(f"[Word诊断] ctx构建完成: company_background长度={len(ctx.get('company_background',''))} pain_points_text长度={len(ctx.get('pain_points_text',''))} step2_summary长度={len(ctx.get('step2_summary',''))} xlsx_summary长度={len(ctx.get('xlsx_summary',''))} step3_transcript_full长度={len(ctx.get('step3_transcript_full',''))} step3_summary_full长度={len(ctx.get('step3_summary_full',''))} step4_input_draft_full长度={len(ctx.get('step4_input_draft_full',''))}")

    return ctx


def _build_requirement_data_from_context(ctx: dict) -> dict:
    """
    从 client 上下文中构建 requirementData（供 Step5 generate-demo 使用）。
    新流程（generate-html / generate-word）不走旧的 /api/step4/generate，
    因此 requirementData 不会自动产生，需要在此手动构建。
    """
    import logging
    logger = logging.getLogger("uvicorn")

    s4draft_str = ctx.get("step4_input_draft_json", "{}")
    s4draft = {}
    try:
        s4draft = json.loads(s4draft_str) if s4draft_str and s4draft_str != "{}" else {}
    except:
        s4draft = {}

    s3_str = ctx.get("step3_summary_full", "")
    # step3_summary_full 是拼接的文本，不是 JSON，跳过

    # 从 step4_input_draft 提取字段
    def safe_list(val):
        if not val:
            return []
        if isinstance(val, list):
            return val
        return [val]

    confirmed_needs = safe_list(s4draft.get("confirmedNeeds", []))
    pain_points = safe_list(s4draft.get("painPoints", []))
    involved_roles = safe_list(s4draft.get("involvedRoles", []))
    phase_one = safe_list(s4draft.get("phaseOneScope", []))
    phase_two = safe_list(s4draft.get("phaseTwoScope", []))
    pending_questions = safe_list(s4draft.get("pendingQuestions", []))

    # 从 step3_summary_full 尝试提取表格信息（文本片段匹配）
    # 实际表格字段由 AI 在 Step5 prompt 中生成，这里只构建结构
    confirmed_tables = []
    fields_by_table = []

    # 构建 confirmedNeeds 为 requirements 格式
    requirements = []
    for i, n in enumerate(confirmed_needs):
        title = n.get("title", "") if isinstance(n, dict) else str(n)
        desc = n.get("description", "") if isinstance(n, dict) else ""
        requirements.append({
            "requirementId": f"REQ-{i+1}",
            "title": title,
            "description": desc,
            "priority": "高",
            "phase": "一期",
            "status": "confirmed"
        })

    # 构建 painPoints
    parsed_pains = []
    for i, p in enumerate(pain_points):
        title = p.get("title", "") if isinstance(p, dict) else str(p)
        desc = p.get("description", "") if isinstance(p, dict) else ""
        parsed_pains.append({
            "painId": f"PAIN-{i+1}",
            "title": title,
            "description": desc,
            "priority": "high"
        })

    # 构建 scope
    def scope_item_builder(items, phase):
        result = []
        for it in items:
            if isinstance(it, dict):
                result.append({
                    "item": it.get("item", "") or it.get("title", ""),
                    "description": it.get("description", ""),
                    "reason": it.get("reason", ""),
                    "phase": phase
                })
            else:
                result.append({"item": str(it), "description": "", "reason": "", "phase": phase})
        return result

    phase_one_scope = scope_item_builder(phase_one, "一期")
    phase_two_scope = scope_item_builder(phase_two, "二期评估")

    # roles
    roles_list = []
    for r in involved_roles:
        if isinstance(r, dict):
            roles_list.append({"role": r.get("role", ""), "responsibility": r.get("responsibility", "")})
        else:
            roles_list.append({"role": str(r), "responsibility": ""})

    requirement_data = {
        "meta": {
            "customerName": ctx.get("customer_name", ""),
            "mainScenario": ctx.get("industry", ""),
        },
        "customerFacts": {
            "involvedRoles": roles_list,
            "currentProcess": s4draft.get("currentProcess", ""),
            "expectedOutcome": s4draft.get("expectedOutcome", ""),
        },
        "smartTableSpec": {
            "scenarioComplexity": "简单流程型",
            "confirmedTables": confirmed_tables,
            "fieldsByTable": fields_by_table,
            "suggestedTables": [],
            "phaseTwoTables": [],
            "automations": [],
            "permissions": [],
            "dashboards": [],
        },
        "scope": {
            "phaseOne": phase_one_scope,
            "phaseTwo": phase_two_scope,
            "notRecommended": []
        },
        "requirements": requirements,
        "painPoints": parsed_pains,
        "openQuestions": [{"question": str(q)} for q in pending_questions],
        "implementationPlanTable": [],
        "quoteScopeTable": [],
        "changeManagementTable": [],
        "deliveryList": []
    }

    logger.info(f"[requirementData] built from context: confirmedTables={len(confirmed_tables)}, requirements={len(requirements)}, phaseOne={len(phase_one_scope)}")
    return requirement_data


def _build_html_prompt(template_html: str, golden_rules: str, example_html: str, ctx: dict, feedback: str = "") -> tuple:
    """构建 HTML 生成的 prompt，返回 (system_prompt, user_prompt)"""
    system = (
        "你是企业微信定制开发服务商的方案专家。请严格使用我给的 HTML 模板骨架，"
        "只替换 {{双花括号}} 占位符为真实业务内容，禁止改动任何 CSS/版式/配色/章节结构。\n"
        "产物中不允许残留任何 {{}} 占位符或'（待补充）'/'（待识别）'/'（暂无）'等空壳话术。\n"
        "务必参考我给的范例风格（Hero 仿真看板、分层架构图、厚场景卡）。\n"
        "黄金规则：\n" + golden_rules
    )
    feedback_block = ("\n\n【修改反馈（必须全部采纳）】\n" + feedback + "\n") if feedback else ""
    user = (
        "【模板骨架】\n" + template_html + "\n\n"
        "【客户材料正文】\n"
        f"客户名称：{ctx['customer_name']}\n"
        f"所属行业：{ctx['industry']}\n"
        f"企业规模：{ctx['scale']}\n"
        f"客户标签：{ctx['tags']}\n"
        f"初始需求：{ctx['initial_demand']}\n\n"
        f"【Step1 客户背景与痛点】\n{ctx['company_background']}\n\n"
        f"【Step1 痛点列表】\n{ctx['pain_points_text']}\n\n"
        f"【Step1 调研缺口】\n{ctx['gaps_text']}\n\n"
        f"【Step1 必问问题清单】\n{ctx['must_ask_text']}\n\n"
        f"【Step2 调研结果摘要】\n{ctx['step2_summary']}\n\n"
        f"【Step2 调研完整记录】\n{ctx['step2_report_full']}\n\n"
        f"【Step2 XLSX 表格结构】\n{ctx['xlsx_summary']}\n\n"
        f"【Step3 沟通记录全文】\n{ctx['step3_transcript_full']}\n\n"
        f"【Step3 需求摘要】\n{ctx['step3_summary_full']}\n\n"
        f"【Step4 方案输入确认内容】\n{ctx['step4_input_draft_full']}\n\n"
        "【参考范例（城邦美商 HTML）】\n" + example_html + "\n\n"
        "【任务】\n"
        "请基于以上客户材料，把模板填充成完整的售前 HTML 方案。"
        "所有 {{}} 占位符必须全部替换为真实内容，禁止残留任何占位符或空壳话术。"
        + feedback_block +
        ("【重要】请务必按照【修改反馈】中的每一条意见进行修改后再输出。" if feedback else "")
    )
    return system, user


def _build_word_prompt(template_docx_path: str, golden_rules: str, example_word_text: str, ctx: dict, feedback: str = "") -> tuple:
    """构建 Word 生成的 prompt，返回 (system_prompt, user_prompt)"""
    system = (
        "你是企业微信定制开发服务商的需求确认文档专家。请严格按 11 节结构输出内容，"
        "每节内按表填充，格式为：【节标题】\\n字段名=内容。\n"
        "禁止输出任何 {{}} 占位符，必须全部替换为真实内容。\n"
        "禁止使用'（待补充）'/'（暂无）'/'（待确认）'等空壳话术。\n"
        "黄金规则：\n" + golden_rules
    )
    user = (
        "【Word 模板结构（锁死骨架，不许改表头）】\n"
        "一、客户基础信息与当前现状\n"
        "  1.1 客户基础信息确认表：公司全称、所属行业、企业规模、企业微信使用情况、已有业务系统、对接人/决策人、本次诉求一句话\n"
        "  1.2 当前业务运转方式：环节、当前做法、主要问题\n"
        "  1.3 核心痛点与优先级：痛点编号、痛点描述、业务影响、优先级\n"
        "二、场景类型判断与方案边界\n"
        "  2.1 场景类型判断：判断项、内容\n"
        "  2.2 一期/二期/不建议范围：类型、范围说明、原因\n"
        "三、需求理解与优先级确认\n"
        "  3.1 需求清单：需求项、客户描述、业务影响、优先级、一期/二期、企业微信实现方式\n"
        "  3.2 客户原话与业务翻译：客户原话、业务语言翻译、是否已确认\n"
        "四、业务流程设计\n"
        "  4.1 当前流程与目标流程：阶段、当前流程、优化后流程、企业微信动作\n"
        "  4.2 流程节点确认表：节点序号、节点名称、操作角色、输入信息、输出结果、是否提醒\n"
        "五、企业微信方案总览\n"
        "  5.1 能力架构：层级、能力、本项目使用方式\n"
        "六、智能表格交付设计\n"
        "  6.1 智能表格总览：表名、表类型、用途、使用对象、一期必做\n"
        "  6.2 核心字段设计表：所属表、字段分组、字段名称、字段类型、是否必填、填写角色、规则说明\n"
        "  6.3 表间关联关系：主表、关联表、关联字段、自动带出/汇总内容、注意事项\n"
        "七、审批与自动化设计\n"
        "  7.1 自动化规则表：规则名称、触发条件、执行动作、通知对象、优先级\n"
        "  7.2 审批流程设计：审批名称、发起角色、审批人、通过后动作、同步表格\n"
        "八、权限与数据看板设计\n"
        "  8.1 权限矩阵：角色、新增、查看范围、可编辑字段、看板权限、敏感字段\n"
        "  8.2 数据看板设计：看板名称、使用对象、核心指标、筛选维度、是否支持下钻\n"
        "九、数据来源、系统对接与交付边界\n"
        "  9.1 数据来源与接入方式：数据对象、来源系统/来源方式、一期接入方式、二期评估事项\n"
        "  9.2 交付清单：类别、交吽内容、是否包含、备注\n"
        "十、实施计划、报价口径与变更机制\n"
        "  10.1 实施计划：阶段、工作内容、客户配合事项、输出物\n"
        "  10.2 报价口径建议：费用模块、范围说明、是否本次包含、备注\n"
        "  10.3 范围变更机制：事项、是否范围内、处理方式\n"
        "十一、待客户确认问题与签署\n"
        "  11.1 待确认问题清单：问题编号、待确认问题、负责人、截止时间、确认结果\n"
        "  11.2 客户确认：确认事项、确认说明\n\n"
        "【客户材料正文】\n"
        f"客户名称：{ctx['customer_name']}\n"
        f"所属行业：{ctx['industry']}\n"
        f"企业规模：{ctx['scale']}\n"
        f"服务商名称：{ctx['provider_name']}\n\n"
        f"【Step1 客户背景与痛点】\n{ctx['company_background']}\n\n"
        f"【Step1 痛点列表】\n{ctx['pain_points_text']}\n\n"
        f"【Step2 调研结果摘要】\n{ctx['step2_summary']}\n\n"
        f"【Step2 调研完整记录】\n{ctx['step2_report_full']}\n\n"
        f"【Step3 沟通记录全文】\n{ctx['step3_transcript_full']}\n\n"
        f"【Step3 需求摘要】\n{ctx['step3_summary_full']}\n\n"
        f"【Step4 方案输入确认内容】\n{ctx['step4_input_draft_full']}\n\n"
        "【参考范例（省心住 Word 文档结构）】\n" + example_word_text + "\n\n"
        "【任务】\n"
        "请基于以上客户材料，按 11 节结构输出 Word 文档内容。"
        "格式：【节标题】\\n字段名=内容，字段名与 docx 模板表头对齐。"
        + ("\n\n【修改反馈（必须全部采纳）】\n" + feedback + "\n\n【重要】请务必按照【修改反馈】中的每一条意见进行修改后再输出。" if feedback else "")
    )
    return system, user


@app.post("/api/step4/generate-html")
async def generate_step4_html(body: dict, user: dict = Depends(require_auth)):
    """按钮①：生成 HTML 售前解决方案（模板 + 黄金规则 + 真实材料）"""
    client_id = body.get("client_id")
    if not client_id:
        raise HTTPException(status_code=400, detail="client_id is required")

    # 取客户数据
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM clients WHERE id = ? AND user_id = ?", (client_id, user["user_id"]))
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="客户不存在")

    client = dict(row)
    feedback = body.get("feedback", "")

    # 读模板文件
    template_html = _read_template(TEMPLATE_BASE + "/templates/售前解决方案_HTML模板.html")
    golden_rules = _read_template(TEMPLATE_BASE + "/黄金规则.md")
    example_html = _read_template(TEMPLATE_BASE + "/examples/范例1_城邦美商_HTML.html")

    if not template_html or not golden_rules:
        return {"success": False, "error": "模板文件读取失败，请在服务器上确认模板包已正确上传"}

    # 构建上下文
    ctx = _load_client_context(client, user)

    # 构建 prompt
    system_prompt, user_prompt = _build_html_prompt(template_html, golden_rules, example_html, ctx, feedback)

    # 调用 CodeBuddy（最多重试1次）
    html_content = None
    for attempt in range(2):
        ai_result = call_codebuddy(system_prompt, user_prompt, max_tokens=12000)
        raw = ai_result["content"]
        record_ai_tokens(client_id, ai_result["usage"]["total_tokens"])
        if raw.startswith("Error:"):
            if attempt == 0:
                continue  # 重试一次
            return {"success": False, "error": raw}

        # 检查残留占位符
        import subprocess, tempfile, os
        with tempfile.NamedTemporaryFile(suffix=".html", mode="w", encoding="utf-8", delete=False) as tmp:
            tmp.write(raw)
            tmp_path = tmp.name
        try:
            result = subprocess.run(
                ["bash", TEMPLATE_BASE + "/scripts/check_placeholder.sh", tmp_path],
                capture_output=True, text=True, timeout=60
            )
            placeholder_ok = (result.returncode == 0)
        except Exception:
            placeholder_ok = True  # 自检失败不阻止
        finally:
            os.unlink(tmp_path)

        if placeholder_ok:
            html_content = raw
            break

        if attempt == 0:
            # 重试，换个方式请求
            continue

    if not html_content:
        return {"success": False, "error": "生成内容未通过自检（残留占位符），请稍后重试"}

    # 保存文件
    client_id_str = str(client_id)
    output_dir = OUTPUTS_BASE + "/" + client_id_str
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"售前解决方案_{ctx['customer_name']}_{timestamp}.html"
    filepath = os.path.join(output_dir, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html_content)

    file_url = "/outputs/" + client_id_str + "/" + filename

    # 更新版本数组（同时构建 requirementData 供 Step5 使用）
    now = datetime.now().isoformat()
    versions = client.get("step4_presales_versions") or []
    if isinstance(versions, str):
        try:
            versions = json.loads(versions)
        except:
            versions = []
    next_ver = len(versions) + 1
    # 从 ctx 构建 requirementData（新流程无旧数据，需手动构建）
    requirement_data = _build_requirement_data_from_context(ctx)
    versions.append({
        "version": next_ver,
        "content": {"htmlContent": html_content, "requirementData": requirement_data},
        "created_at": now,
        "file_url": file_url,
        "filename": filename,
    })
    conn2 = get_db()
    cursor2 = conn2.cursor()
    cursor2.execute(
        "UPDATE clients SET step4_presales_versions = ? WHERE id = ?",
        (json.dumps(versions, ensure_ascii=False), client_id)
    )
    conn2.commit()
    conn2.close()

    return {
        "success": True,
        "file_url": file_url,
        "filename": filename,
        "version": next_ver,
        "html_content": html_content,
    }


@app.post("/api/step4/generate-word")
async def generate_step4_word(body: dict, user: dict = Depends(require_auth)):
    """按钮②：生成 Word 报价方案（模板 + 黄金规则 + 真实材料 → minimax → 回填 docx）"""
    client_id = body.get("client_id")
    if not client_id:
        raise HTTPException(status_code=400, detail="client_id is required")

    # 取客户数据
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM clients WHERE id = ? AND user_id = ?", (client_id, user["user_id"]))
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="客户不存在")

    client = dict(row)
    feedback = body.get("feedback", "")

    # 读模板文件
    golden_rules = _read_template(TEMPLATE_BASE + "/黄金规则.md")
    example_word_text = _read_template(TEMPLATE_BASE + "/examples/范例2_省心住_Word.docx")
    # example_word 是 docx，需要提取文本
    try:
        from docx import Document
        doc_ex = Document(TEMPLATE_BASE + "/examples/范例2_省心住_Word.docx")
        example_word_text = "\n".join(p.text for p in doc_ex.paragraphs if p.text.strip())
    except Exception:
        example_word_text = golden_rules  # fallback

    if not golden_rules:
        return {"success": False, "error": "模板文件读取失败，请在服务器上确认模板包已正确上传"}

    # 构建上下文
    ctx = _load_client_context(client, user)

    # 构建 prompt
    system_prompt, user_prompt = _build_word_prompt("", golden_rules, example_word_text, ctx, feedback)

    # 调用 CodeBuddy（最多重试1次）
    word_text_raw = None
    for attempt in range(2):
        ai_result = call_codebuddy(system_prompt, user_prompt, max_tokens=15000)
        raw = ai_result["content"]
        record_ai_tokens(client_id, ai_result["usage"]["total_tokens"])
        if raw.startswith("Error:"):
            if attempt == 0:
                continue
            return {"success": False, "error": raw}
        word_text_raw = raw
        break

    if not word_text_raw:
        return {"success": False, "error": "Word 内容生成失败，请稍后重试"}

    import logging
    logger = logging.getLogger("uvicorn")
    logger.info(f"[Word诊断] 模型原始输出长度={len(word_text_raw)} 前500字={word_text_raw[:500]}")

    # ---- 极简回填：只替换单元格内容恰好等于 {{key}} 的情况 ----
    try:
        from docx import Document
    except Exception as e:
        return {"success": False, "error": f"Word 模板读取失败: {str(e)}"}

    doc = Document(TEMPLATE_BASE + "/templates/技术路线及报价方案_Word模板.docx")

    # 解析模型输出，构建 key→value 映射
    key_map = {}
    for line in word_text_raw.split('\n'):
        line = line.strip()
        if not line or line.startswith('【'):
            continue
        eq_idx = line.find('=')
        if eq_idx > 0:
            key = line[:eq_idx].strip()
            val = line[eq_idx+1:].strip()
            if key and val and val != 'None' and key not in key_map:
                key_map[key] = val

    # docx 占位符名 → 模型输出字段名 的别名映射（完全覆盖所有90+占位符）
    alias = {
        # 封面/基本信息
        "客户名": "公司全称", "服务商名": "服务商名称", "行业": "所属行业",
        "日期": "输出日期", "联系人": "客户联系人", "第几次沟通": "沟通阶段",
        "人数/门店数": "企业规模",
        "是否已用、用到什么程度": "企业微信使用情况",
        "系统名，如装修云管家/ERP": "已有业务系统",
        "角色": "对接人/决策人", "老板最想要的": "本次诉求一句话",
        "项目场景": "项目场景",
        # 1.2 当前业务运转
        "环节1": "环节", "环节2": "环节", "环节3": "环节", "环节4": "环节",
        "当前做法": "当前做法", "当前": "当前流程", "优化后": "优化后流程",
        "企微动作": "企业微信动作",
        # 1.3 痛点
        "问题": "主要问题", "影响": "业务影响", "方案": "对应企微方案",
        # 2.1 场景
        "如：客户销售管理/订单交付/数据看板": "场景大类",
        "前端看板 / 全流程线上化 / 局部优化": "本期定位",
        "不替代什么、不重建什么": "交付边界",
        # 2.2 范围
        "下一阶段": "范围说明", "暂不做": "范围说明",
        "为什么先做": "原因", "为什么放二期": "原因", "为什么不做什么": "原因",
        # 3.1 需求
        "描述": "客户描述", "实现方式": "企业微信实现方式",
        # 3.2 原话翻译
        "客户原话": "客户原话", "翻译成业务需求": "业务语言翻译",
        # 4.1 流程
        "阶段1": "阶段", "阶段2": "阶段", "阶段3": "阶段", "阶段4": "阶段",
        # 4.2 节点
        "节点": "节点名称", "输入": "输入信息", "输出": "输出结果",
        # 5.1 能力架构
        "怎么用": "本项目使用方式",
        # 6.1 表格总览
        "表1": "表名", "表2": "表名", "表3": "表名", "表名": "表名",
        "用途": "用途",
        # 6.2 字段
        "字段": "字段名称",
        # 6.3 关联
        "主表": "主表", "关联表": "关联表", "带出内容": "自动带出/汇总内容", "注意": "注意事项",
        # 7.1 自动化
        "触发": "触发条件", "动作": "执行动作", "对象": "通知对象",
        "规则1": "规则名称", "规则2": "规则名称", "规则3": "规则名称",
        "规则": "规则名称",
        # 7.2 审批
        "审批1": "审批名称", "审批2": "审批名称", "发起角色": "发起角色", "审批人": "审批人",
        "表": "同步表格",
        # 8.1 权限
        "部门负责人": "角色", "一线员工": "角色", "本人": "查看范围", "本部门": "查看范围",
        "看板": "看板权限", "脱敏规则": "敏感字段",
        # 8.2 看板
        "看板1": "看板名称", "看板2": "看板名称", "看板3": "看板名称",
        "指标": "核心指标", "维度": "筛选维度",
        # 9.1 数据来源
        "对象1": "数据对象", "对象2": "数据对象", "对象3": "数据对象",
        "源系统": "来源系统/来源方式",
        # 9.2 交付清单
        "底表+字段配置": "交付内容", "看板列表": "交付内容",
        "规则/审批": "交付内容", "角色权限": "交付内容",
        # 10.1 实施计划
        "字段口径确认": "工作内容", "底表搭建": "工作内容",
        "看板与权限配置": "工作内容", "试跑与培训": "工作内容",
        "确认表": "输出物", "表+字段": "输出物", "操作说明": "输出物",
        "安排试用人员": "客户配合事项", "提供枚举/样例数据": "客户配合事项",
        # 10.2 报价
        "表数量与字段": "范围说明", "看板数量": "范围说明",
        "规则数量": "范围说明", "对接复杂度": "范围说明", "运维方式": "范围说明",
        "报价或口径": "备注", "是/否": "是否本次包含", "是/否 + 系统名": "是否已有源系统",
        # 10.3 变更
        "视情况": "是否范围内", "走变更评估、单独报价": "处理方式",
        # 11.1 问题清单
        "客户方": "负责人",
        # 11.2 确认
        "客户签字确认一期范围": "确认说明",
        "客户确认报价方式": "确认说明", "客户确认接入方式": "确认说明",
        "客户确认排期": "确认说明",
        # 特殊
        "一期必做": "一期必做", "二期评估": "二期评估事项",
        "按计划交付": "处理方式", "培训次数": "范围说明", "培训说明": "备注",
        "确认验收": "输出物",
        "汇总内容": "自动带出/汇总内容",
        "需求1": "需求项", "需求2": "需求项", "需求3": "需求项", "需求4": "需求项",
        "本次必做": "范围说明",
    }

    skip_keys = {'来源', '备注', '注意', '自动带出/汇总内容', '一期必做', '确认说明', '处理方式', '是否范围内'}

    # 直接替换：单元格文字恰好等于 {{key}} 时才替换
    filled_count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text.startswith('{{') and cell_text.endswith('}}') and cell_text.count('{{') == 1:
                    ph_key = cell_text[2:-2]
                    resolved_key = None
                    if ph_key in key_map:
                        resolved_key = ph_key
                    elif ph_key in alias and alias[ph_key] in key_map:
                        resolved_key = alias[ph_key]
                    if resolved_key:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.text == cell_text or run.text == '{{' + ph_key + '}}':
                                    run.text = key_map[resolved_key]
                                    filled_count += 1
                                    break
                    elif ph_key not in skip_keys:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.text == cell_text or run.text == '{{' + ph_key + '}}':
                                    run.text = '⚠️ 待确认'
                                    break

    # 保存文件
    client_id_str = str(client_id)
    output_dir = OUTPUTS_BASE + "/" + client_id_str
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"技术路线及报价方案_{ctx['customer_name']}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    file_url = "/outputs/" + client_id_str + "/" + filename

    # 更新版本数组（同时构建 requirementData 供 Step5 使用）
    now = datetime.now().isoformat()
    versions = client.get("step4_technical_versions") or []
    if isinstance(versions, str):
        try:
            versions = json.loads(versions)
        except:
            versions = []
    next_ver = len(versions) + 1
    # 从 ctx 构建 requirementData（供 Step5 generate-demo 使用）
    requirement_data = _build_requirement_data_from_context(ctx)
    versions.append({
        "version": next_ver,
        "content": {"wordText": word_text_raw, "requirementData": requirement_data},
        "created_at": now,
        "file_url": file_url,
        "filename": filename,
    })
    conn2 = get_db()
    cursor2 = conn2.cursor()
    cursor2.execute(
        "UPDATE clients SET step4_technical_versions = ? WHERE id = ?",
        (json.dumps(versions, ensure_ascii=False), client_id)
    )
    conn2.commit()
    conn2.close()

    return {
        "success": True,
        "file_url": file_url,
        "filename": filename,
        "version": next_ver,
    }


# ==================== Step4 技术文档 HTML 生成（文档风，11章26表）====================
@app.post("/api/step4/generate-technical")
async def generate_step4_technical(body: dict, user: dict = Depends(require_auth)):
    """按钮②（技术文档模式）：生成技术路线及报价方案 HTML（文档风，可下载 Word）"""
    client_id = body.get("client_id")
    if not client_id:
        raise HTTPException(status_code=400, detail="client_id is required")

    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM clients WHERE id = ? AND user_id = ?", (client_id, user["user_id"]))
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="客户不存在")

    client = dict(row)

    # 加载上下文
    ctx = _load_client_context(client, user)

    # 构建 requirementData 供 prompt 使用
    requirement_data = _build_requirement_data_from_context(ctx)
    req_json_str = json.dumps(requirement_data, ensure_ascii=False)

    # 调用 CodeBuddy 生成技术文档 HTML
    techdoc_prompt = STEP4_TECHDOC_PROMPT.replace("{requirement_data}", req_json_str)
    html_content = None
    for attempt in range(2):
        ai_result = call_codebuddy(STEP4_TECHDOC_PROMPT, techdoc_prompt, max_tokens=12000)
        raw = ai_result["content"]
        record_ai_tokens(client_id, ai_result["usage"]["total_tokens"])
        if raw.startswith("Error:"):
            if attempt == 0:
                continue
            return {"success": False, "error": raw}
        html_content = raw
        break

    if not html_content:
        return {"success": False, "error": "技术文档生成失败，请稍后重试"}

    # 保存文件
    client_id_str = str(client_id)
    output_dir = OUTPUTS_BASE + "/" + client_id_str
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"技术路线及报价方案_{ctx['customer_name']}_{timestamp}.html"
    filepath = os.path.join(output_dir, filename)

    # 注入访客跟踪脚本（与 presales HTML 保持一致）
    tracking_js = (
        '<script>'
        '(function(){'
        'var vid=localStorage.getItem("pa_vid")||(localStorage.setItem("pa_vid","v"+Math.random().toString(36).substr(2,9)+Date.now()),localStorage.getItem("pa_vid"));'
        'var fu=location.pathname;'
        'var cid=' + str(client_id) + ';'
        'var sd=0;'
        'function tk(a,e){var p={visitor_id:vid,file_url:fu,client_id:cid,referer:document.referrer,scroll_depth:sd};if(e)Object.assign(p,e);fetch("/api/track/"+a,{method:"POST",headers:{"Content-Type":"application/json"},body:JSON.stringify(p)}).catch(function(){});}'
        'tk("visit");'
        'window.addEventListener("scroll",function(){var h=document.documentElement,b=document.body,pct=Math.round(100*(h.scrollTop||b.scrollTop)/(h.scrollHeight-h.clientHeight));if(pct>sd)sd=pct;},{passive:true});'
        'setInterval(function(){tk("heartbeat");},30000);'
        '})();'
        '</script>'
    )
    if "</body>" in html_content.lower():
        html_content = html_content.replace("</body>", tracking_js + "</body>")
    else:
        html_content += tracking_js

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html_content)

    file_url = "/outputs/" + client_id_str + "/" + filename

    # 更新版本数组
    now = datetime.now().isoformat()
    versions = client.get("step4_technical_versions") or []
    if isinstance(versions, str):
        try:
            versions = json.loads(versions)
        except:
            versions = []
    next_ver = len(versions) + 1
    versions.append({
        "version": next_ver,
        "content": {"htmlContent": html_content, "requirementData": requirement_data},
        "created_at": now,
        "file_url": file_url,
        "filename": filename,
    })
    conn2 = get_db()
    cursor2 = conn2.cursor()
    cursor2.execute(
        "UPDATE clients SET step4_technical_versions = ? WHERE id = ?",
        (json.dumps(versions, ensure_ascii=False), client_id)
    )
    conn2.commit()
    conn2.close()

    return {
        "success": True,
        "file_url": file_url,
        "filename": filename,
        "version": next_ver,
        "html_content": html_content,
    }


# ==================== Step4 Word .docx 生成 ====================
@app.post("/api/step4/generate-docx")
async def generate_step4_docx(body: dict, user: dict = Depends(require_auth)):
    """用 python-docx 生成真正的 .docx，对齐 11 节主结构"""
    client_id = body.get("client_id")
    if not client_id:
        return {"success": False, "error": "缺少 client_id"}
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM clients WHERE id = ? AND user_id = ?", (client_id, user["user_id"]))
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="客户不存在")
    client = dict(row)
    for field in ("step4_presales_versions",):
        if client.get(field) and isinstance(client[field], str):
            try:
                client[field] = json.loads(client[field])
            except:
                pass
    versions = client.get("step4_presales_versions") or []
    if not versions:
        return {"success": False, "error": "请先生成 Step4 售前方案"}
    latest = versions[-1]
    wc = latest.get("content", {})
    if isinstance(wc, str):
        try:
            wc = json.loads(wc)
        except:
            return {"success": False, "error": "方案内容格式异常"}
    rd = wc.get("requirementData") or {}
    cust_name = rd.get("meta", {}).get("customerName", client.get("name", "客户"))
    sp_name = user.get("provider_name", "{{SP_FULL_NAME}}")
    output_date = datetime.now().strftime("%Y年%m月%d日")
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return {"success": False, "error": "服务器未安装 python-docx"}

    def ls(v):
        return v if isinstance(v, list) else ([v] if v else [])

    def ss(v):
        if not v or isinstance(v, bool):
            return ""
        if isinstance(v, (int, float)):
            return str(v)
        if isinstance(v, str):
            return v.strip()
        if isinstance(v, list):
            return "、".join(ss(x) for x in v if x)
        if isinstance(v, dict):
            return v.get("title") or v.get("name") or v.get("item") or str(v)
        return str(v)

    def yn(v):
        if isinstance(v, bool):
            return "是" if v else "否"
        t = ss(v).lower()
        if t in ("true", "yes", "是"):
            return "是"
        if t in ("false", "no", "否"):
            return "否"
        return ss(v)

    def shd(cell, fill):
        tc = cell._tc
        pr = tc.get_or_add_tcPr()
        s = OxmlElement("w:shd")
        s.set(qn("w:val"), "clear")
        s.set(qn("w:color"), "auto")
        s.set(qn("w:fill"), fill)
        pr.append(s)

    def add_t(headers, rows):
        if not rows:
            return
        tb = doc.add_table(rows=1, cols=len(headers))
        tb.style = "Table Grid"
        for i, h in enumerate(headers):
            c = tb.rows[0].cells[i]
            c.text = ""
            r = c.paragraphs[0].add_run(ss(h))
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(255, 255, 255)
            shd(c, "1E5AFF")
        for ri, row in enumerate(rows):
            cs = tb.add_row().cells
            for ci, val in enumerate(row):
                if ci < len(cs):
                    cs[ci].text = ss(val)
                    if cs[ci].paragraphs[0].runs:
                        cs[ci].paragraphs[0].runs[0].font.size = Pt(10)
                    if ri % 2 == 1:
                        shd(cs[ci], "F7F8FA")
        doc.add_paragraph("").paragraph_format.space_after = Pt(2)

    def h1(t):
        p = doc.add_heading(t, level=1)
        for r in p.runs:
            r.font.size = Pt(15)
            r.font.color.rgb = RGBColor(0x1E, 0x5A, 0xFF)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)

    def h2(t):
        p = doc.add_heading(t, level=2)
        for r in p.runs:
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x18, 0x22, 0x35)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

    def txt(t, bold=False, sz=10, color=None, align=None):
        p = doc.add_paragraph()
        if align:
            p.alignment = align
        r = p.add_run(ss(t))
        r.font.size = Pt(sz)
        r.bold = bold
        if color:
            r.font.color.rgb = RGBColor(*color)
        p.paragraph_format.space_after = Pt(4)

    def sp():
        p = doc.add_paragraph("")
        p.paragraph_format.space_after = Pt(6)

    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21)
    s.page_height = Cm(29.7)
    s.left_margin = Cm(2)
    s.right_margin = Cm(2)
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)

    # Cover
    txt("企业微信定制开发", bold=True, sz=20, color=(0x1E, 0x5A, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    txt("需求确认 & 方案设计表", bold=True, sz=16, color=(0x18, 0x22, 0x35), align=WD_ALIGN_PARAGRAPH.CENTER)
    scenario = ss(rd.get("meta", {}).get("mainScenario", ""))
    cust_name_val = ss(rd.get("meta", {}).get("customerName", cust_name))
    txt(cust_name_val + "｜" + scenario, sz=11, color=(0x66, 0x70, 0x85), align=WD_ALIGN_PARAGRAPH.CENTER)
    sp()

    meta = rd.get("meta", {})
    facts = rd.get("customerFacts", {})
    industry = ss(meta.get("industry", ""))
    scale = ss(meta.get("companyScale", ""))
    roles3 = ls(facts.get("involvedRoles", []))
    existing_tools = ls(facts.get("existingTools", []))
    add_t(
        ["信息项", "内容"],
        [
            ("客户名称", cust_name_val),
            ("所属行业", industry),
            ("公司规模", scale),
            ("项目场景", scenario),
            ("服务商名称", sp_name),
            ("文档版本", ss(meta.get("version", "V1.0"))),
            ("编制日期", output_date),
        ]
    )
    notice = doc.add_paragraph()
    r = notice.add_run("本文件用于服务商与客户共同确认需求范围、一期边界、智能表格口径、权限与待确认问题。客户未确认内容不得写入一期交付承诺。")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x7C, 0x59, 0x14)
    pr = notice._p.get_or_add_pPr()
    ns = OxmlElement("w:shd")
    ns.set(qn("w:val"), "clear")
    ns.set(qn("w:color"), "auto")
    ns.set(qn("w:fill"), "FFFBE6")
    pr.append(ns)
    notice.paragraph_format.space_after = Pt(12)

    # 一
    h1("一、客户基础信息与当前现状")
    h2("1.1 客户基础信息确认表")
    current_state = ss(facts.get("customerCurrentState", ""))
    current_state_status = "已确认" if current_state else "⚠️ 待确认"
    add_t(
        ["信息项", "确认内容", "状态"],
        [
            ("行业", industry, "已确认"),
            ("规模", scale, "已确认"),
            ("现有系统", "、".join(existing_tools) if existing_tools else "⚠️ 待确认", "⚠️ 待确认"),
            ("核心团队角色", "、".join(roles3) if roles3 else "⚠️ 待确认", "⚠️ 待确认"),
            ("当前业务现状", current_state, current_state_status),
        ]
    )

    h2("1.2 当前业务运转方式")
    pains3 = ls(rd.get("painPoints", []))
    bp3 = rd.get("businessProcess", {})
    cur3 = ls(bp3.get("currentFlow", []))
    tgt3 = ls(bp3.get("targetFlow", []))
    proc3 = []
    for i, cf in enumerate(cur3):
        ti = tgt3[i] if i < len(tgt3) else {}
        proc3.append([
            ss(cf.get("stepName", "")),
            ss(cf.get("currentMethod", "")),
            ss(ti.get("systemAction", cf.get("currentMethod", ""))),
            ss(cf.get("role", ""))
        ])
    if proc3:
        add_t(["环节", "当前做法", "企业微信动作", "操作角色"], proc3)
    else:
        txt("（业务流程待从 Step3 补充）")

    h2("1.3 核心痛点与优先级")
    if pains3:
        pr2 = []
        for i, p in enumerate(pains3):
            ev = ss(p.get("evidence", ""))[:40]
            pr2.append([
                str(i + 1),
                ss(p.get("title", "")),
                ss(p.get("description", "")),
                ss(p.get("businessImpact", "")),
                ss(p.get("priority", "")),
                ev
            ])
        add_t(["痛点编号", "痛点（业务语言）", "描述", "业务影响", "优先级", "依据"], pr2)
    else:
        txt("（未识别出痛点）")

    # 二
    h1("二、场景类型判断与方案边界")
    h2("2.1 场景类型判断")
    if scenario:
        txt("本项目应识别为「" + scenario + "」场景。")
    else:
        txt("⚠️ 场景类型待确认。")

    h2("2.2 方案边界")
    scope3 = rd.get("scope", {})
    bdry = []
    for it in ls(scope3.get("phaseOne", [])):
        bdry.append(["✅ 本期范围内", ss(it.get("item", "")), ss(it.get("reason", ""))])
    for it in ls(scope3.get("phaseTwo", [])):
        bdry.append(["🔄 二期扩展", ss(it.get("item", "")), ss(it.get("prerequisites", ""))])
    for it in ls(scope3.get("notRecommended", [])):
        bdry.append(["❌ 暂不建议", ss(it.get("item", "")), ss(it.get("reason", ""))])
    if bdry:
        add_t(["类型", "内容", "前提/原因"], bdry)
    else:
        txt("（一期/二期边界待确认）")

    # 三
    h1("三、需求理解与优先级确认")
    h2("3.1 需求清单")
    reqs3 = ls(rd.get("requirements", []))
    if reqs3:
        rr = []
        for r in reqs3:
            rr.append([
                ss(r.get("requirementName", "")),
                ss(r.get("customerExpression", "")),
                ss(r.get("businessTranslation", "")),
                ss(r.get("priority", "")),
                ss(r.get("phase", "")),
                ss(r.get("confirmedStatus", ""))
            ])
        add_t(["需求", "客户表达", "业务语言", "优先级", "一期/二期", "确认状态"], rr)
    else:
        txt("（需求待从 Step2/3 补充）")

    h2("3.2 客户原话与业务翻译")
    oqt_src = []
    for r in reqs3:
        ce = ss(r.get("customerExpression", ""))
        if ce:
            oqt_src.append([
                ce,
                ss(r.get("businessTranslation", "")),
                ss(r.get("confirmedStatus", "⚠️ 待确认"))
            ])
    if oqt_src:
        add_t(["客户原话", "业务翻译", "确认状态"], oqt_src)
    elif reqs3:
        txt("（原话翻译待补充）")
    else:
        txt("（需求原话待补充）")

    # 四
    h1("四、业务流程设计")
    h2("4.1 当前流程与目标流程")
    if proc3:
        add_t(["环节", "当前做法", "企业微信动作", "操作角色"], proc3)
    else:
        txt("（业务流程待补充）")

    h2("4.2 流程节点确认表")
    nodes3 = ls(bp3.get("processNodes", []))
    if nodes3:
        nr = []
        for i, n in enumerate(nodes3):
            nr.append([
                str(i + 1),
                ss(n.get("nodeName", "")),
                ss(n.get("responsibleRole", "")),
                ss(n.get("input", "")),
                ss(n.get("output", "")),
                yn(n.get("reminderNeeded", False))
            ])
        add_t(["序号", "节点名称", "操作角色", "输入信息", "输出结果", "是否提醒"], nr)
    else:
        txt("（流程节点待补充）")

    # 五
    h1("五、企业微信方案总览")
    mods = ls(rd.get("moduleRecommendation", []))
    if mods:
        mr = []
        for m in mods:
            mr.append([
                ss(m.get("moduleType", "")),
                ss(m.get("moduleName", "")),
                ss(m.get("notes", ""))
            ])
        add_t(["能力层", "模块名称", "本项目使用方式"], mr)
    else:
        txt("企业微信入口 + 智能表格数据底座 + 审批 / 自动化 / 权限 / 看板轻量定制方案。")

    # 六
    st3 = rd.get("smartTableSpec", {})
    tabs3 = ls(st3.get("confirmedTables", []))
    h1("六、智能表格交付设计")
    h2("6.1 智能表格总览")
    if tabs3:
        tr = []
        for t in tabs3:
            phase_val = ss(t.get("phase", ""))
            phase_label = "✅ 一期" if "一期" in phase_val and "二期" not in phase_val else "🔄 二期"
            tr.append([
                ss(t.get("tableName", "")),
                ss(t.get("tablePurpose", "")),
                "、".join(ls(t.get("roles", []))),
                phase_label
            ])
        add_t(["表名", "用途", "使用角色", "一期/二期"], tr)
    else:
        txt("（智能表格待从 Step2 xlsx 补充）")

    h2("6.2 核心字段设计表")
    fmap = {}
    for ft in ls(st3.get("fieldsByTable", [])):
        ft_name = ss(ft.get("tableName", ""))
        fmap[ft_name] = ls(ft.get("fields", []))
    for t in tabs3:
        tn = ss(t.get("tableName", ""))
        if not tn:
            continue
        txt("  " + tn, bold=True, sz=11, color=(0x18, 0x22, 0x35))
        flds = ls(fmap.get(tn, []))
        if flds:
            fr = []
            for f in flds:
                fr.append([
                    ss(f.get("fieldName", "")),
                    ss(f.get("fieldType", "")),
                    yn(f.get("required", False)),
                    ss(f.get("source", "")),
                    ss(f.get("rule", ""))
                ])
            add_t(["字段名称", "字段类型", "必填", "维护角色", "规则说明"], fr)
        else:
            txt("（字段待补充）")

    h2("6.3 表间关联关系")
    rels = ls(st3.get("relations", []))
    if rels:
        rr2 = []
        for r in rels:
            rr2.append([
                ss(r.get("mainTable", "")),
                ss(r.get("relatedTable", "")),
                ss(r.get("relationField", "")),
                ss(r.get("autoFill", "")),
                ss(r.get("note", ""))
            ])
        add_t(["主表", "关联表", "关联字段", "自动带出", "注意事项"], rr2)
    else:
        txt("（表间关联关系待补充）")

    # 七
    h1("七、审批与自动化设计")
    autos3 = ls(st3.get("automations", []))
    if autos3:
        ar = []
        for a in autos3:
            ar.append([
                ss(a.get("name", "")),
                ss(a.get("trigger", "")),
                ss(a.get("action", "")),
                ss(a.get("notifyTarget", "")),
                ss(a.get("priority", "中"))
            ])
        add_t(["规则名称", "触发条件", "执行动作", "通知对象", "优先级"], ar)
    else:
        txt("（自动化规则待补充，可结合审批流配置）")

    if mods:
        phase_one_mods = [m for m in mods if "一期" in ss(m.get("phase", ""))]
        if phase_one_mods:
            names = "、".join(ss(m.get("moduleName", "")) for m in phase_one_mods)
            txt("建议对「" + names + "」配置审批流程，具体审批节点待与客户确认。")

    # 八
    h1("八、权限与数据看板设计")
    h2("8.1 权限矩阵")
    perms3 = ls(st3.get("permissions", []))
    if perms3:
        pr3 = []
        for p in perms3:
            pr3.append([
                ss(p.get("role", "")),
                ss(p.get("addScope", "")),
                ss(p.get("viewScope", "")),
                ss(p.get("editableFields", "")),
                ss(p.get("sensitiveFields", ""))
            ])
        add_t(["角色", "新增权限", "查看范围", "可编辑字段", "敏感字段"], pr3)
    elif roles3:
        pr3 = []
        table_names = "、".join(ss(t.get("tableName", "")) for t in tabs3[:3])
        for r in roles3:
            pr3.append([ss(r), "本职范围内", table_names, "本职相关字段", "——"])
        add_t(["角色", "新增权限", "查看范围", "可编辑字段", "敏感字段"], pr3)
        txt("（以上为推断值，建议根据实际情况调整）", sz=9, color=(0x98, 0xA2, 0xB3))
    else:
        txt("（权限矩阵待补充）")

    h2("8.2 数据看板设计")
    dashes = ls(st3.get("dashboards", []))
    if dashes:
        dr = []
        for d in dashes:
            dr.append([
                ss(d.get("dashboardName", "")),
                ss(d.get("users", "")),
                ss(d.get("metrics", "")),
                ss(d.get("filters", ""))
            ])
        add_t(["看板名称", "使用对象", "核心指标", "筛选维度"], dr)
    else:
        txt("（数据看板待规划，建议以「销售情况看板」起步）")

    # 九
    h1("九、数据来源、系统对接与交付边界")
    h2("9.1 数据来源与接入方式")
    if tabs3:
        dr2 = []
        for t in tabs3:
            phase_val = ss(t.get("phase", ""))
            phase_one = "一期" in phase_val and "二期" not in phase_val
            phase_two = "二期" in phase_val
            dr2.append([
                ss(t.get("tableName", "")),
                "客户现有 Excel/企微智能表格",
                "一期接入" if phase_one else "⚠️ 待评估",
                "ERP对接" if phase_two else "——"
            ])
        add_t(["数据对象", "来源", "一期处理方式", "二期评估"], dr2)
    else:
        txt("（数据来源待补充）")

    h2("9.2 交付边界说明")
    for it in ls(scope3.get("phaseOne", [])):
        txt("✅ " + ss(it.get("item", "")) + " — " + ss(it.get("reason", "")))
    for it in ls(scope3.get("phaseTwo", [])):
        item = ss(it.get("item", ""))
        prereq = ss(it.get("prerequisites", ""))
        if prereq:
            txt("🔄 " + item + "（前提：" + prereq + "）")
        else:
            txt("🔄 " + item)

    # 十
    h1("十、实施计划、报价口径与变更机制")
    h2("10.1 实施计划")
    impl3 = ls(rd.get("implementationPlanTable", []))
    if impl3:
        ir = []
        for p in impl3:
            ir.append([
                ss(p.get("phase", "")),
                ss(p.get("workContent", "")),
                ss(p.get("customerCooperation", "")),
                ss(p.get("output", ""))
            ])
        add_t(["阶段", "工作内容", "客户配合事项", "输出物"], ir)
    else:
        add_t(["阶段", "工作内容", "客户配合事项", "输出物"], [
            ("需求确认", "确认需求范围、一期二期边界、智能表格结构", "提供需求，参与评审", "需求确认文档"),
            ("智能表格搭建", "按确认的字段设计搭建智能表格", "参与字段确认、提供基础数据", "可用的智能表格 Demo"),
            ("规则配置", "配置自动化规则、审批流、权限", "参与规则评审、测试确认", "配置完成的规则"),
            ("试运行与优化", "上线试运行，收集反馈并优化", "提供试运行数据、反馈问题", "上线文档+优化建议"),
        ])

    h2("10.2 报价口径建议（只写口径，不写金额）")
    qt3 = ls(rd.get("quoteScopeTable", []))
    if qt3:
        qr = []
        for q in qt3:
            qr.append([
                ss(q.get("feeModule", "")),
                ss(q.get("scope", "")),
                yn(q.get("included", "")),
                ss(q.get("note", ""))
            ])
        add_t(["费用模块", "范围说明", "是否本次包含", "备注"], qr)
    else:
        txt("（报价口径待与服务商确认，⚠️ 本文仅提供口径说明，不写具体金额）", color=(0xFF, 0x6B, 0x35))

    h2("10.3 范围变更机制")
    chg = ls(rd.get("changeManagementTable", []))
    if chg:
        chr_list = []
        for c in chg:
            chr_list.append([
                ss(c.get("item", "")),
                yn(c.get("inScope", "")),
                ss(c.get("handling", ""))
            ])
        add_t(["事项", "是否范围内", "处理方式"], chr_list)
    else:
        add_t(["事项", "是否范围内", "处理方式"], [
            ("新增智能表格字段", "否", "评估后进入变更单处理"),
            ("跨表关联新增", "否", "评估工作量后单独报价"),
            ("自动化规则增加", "否", "评估后进入变更单处理"),
        ])

    # 十一
    h1("十一、待客户确认问题与签署")
    h2("11.1 待确认问题")
    oqs3 = ls(rd.get("openQuestions", []))
    if oqs3:
        oq3r = []
        for i, o in enumerate(oqs3):
            oq3r.append([
                str(i + 1),
                ss(o.get("question", "")),
                ss(o.get("owner", "服务商")),
                ss(o.get("priority", "中"))
            ])
        add_t(["编号", "问题", "负责人", "优先级"], oq3r)
    else:
        txt("（暂无待确认问题，或从 Step3 沟通记录补充）")

    h2("11.2 确认说明")
    note_p = doc.add_paragraph()
    note_r = note_p.add_run("声明：本文档为售前需求确认稿，以上待确认项（标注 ⚠️）未纳入一期交付承诺。ERP 系统对接、AI 智能填报、历史数据清洗等未经客户明确确认的能力，不在本次一期交付范围内，建议以二期独立项目评估。感谢贵司信任。")
    note_r.font.size = Pt(10)
    note_r.font.color.rgb = RGBColor(0x34, 0x40, 0x54)
    note_p.paragraph_format.space_after = Pt(12)

    h2("11.3 签署")
    add_t(["确认方", "姓名/职务", "确认意见", "日期"], [
        ["客户方", "", "", ""],
        ["服务商方", "", "", ""]
    ])

    # 保存
    import os
    filename = cust_name_val + "_需求确认与方案设计表.docx"
    savedir = "/var/www/provider-assist/step4-docx"
    os.makedirs(savedir, exist_ok=True)
    filepath = os.path.join(savedir, filename)
    doc.save(filepath)
    return {"success": True, "url": "/step4-docx/" + filename, "filename": filename, "customer": cust_name_val}


STEP5_SCHEMA_SYSTEM_PROMPT = """你是一个企业微信智能表格搭建专家。请基于需求确认数据中的 smartTableSpec，生成可直接用于建表的 JSON Schema。

请直接输出 JSON，不要任何解释文字，不要 markdown 代码块。"""


STEP5_SCHEMA_USER_PROMPT = """【smartTableSpec（需求结构化中的智能表格规格）】
```json
{smart_table_spec}
```

【scope（交付边界）】
```json
phaseOne：{phase_one_scope}
phaseTwo：{phase_two_scope}
notRecommended：{not_recommended_scope}
```

请严格按以下 JSON Schema 输出（直接输出 JSON，不要任何前缀）：

【强制要求 - 最高优先级】每个子表的 sample_records 必须恰好填写 10 条真实业务数据，数据要贴合行业和客户场景，禁止填虚假或无关数据。少于 10 条视为严重不合格，必须补齐。输出前必须自检每张表的 sample_records 数量是否 == 10。

{
  "doc_name": "智能表格名称",
  "sheets": [
    {
      "sheet_name": "子表名称",
      "fields": [
        {
          "field_title": "字段名",
          "field_type": "文本|多行文本|单选|多选|数字|金额|日期|日期时间|人员|手机|附件|图片|关联记录|公式|自动编号|进度|勾选|URL",
          "required": true
        }
      ],
      "sample_records": [
        { "字段1": "示例值1", "字段2": "示例值2", "字段3": "示例值3" },
        { "字段1": "示例值1", "字段2": "示例值2", "字段3": "示例值3" },
        { "字段1": "示例值1", "字段2": "示例值2", "字段3": "示例值3" },
        { "字段1": "示例值1", "字段2": "示例值2", "字段3": "示例值3" },
        { "字段1": "示例值1", "字段2": "示例值2", "字段3": "示例值3" },
        { "字段1": "示例值1", "字段2": "示例值2", "字段3": "示例值3" },
        { "字段1": "示例值1", "字段2": "示例值2", "字段3": "示例值3" },
        { "字段1": "示例值1", "字段2": "示例值2", "字段3": "示例值3" },
        { "字段1": "示例值1", "字段2": "示例值2", "字段3": "示例值3" },
        { "字段1": "示例值1", "字段2": "示例值2", "字段3": "示例值3" }
      ]
    }
  ]
}

【填写规则 - 必须遵守】
1. 只包含 phase="一期" 的表，不含二期评估。
2. field_type 只能取以下中文名之一：**文本 / 多行文本 / 数字 / 单选 / 多选 / 日期时间 / 金额 / 百分比 / 进度 / 手机 / 邮箱 / 链接 / 勾选 / 人员 / 附件 / 图片 / 关联记录 / 公式 / 自动编号**。
3. 🆕 **字段数贴合真实业务，不做轻量删减**：核心业务表 12-20 个字段，覆盖该对象完整信息（基础信息+状态+时间+责任人+金额/数量+备注等）；核心表字段少于 10 个视为不合格。仅字典表/辅助表可少（6-10 个）。
4. 🆕 **覆盖核心表清单**：smartTableSpec / 需求报告 coreTablesExpected 里列出的每一张核心表都必须建，缺表视为不合格。
5. 【sample_records 恰好 10 条，且 10 条内容各不相同】——禁止把同一条复制 10 遍，禁止填虚假/无关数据。数据必须贴合本行业本客户场景，10 条应体现不同记录（不同客户/项目/日期/金额等）。
6. 字段命名专业、贴合行业术语；一张表聚焦一个业务对象。

## 输出前自检（不通过就重写）
1. 合法 JSON：无注释、无尾逗号、无代码块。
2. 每个子表 sample_records 正好 10 条且互不相同。
3. field_type 全部来自上述清单。
4. 🆕 核心业务表字段数 ≥12；coreTablesExpected 中的表张张都在，无遗漏。

直接输出有效 JSON，不要 markdown 代码块包裹。"""


@app.post("/api/step5/generate-demo")
async def generate_step5_demo(body: dict, user: dict = Depends(require_auth)):
    """生成 Step5 企业微信智能表格 JSON Schema（从 Step4 smartTableSpec 派生）"""
    client_id = body.get("client_id")
    if not client_id:
        return {"success": False, "error": "缺少 client_id"}

    client = db_get_client(client_id)
    if not client:
        return {"success": False, "error": "客户不存在"}

    import logging
    logger = logging.getLogger("uvicorn")
    logger.warning(f"[Step5Demo] client_id={client_id} client_name={client.get('name')} client_keys={list(client.keys())}")
    logger.warning(f"[Step5Demo] step2_schema type={type(client.get('step2_schema'))} step3_summary type={type(client.get('step3_summary'))}")
    s2 = client.get("step2_schema")
    if isinstance(s2, dict):
        logger.warning(f"[Step5Demo] step2_schema sheets={s2.get('sheets')}")
    elif isinstance(s2, str) and s2:
        try:
            parsed = json.loads(s2)
            logger.warning(f"[Step5Demo] step2_schema parsed sheets={parsed.get('sheets')}")
        except:
            logger.warning(f"[Step5Demo] step2_schema parse failed, raw={s2[:200]}")

    # 从 Step4 产物1的最新版本获取 requirementSolutionData.smartTableSpec
    presales_versions = client.get("step4_presales_versions") or []
    logger.warning(f"[Step5Demo] presales_versions count={len(presales_versions)}")
    if not presales_versions:
        return {"success": False, "error": "请先生成 Step4 售前方案"}

    latest = presales_versions[-1]
    content = latest.get("content") or {}
    logger.warning(f"[Step5Demo] content keys={list(content.keys()) if isinstance(content, dict) else type(content)}")
    requirement_data = content.get("requirementData")
    logger.warning(f"[Step5Demo] requirementData from presales={bool(requirement_data)}")

    # Fallback: 从 step2_schema 和 step3_summary 构建 requirementData
    if not requirement_data:
        step2_schema = client.get("step2_schema") or {}
        step3_summary = client.get("step3_summary") or {}
        if isinstance(step2_schema, str):
            try: step2_schema = json.loads(step2_schema)
            except: step2_schema = {}
        if isinstance(step3_summary, str):
            try: step3_summary = json.loads(step3_summary)
            except: step3_summary = {}

        confirmed_tables = []
        fields_by_table = []
        sheets = step2_schema.get("sheets") or []
        for s in sheets:
            name = s.get("name", "") or s.get("sheet_name", "")
            cols = s.get("columns") or s.get("fields") or []
            if not name: continue
            confirmed_tables.append({"tableName": name, "tablePurpose": "业务管理表", "source": "xlsx", "phase": "一期", "roles": []})
            fields = []
            for c in cols:
                fname = c.get("name", "") or c.get("field_title", "") or c.get("fieldName", "")
                ftype = c.get("type", "") or c.get("field_type", "") or "文本"
                if not fname: continue
                ft_map = {"文本":"文本","text":"文本","多行文本":"多行文本","NUMBER":"数字","数字":"数字","SINGLE_SELECT":"单选","单选":"单选","DATE":"日期","日期":"日期","DATE_TIME":"日期时间","金额":"金额","CURRENCY":"金额"}
                fields.append({"fieldName": fname, "fieldType": ft_map.get(ftype, "文本"), "required": False, "rule": "", "source": "xlsx"})
            if fields:
                fields_by_table.append({"tableName": name, "fields": fields})

        # 构建 phase_one_scope
        _raw_s3 = step3_summary.get("_raw") or {}
        _cn = step3_summary.get("confirmedNeeds") or []
        _pp = step3_summary.get("painPoints") or []
        _s4i = {}
        step4_input = client.get("step4_input_draft") or {}
        if isinstance(step4_input, str):
            try: _s4i = json.loads(step4_input)
            except: _s4i = {}
        else: _s4i = step4_input or {}
        _cn = _cn or _s4i.get("confirmedNeeds") or []
        _pp = _pp or _s4i.get("painPoints") or []
        phase_one_scope = []
        for n in (_cn if isinstance(_cn, list) else [_cn]):
            title = n.get("title") or n.get("name") or "" if isinstance(n, dict) else ""
            reason = n.get("description") or "客户确认需求" if isinstance(n, dict) else ""
            if isinstance(n, str): title, reason = n, "客户确认需求"
            if title and len(phase_one_scope) < 8: phase_one_scope.append({"item": title, "reason": reason, "deliveryForm": "智能表格"})
        for p in (_pp if isinstance(_pp, list) else [_pp]):
            title = p.get("title") or "" if isinstance(p, dict) else ""
            reason = p.get("description") or "痛点解决" if isinstance(p, dict) else ""
            if isinstance(p, str): title, reason = p, "痛点解决"
            if title and len(phase_one_scope) < 8: phase_one_scope.append({"item": title, "reason": reason, "deliveryForm": "智能表格"})
        if not phase_one_scope and _s4i.get("customerCurrentState"):
            phase_one_scope = [{"item": "核心业务管理", "reason": _s4i["customerCurrentState"][:80], "deliveryForm": "智能表格"}]
        elif not phase_one_scope and _raw_s3.get("customerCurrentState"):
            phase_one_scope = [{"item": "核心业务管理", "reason": _raw_s3["customerCurrentState"][:80], "deliveryForm": "智能表格"}]

        requirement_data = {
            "smartTableSpec": {
                "scenarioComplexity": "简单流程型",
                "confirmedTables": confirmed_tables,
                "fieldsByTable": fields_by_table,
                "suggestedTables": [],
                "phaseTwoTables": []
            },
            "scope": {
                "phaseOne": phase_one_scope,
                "phaseTwo": [],
                "notRecommended": []
            }
        }
        logger.warning(f"[Step5Demo] fallback built requirement_data confirmedTables={len(confirmed_tables)}")

    # ===== 兼容 _raw 嵌套格式：无论 requirement_data 从哪来，都尝试从 _raw 补全 scope =====
    # _raw 可能是 dict，也可能是 JSON 字符串，需要统一解析
    _raw = requirement_data.get("_raw") or {}
    if not _raw:
        s3 = client.get("step3_summary") or {}
        if isinstance(s3, str):
            try: s3 = json.loads(s3)
            except: s3 = {}
        _raw = s3.get("_raw") or {}
    if isinstance(_raw, str):
        try: _raw = json.loads(_raw)
        except: _raw = {}
    scope = requirement_data.get("scope") or {}
    if not scope.get("phaseOne") and _raw:
        # 从 _raw.painPoints 提取 phase_one_scope
        pts = _raw.get("painPoints") or []
        phase_one = []
        for p in (pts if isinstance(pts, list) else [pts]):
            title = (p.get("title") or p.get("name") or "") if isinstance(p, dict) else ""
            desc = (p.get("description") or "") if isinstance(p, dict) else ""
            if title:
                phase_one.append({"item": title, "reason": desc[:60] if desc else "客户痛点", "deliveryForm": "智能表格"})
        if phase_one:
            scope["phaseOne"] = phase_one
            logger.warning(f"[Step5Demo] _raw 补全 phaseOne: {len(phase_one)} 项")
        # 如果仍空，用 customerCurrentState 作为兜底
        if not scope.get("phaseOne") and _raw.get("customerCurrentState"):
            scope["phaseOne"] = [{"item": "核心业务管理", "reason": _raw["customerCurrentState"][:80], "deliveryForm": "智能表格"}]
    requirement_data["scope"] = scope

    smart_table_spec = requirement_data.get("smartTableSpec") or {}
    confirmed_tables = smart_table_spec.get("confirmedTables") or []
    logger.warning(f"[Step5Demo] smart_table_spec confirmedTables={len(confirmed_tables)}, scope phaseOne={len(scope.get('phaseOne', []))}")

    # confirmedTables 为空时：让 AI 根据 scope 和客户描述自行推断表结构
    if not confirmed_tables:
        raw_ctx = ""
        if _raw.get("customerCurrentState"):
            raw_ctx += f"\n客户现状：{_raw['customerCurrentState']}"
        pts = _raw.get("painPoints") or []
        if pts:
            raw_ctx += "\n痛点：" + "；".join([
                (p.get("title","") + "：" + p.get("description","")[:30]) if isinstance(p, dict) else str(p)
                for p in pts[:5]
            ])
        if not scope.get("phaseOne"):
            logger.warning("[Step5Demo] confirmedTables 和 phaseOne 均为空，无法推断")
            return {"success": False, "error": "无法生成：缺少需求数据。请到 Step3 填写「确认需求」和「一期范围」，再重试。"}
        user_prompt = STEP5_SCHEMA_USER_PROMPT.replace(
            "{smart_table_spec}",
            json.dumps({
                "scenarioComplexity": smart_table_spec.get("scenarioComplexity", "简单流程型"),
                "confirmedTables": [],
                "scope": scope
            }, ensure_ascii=False, indent=2)
        ).replace(
            "{phase_one_scope}", json.dumps(scope.get("phaseOne") or [], ensure_ascii=False, indent=2)
        ).replace(
            "{phase_two_scope}", json.dumps(scope.get("phaseTwo") or [], ensure_ascii=False, indent=2)
        ).replace(
            "{not_recommended_scope}", json.dumps(scope.get("notRecommended") or [], ensure_ascii=False, indent=2)
        )
        user_prompt += f"\n\n【客户业务描述（请据此推断表结构）】{raw_ctx}\n\n【重要】confirmedTables 为空，请根据上方客户业务描述，自行推断需要创建哪些智能表格（建议2-4张，覆盖核心业务流程），每个表必须有实际字段名称，禁止返回空表或占位表。"
    else:
        user_prompt = STEP5_SCHEMA_USER_PROMPT.replace(
            "{smart_table_spec}", json.dumps(smart_table_spec, ensure_ascii=False, indent=2)
        ).replace(
            "{phase_one_scope}", json.dumps(scope.get("phaseOne") or [], ensure_ascii=False, indent=2)
        ).replace(
            "{phase_two_scope}", json.dumps(scope.get("phaseTwo") or [], ensure_ascii=False, indent=2)
        ).replace(
            "{not_recommended_scope}", json.dumps(scope.get("notRecommended") or [], ensure_ascii=False, indent=2)
        )

    ai_result = call_codebuddy(STEP5_SCHEMA_SYSTEM_PROMPT, user_prompt, max_tokens=25000, timeout=600)
    raw = ai_result["content"]
    record_ai_tokens(client_id, ai_result["usage"]["total_tokens"])
    schema = parse_json_response(raw)

    # 后处理：确保每表恰好 10 条样例数据
    if schema and schema.get("sheets"):
        import logging
        logger2 = logging.getLogger("uvicorn")
        # 常用填充值库（用于生成更真实的补录数据）
        _name_pool = ["张伟", "李娜", "王芳", "刘洋", "陈静", "杨勇", "赵磊", "黄丽", "周强", "吴敏"]
        _status_pool = ["进行中", "已完成", "待处理", "已取消", "暂停中"]
        _dept_pool = ["销售部", "市场部", "运营部", "技术部", "财务部", "行政部"]
        _priority_pool = ["高", "中", "低"]
        _date_pool = ["2026-07-01", "2026-07-05", "2026-07-08", "2026-07-10", "2026-07-12", "2026-07-15"]
        _amount_pool = [3500, 8200, 15000, 28000, 45000, 63000, 98000]

        def _smart_pad_record(template, idx, field_types):
            """根据字段类型生成更真实的第 idx 条补录数据"""
            result = {}
            for k, v in template.items():
                ftype = field_types.get(k, '文本')
                if isinstance(v, str):
                    if ftype == '单选' and v in _status_pool:
                        result[k] = _status_pool[idx % len(_status_pool)]
                    elif ftype == '单选' and v in _priority_pool:
                        result[k] = _priority_pool[idx % len(_priority_pool)]
                    elif ftype == '单选' and v in _dept_pool:
                        result[k] = _dept_pool[idx % len(_dept_pool)]
                    elif '名称' in k or '负责人' in k or '客户' in k:
                        result[k] = _name_pool[idx % len(_name_pool)]
                    elif '部门' in k or '部门' in k:
                        result[k] = _dept_pool[idx % len(_dept_pool)]
                    elif '状态' in k:
                        result[k] = _status_pool[idx % len(_status_pool)]
                    elif '日期' in k or '时间' in k:
                        result[k] = _date_pool[idx % len(_date_pool)]
                    elif '金额' in k or '预算' in k or '费用' in k:
                        result[k] = _amount_pool[idx % len(_amount_pool)]
                    elif '备注' in k or '说明' in k:
                        result[k] = v + f"（第{idx+1}条）"
                    else:
                        result[k] = v + f"_{idx+1}"
                elif isinstance(v, (int, float)):
                    if '金额' in k or '预算' in k:
                        result[k] = _amount_pool[idx % len(_amount_pool)]
                    else:
                        result[k] = v + idx
                else:
                    result[k] = v
            return result

        def _guess_field_type(field_title):
            """根据字段名推测类型"""
            t = field_title or ''
            if any(x in t for x in ['状态', '类型', '阶段', '等级']): return '单选'
            if any(x in t for x in ['日期', '时间']): return '日期'
            if any(x in t for x in ['金额', '预算', '费用', '报价', '成本']): return '金额'
            if any(x in t for x in ['名称', '客户', '负责人', '联系人', '员工']): return '文本'
            if any(x in t for x in ['数量', '数量', '次数']): return '数字'
            if any(x in t for x in ['备注', '说明', '描述', '内容']): return '多行文本'
            if any(x in t for x in ['手机', '电话']): return '手机'
            if any(x in t for x in ['邮箱']): return '邮箱'
            if any(x in t for x in ['URL', '链接', '网址']): return '链接'
            return '文本'

        for sheet in schema["sheets"]:
            records = sheet.get("sample_records") or []
            current = len(records)
            if current < 10 and current > 0:
                logger2.warning(f"[Step5Demo] sheet '{sheet.get('sheet_name','')}' has {current} records, padding to 10 with smart variation")
                template = records[0] if records else {}
                # 建立字段名→类型的映射（从 fields 定义获取，否则靠猜）
                field_type_map = {}
                for f in (sheet.get("fields") or []):
                    ft = f.get("field_type", "")
                    for fn in ("field_title", "fieldName", "name"):
                        if f.get(fn):
                            field_type_map[f.get(fn)] = ft
                            break
                # 如果 fields 没有定义，从 template 的值猜测
                for k in template.keys():
                    if k not in field_type_map:
                        field_type_map[k] = _guess_field_type(k)
                idx = 0
                while len(records) < 10:
                    new_record = _smart_pad_record(template, idx, field_type_map)
                    records.append(new_record)
                    idx += 1
                sheet["sample_records"] = records
                logger2.info(f"[Step5Demo] sheet '{sheet.get('sheet_name','')}' final record count: {len(records)}")
            elif current == 0:
                logger2.warning(f"[Step5Demo] sheet '{sheet.get('sheet_name','')}' has ZERO records, cannot pad - skipping")
            else:
                logger2.info(f"[Step5Demo] sheet '{sheet.get('sheet_name','')}' has {current} records (OK)")

    if not schema:
        return {"success": False, "error": "Step5 Schema 生成失败：" + (raw[:200] if raw else "空响应")}

    # 保存 JSON Schema 到 client
    db_update_client(client_id, {"step5_schema": schema})

    return {"success": True, "demo": schema}


STEP5_AGENT_PROMPT = """你是一个企业微信定制开发增项顾问。基于客户背景和已规划的智能表格方案，推荐可打包进定制开发的 AI 助手产品（企微群机器人 + 单聊机器人）。

【客户背景】
客户名称：{customer_name}
行业：{industry}
需求：{initial_demand}

【现有智能表格 Schema（已规划的一期交付内容）】
{schema_summary}

请生成 4-6 条增项建议，每条对应一个可独立交付的 AI 助手功能模块，包含：
- title：功能名称（如"智能问答机器人"）
- type：产品类型（群机器人/单聊机器人/混合）
- description：2-3 句话说明实现方式和价值
- example：该功能在当前客户场景中的具体对话示例（机器人收到什么消息、返回什么结果）
- difficulty：实现难度（低/中/高）
- phase：建议时机（一期/二期）

【企微AI助手增项方向参考】
- 智能问答：员工/客户发消息给机器人，机器人查询智能表格数据后回答（如"帮我查一下本周新增的面试候选人"）
- 快捷指令：发送特定指令，机器人执行写表/查表/推送操作（如发送"催款"触发应收提醒流程）
- 主动推送：当表格数据满足某条件，机器人自动发群消息或私信通知（如逾期未到账提醒项目经理）
- 自然语言写表：发送"帮我加一条 xxx 记录"，机器人解析后写入智能表格
- FAQ 知识库：基于行业/公司知识库，员工随时问机器人（如"合同到期前多久可以续签"）

直接输出 JSON 数组，不要 markdown 代码块，不要任何前缀文字。"""


@app.post("/api/step4/publish")
async def publish_step4_report(body: dict, user: dict = Depends(require_auth)):
    """将 Step4 HTML 方案发布为可分享的外链"""
    client_id = body.get("client_id")
    html_content = body.get("html_content", "")
    doc_type = body.get("type", "presales")  # "presales" or "technical"

    if not client_id:
        return {"success": False, "error": "缺少 client_id"}
    if not html_content:
        return {"success": False, "error": "缺少 html_content"}

    # 保存到本地 public/s/ 目录（用于本地预览）
    share_dir = Path(__file__).parent / "public" / "s"
    share_dir.mkdir(parents=True, exist_ok=True)
    suffix = "_technical" if doc_type == "technical" else ""
    filepath = share_dir / f"{client_id}_step4{suffix}.html"
    # 注入访客追踪 JS
    tracking_js = (
        '<script>'
        '(function(){'
        'var vid=localStorage.getItem("pa_vid")||(localStorage.setItem("pa_vid","v"+Math.random().toString(36).substr(2,9)+Date.now()),localStorage.getItem("pa_vid"));'
        'var fu=location.pathname;'
        'var cid=' + str(client_id) + ';'
        'var sd=0;'
        'function tk(a,e){var p={visitor_id:vid,file_url:fu,client_id:cid,referer:document.referrer,scroll_depth:sd};if(e)Object.assign(p,e);fetch("/api/track/"+a,{method:"POST",headers:{"Content-Type":"application/json"},body:JSON.stringify(p)}).catch(function(){});}'
        'tk("visit");'
        'window.addEventListener("scroll",function(){var h=document.documentElement,b=document.body,pct=Math.round(100*(h.scrollTop||b.scrollTop)/(h.scrollHeight-h.clientHeight));if(pct>sd)sd=pct;},{passive:true});'
        'setInterval(function(){tk("heartbeat");},30000);'
        '})();'
        '</script>'
    )
    if "</body>" in html_content.lower():
        html_content = html_content.replace("</body>", tracking_js + "</body>")
    else:
        html_content += tracking_js

    filepath.write_text(html_content, encoding="utf-8")

    url = f"/public/s/{client_id}_step4{suffix}.html"
    return {"success": True, "url": url}


# ==================== Step4 AI 对话建议 ====================
@app.post("/api/step4/chat-suggest")
async def step4_chat_suggest(body: dict, user: dict = Depends(require_auth)):
    """多轮对话：AI 诊断5维度 + 追问机制"""
    import logging
    logger = logging.getLogger("uvicorn")
    client_id = body.get("client_id")
    doc_type = body.get("doc_type")  # 'presales' | 'technical'
    content = body.get("content", "")[:3000]
    user_input = body.get("user_input", "").strip()
    history = body.get("history", [])  # [{role:'user'|'assistant', content:'...'}]

    if not client_id:
        raise HTTPException(status_code=400, detail="client_id required")

    # 判断是售前(5维度)还是技术路线(4维度)
    is_presales = doc_type == 'presales'
    dimensions = "A内容完整度、B案例可信度、C语言风格、D差异化、E其他" if is_presales else "A内容完整度、B案例可信度、D差异化、E其他"

    # 如果没有内容，返回诊断引导
    if not content:
        return {
            "success": True,
            "type": "suggestions",
            "suggestions": [
                {"title": "暂无方案内容", "desc": "请先生成售前方案", "prompt": "请先生成售前方案"}
            ]
        }

    # 构建多轮对话 prompt
    system_prompt = """你是一个经验丰富的售前方案优化顾问，用苏格拉底式追问法帮助客户明确修改方向。

## 绝对禁止
- 禁止问"是否/是不是/要不要"这类yes/no问题
- 禁止选项中出现"是"、"否"、"要"、"不要"
- 禁止说"内容不够完整/不具体"这种泛泛的话
- 必须基于文档具体内容来问

## 你的工作方式（严格遵守）
用户给出修改意见后，你**绝对不要**直接输出建议卡片。你必须：
1. 读取方案内容，找到最需要明确的一个具体点
2. 问一个**具体到能让客户直接回答**的问题
3. 给出2-4个**具体的选项**让客户选

## 好的追问例子（基于方案内容）
- "文档中提到了A客户的案例，请问贵公司更关注哪个行业的案例？"
- "痛点中提到'效率低'，具体是哪个环节最耗时？"
- "报价写了15万，这个预算贵公司主要考虑哪方面的投入？"
- "案例里已有公牛集团的案例，您觉得还需要哪个同行案例？"
- "方案中没有提到实施周期，您认为一期希望多久完成？"

## 坏的追问（禁止出现）
- "痛点分析是否需要更深入？" ❌
- "要不要增加案例？" ❌
- "内容是否完整？" ❌

## 输出格式（严格JSON，不要任何前缀，不要有任何解释文字）
追问时：{"type":"probing","probing":"针对方案具体内容的追问（30字内）","options":["具体选项A","具体选项B","具体选项C"]}
直接指令时：{"type":"directive","directive":"具体的修改指令（30字内）"}
确认时：{"type":"confirm","confirm":"复述客户意图（20字内）","options":["A. 是的", "B. 不对"]}"""

    # 构建对话上下文（最近4轮）
    recent = history[-4:] if history else []
    ctx = "\n".join([f"{'顾问' if m.get('role')=='assistant' else '用户'}：{m.get('content','')}" for m in recent])

    user_prompt = f"""## 对话历史
{ctx}

## 方案内容摘要
{content[:2000] if content else '（暂无方案内容）'}

## 用户最新反馈
{user_input if user_input else '（分析方案，给出追问，不要直接给建议卡片）'}"""

    try:
        result = call_minimax(system_prompt, user_prompt, max_tokens=1200)
        raw = result.get("content", "").strip()
        logger.warning(f"[chat-suggest] raw response: {raw[:200]}")
        # 解析 JSON
        try:
            data = json.loads(raw)
        except:
            import re
            m = re.search(r'\{.*\}', raw, re.DOTALL)
            if m:
                data = json.loads(m.group())
            else:
                data = {"type": "raw", "content": raw[:200]}
    except Exception as e:
        logger.warning(f"[chat-suggest] error: {e}")
        data = {"type": "raw", "content": str(e)}

    return {"success": True, **data}


# ==================== Step4 售前对话 · 文字聊天 ====================
@app.post("/api/step4/chat")
async def step4_chat(body: dict, user: dict = Depends(require_auth)):
    """对话式修订：AI 根据用户反馈给出具体修改建议（含内容+位置）"""
    import logging
    logger = logging.getLogger("uvicorn")
    client_id = body.get("client_id")
    doc_type = body.get("doc_type", "presales")
    html_content = body.get("html_content")   # JSON 对象或 HTML 字符串
    requirement_data = body.get("requirement_data", {})  # 原始结构化需求
    user_input = body.get("user_input", "").strip()
    history = body.get("history", [])  # [{role:'user'|'assistant', content:'...'}]

    if not client_id:
        raise HTTPException(status_code=400, detail="client_id required")
    if not user_input:
        raise HTTPException(status_code=400, detail="user_input required")

    # 构建对话上下文（最近6轮）
    recent = history[-6:] if history else []
    ctx = "\n".join([f"{'顾问' if m.get('role')=='assistant' else '用户'}：{m.get('content','')}" for m in recent])

    # 序列化 html_content
    if isinstance(html_content, dict):
        html_str = json.dumps(html_content, ensure_ascii=False)
    else:
        html_str = str(html_content) if html_content else ""

    req_str = json.dumps(requirement_data, ensure_ascii=False) if requirement_data else "暂无结构化需求数据"

    # 根据 doc_type 切换角色
    if doc_type == "technical":
        system_prompt = """你是一个企业微信技术文档优化顾问。

## 你的职责
用户对已生成的技术路线及报价方案（HTML 格式）提出修改意见，你给出**具体的修改建议**（修改内容 + 精确位置）。

## 文档结构（11章26表）
一、客户基础信息与当前现状 / 二、场景类型判断与方案边界 / 三、需求理解与优先级确认 / 四、业务流程设计 / 五、企业微信方案总览 / 六、智能表格交付设计（字段逐字段展开，核心） / 七、审批与自动化设计 / 八、权限与数据看板设计 / 九、数据来源与系统对接 / 十、实施计划、报价口径与变更机制 / 十一、待客户确认问题与签署

## 绝对禁止
- 不要说"内容不够完整/不具体"这种泛泛的话
- 不要说"是否/是不是/要不要"这类 yes/no 问题
- 不要只给方向性建议，必须给出**具体可执行的内容文本**
- 禁止输出任何解释、说明、注释

## 修改位置描述规则
技术文档是 HTML，回答位置时必须具体到：章节号 + 小节号 + 表格/段落位置。
例如："在第三章 3.2 小节的'需求清单'表格第 3 行，将'XX'改为'YY'"
或："在第六章 6.2 的'核心字段设计表'中，为'入库单号'字段增加'自动编号'类型说明"

## 输出格式（严格 JSON，不要前缀，不要 markdown 包裹）
当你认为用户意见可行时：
{"agreed": true, "suggestion": "具体的修改建议文本（含要加入的具体内容）", "location": "精确位置描述（章.节 表格名+行号/位置）"}

当你认为用户意见不明确、需要澄清时：
{"agreed": false, "probing": "追问具体是哪个方面需要调整（20字内）", "options": ["A. 措辞严谨性", "B. 表格内容", "C. 报价口径", "D. 范围边界"]}

直接输出 JSON，不要有任何其他文字。"""
    else:
        system_prompt = """你是一个企业微信智能表格售前方案优化顾问。

## 你的职责
用户对已生成的售前方案提出修改意见，你给出**具体的修改建议**（修改内容 + 放置位置）。

## 绝对禁止
- 不要说"内容不够完整/不具体"这种泛泛的话
- 不要说"是否/是不是/要不要"这类 yes/no 问题
- 不要只给方向性建议，必须给出**具体可执行的内容文本**
- 禁止输出任何解释、说明、注释

## 输出格式（严格 JSON，不要前缀，不要 markdown 包裹）
当你认为用户意见可行时：
{"agreed": true, "suggestion": "具体的修改建议文本（含要加入的内容和位置）", "location": "建议放置的具体位置描述"}

当你认为用户意见不明确、需要澄清时：
{"agreed": false, "probing": "追问具体是哪个方面需要调整（20字内）", "options": ["A. 措辞风格", "B. 案例内容", "C. 价值量化", "D. 结构框架"]}

直接输出 JSON，不要有任何其他文字。"""

    # 技术文档是完整 HTML 字符串（截取前 4000 字符）；售前是 JSON 对象
    if doc_type == "technical":
        doc_summary = html_str[:4000] if html_str else "（暂无技术文档内容）"
    else:
        doc_summary = html_str[:3000]

    user_prompt = f"""## 当前方案摘要
{doc_summary}

## 原始需求数据
{req_str[:2000]}

## 对话历史
{ctx}

## 用户最新反馈
{user_input}

请根据用户反馈，给出具体修改建议或追问。不要输出任何解释文字，只输出 JSON。"""

    try:
        result = call_minimax(system_prompt, user_prompt, max_tokens=1500)
        raw = result.get("content", "").strip()
        logger.warning(f"[step4-chat] raw: {raw[:300]}")
        # 解析 JSON
        try:
            data = json.loads(raw)
        except:
            import re
            m = re.search(r'\{.*\}', raw, re.DOTALL)
            if m:
                data = json.loads(m.group())
            else:
                data = {"agreed": False, "probing": "请更具体描述您的修改需求", "options": ["措辞调整", "补充内容", "删除内容", "结构重组"]}
    except Exception as e:
        logger.warning(f"[step4-chat] error: {e}")
        data = {"agreed": False, "probing": "抱歉，对话处理出现异常，请稍后重试", "options": ["措辞调整", "补充内容", "删除内容", "结构重组"]}

    return {"success": True, **data}


# ==================== Step4 售前对话 · 重新生成 ====================
@app.post("/api/step4/regenerate")
async def step4_regenerate(body: dict, user: dict = Depends(require_auth)):
    """对话确认后，根据改动清单重新生成售前方案 HTML（新版本追加到列表）"""
    import logging
    logger = logging.getLogger("uvicorn")
    client_id = body.get("client_id")
    doc_type = body.get("doc_type", "presales")
    html_content = body.get("html_content")   # 当前版本的 JSON 对象
    requirement_data = body.get("requirement_data", {})
    confirmed_changes = body.get("confirmed_changes", [])  # [{action, location, content}]
    chat_history = body.get("chat_history", [])  # 完整对话历史

    if not client_id:
        raise HTTPException(status_code=400, detail="client_id required")
    if not html_content:
        raise HTTPException(status_code=400, detail="html_content required")

    # 获取客户数据（用于 ctx 构建）
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM clients WHERE id = ? AND user_id = ?", (client_id, user["user_id"]))
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="客户不存在")

    client = dict(row)
    for field in ("step1_result", "step2_report", "step2_todo", "step2_schema", "step3_summary", "uploaded_files", "step4_input_draft"):
        if client.get(field) and isinstance(client[field], str):
            try:
                client[field] = json.loads(client[field])
            except:
                pass

    # 构建完整需求数据（同 generate_step4_artifacts）
    step1 = client.get("step1_result", {}) or {}
    step3 = client.get("step3_summary", {}) or {}
    step4_input = client.get("step4_input_draft", {}) or {}
    uploaded_files = client.get("uploaded_files") or []

    def _build_req_data():
        # 复用 generate_step4_artifacts 中构建 requirement_data 的逻辑片段
        rd = {
            "meta": {
                "company_name": client.get("name", ""),
                "industry": client.get("industry", ""),
                "scale": client.get("scale", ""),
                "initial_demand": client.get("initial_demand", "")
            },
            "step1": step1,
            "step3": step3,
            "step4_input": step4_input,
            "uploaded_files_text": "\n\n".join([
                f"【{f.get('name','记录')}】{f.get('content','') or f.get('text','')}"
                for f in uploaded_files if f.get('content') or f.get('text')
            ])
        }
        return rd

    req_data = _build_req_data()
    req_json_str = json.dumps(req_data, ensure_ascii=False)

    # 对话摘要（用于注入到 prompt）
    changes_text = "\n".join([
        f"- [{c.get('action','')}] 在「{c.get('location','')}」：{c.get('content','')}"
        for c in confirmed_changes
    ]) if confirmed_changes else "（用户无额外修改要求）"

    chat_summary = ""
    if chat_history:
        for msg in chat_history:
            role = "用户" if msg.get("role") == "user" else "AI"
            chat_summary += f"{role}：{msg.get('content','')}\n"

    # 构建专用 regenerate prompt
    REGENERATE_PROMPT = f"""你是一个企业微信智能表格可视化方案顾问。请基于原始需求 + 用户对话确认的修改意见，重新生成售前方案 HTML 内容。

【原始需求数据】
{req_json_str}

【用户确认的修改意见】
{changes_text}

【对话历史摘要】
{chat_summary}

请重新生成完整的方案 JSON，直接输出 JSON，不要任何前缀、说明、markdown 包裹。

输出格式同 STEP4_HTML_PROMPT：
{{
  "pageTitle": "",
  "hero": {{...}},
  "customerStageJudgement": {{...}},
  "insightSection": {{...}},
  "scenarioBreakdown": [...],
  "architecture": {{...}},
  "recommendedModules": [...],
  "roadmap": [...],
  "valuePoints": [...],
  "pendingQuestions": []
}}

**强制要求**：
- 将用户确认的修改意见**全部融入**新版本中
- 保持 JSON 结构完整，所有字段不得为 null 或空数组
- 以 `{{` 开头，以 `}}` 结尾，不输出任何其他文字"""

    html_prompt = REGENERATE_PROMPT  # 数据已通过 f-string 注入
    REGEN_SYSTEM = "你是一个企业微信智能表格可视化方案顾问。请严格根据用户提供的需求数据生成方案，直接输出 JSON，不要任何前缀解释。"
    ai_result = call_codebuddy(REGEN_SYSTEM, html_prompt, max_tokens=10000)
    raw = ai_result.get("content", "").strip()

    if raw.startswith("Error:") or not raw:
        return {"success": False, "error": raw or "生成失败"}

    new_html_content = parse_json_response(raw)
    if not new_html_content:
        return {"success": False, "error": "生成内容解析失败，请重试"}

    # 保存新版本到数据库
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT step4_presales_versions FROM clients WHERE id = ?", (client_id,))
    row = cursor.fetchone()
    existing_versions = []
    if row and row[0]:
        try:
            existing_versions = json.loads(row[0])
        except:
            existing_versions = []
    conn.close()

    new_version = {
        "version": len(existing_versions) + 1,
        "content": {"htmlContent": new_html_content},
        "created_at": datetime.now().isoformat(),
        "chat_history": chat_history,
        "confirmed_changes": confirmed_changes,
        "source": "chat_regenerate"
    }
    existing_versions.append(new_version)

    db_update_client(client_id, {"step4_presales_versions": existing_versions})

    # 发布新版本 HTML
    html_text = json.dumps(new_html_content, ensure_ascii=False)
    publish_result = {"success": False}
    try:
        from pathlib import Path
        out_dir = Path("/Users/laixiangjun/Eco-Wecom/outputs") / str(client_id)
        out_dir.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = out_dir / f"售前解决方案_{client.get('name','客户')}_{ts}_V{len(existing_versions)}.html"
        # 注入访客追踪
        tracking_js = ""
        full_html = f'<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8"><title>{new_html_content.get("pageTitle","售前方案")}</title></head><body>{html_text}{tracking_js}</body></html>'
        filepath.write_text(full_html, encoding="utf-8")
        pub_url = f"/public/s/{client_id}_step4_presales_v{len(existing_versions)}.html"
        publish_result = {"success": True, "url": pub_url}
    except Exception as e:
        logger.warning(f"[regenerate] publish error: {e}")

    return {
        "success": True,
        "version": len(existing_versions),
        "html_content": new_html_content,
        "publish_url": publish_result.get("url", ""),
        "message": f"新版本 V{len(existing_versions)} 已生成并追加到版本列表"
    }


# ==================== Step4 技术文档对话 · 重新生成 ====================
@app.post("/api/step4/regenerate-technical")
async def step4_regenerate_technical(body: dict, user: dict = Depends(require_auth)):
    """对话确认后，根据改动清单重新生成技术路线及报价方案 HTML（新版本追加到列表）"""
    import logging
    logger = logging.getLogger("uvicorn")
    client_id = body.get("client_id")
    html_content = body.get("html_content")   # 当前版本的技术文档 HTML 字符串
    confirmed_changes = body.get("confirmed_changes", [])  # [{action, location, content}]
    chat_history = body.get("chat_history", [])  # 完整对话历史

    if not client_id:
        raise HTTPException(status_code=400, detail="client_id required")
    if not html_content:
        raise HTTPException(status_code=400, detail="html_content required")

    # 获取客户数据
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM clients WHERE id = ? AND user_id = ?", (client_id, user["user_id"]))
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="客户不存在")

    client = dict(row)
    for field in ("step1_result", "step2_report", "step2_todo", "step2_schema", "step3_summary", "uploaded_files", "step4_input_draft"):
        if client.get(field) and isinstance(client[field], str):
            try:
                client[field] = json.loads(client[field])
            except:
                pass

    # 构建完整需求数据
    step1 = client.get("step1_result", {}) or {}
    step3 = client.get("step3_summary", {}) or {}
    step4_input = client.get("step4_input_draft", {}) or {}
    uploaded_files = client.get("uploaded_files") or []

    rd = {
        "meta": {
            "company_name": client.get("name", ""),
            "industry": client.get("industry", ""),
            "scale": client.get("scale", ""),
            "initial_demand": client.get("initial_demand", "")
        },
        "step1": step1,
        "step3": step3,
        "step4_input": step4_input,
        "uploaded_files_text": "\n\n".join([
            f"【{f.get('name','记录')}】{f.get('content','') or f.get('text','')}"
            for f in uploaded_files if f.get('content') or f.get('text')
        ])
    }
    req_json_str = json.dumps(rd, ensure_ascii=False)

    # 对话摘要
    changes_text = "\n".join([
        f"- [{c.get('action','')}] 在「{c.get('location','')}」：{c.get('content','')}"
        for c in confirmed_changes
    ]) if confirmed_changes else "（用户无额外修改要求）"

    chat_summary = ""
    if chat_history:
        for msg in chat_history:
            role = "用户" if msg.get("role") == "user" else "AI"
            chat_summary += f"{role}：{msg.get('content','')}\n"

    # 技术文档专用 regenerate prompt（完整 HTML，不是 JSON）
    REGEN_TECH_PROMPT = f"""你是一个企业微信定制开发技术方案顾问。请基于原始需求 + 用户对话确认的修改意见，重新生成《技术路线及报价方案》完整 HTML 文档。

【原始需求数据】
{req_json_str}

【用户确认的修改意见】
{changes_text}

【对话历史摘要】
{chat_summary}

请重新生成完整 HTML 文档（文档风，华文楷体，可下载 Word），直接输出 HTML，不要任何前缀、说明、markdown 包裹。

## 文档骨架（11 章 26 表，锁死顺序）
封面 / 元信息表×2 / 一、客户基础信息与当前现状 / 二、场景类型判断与方案边界 / 三、需求理解与优先级确认 / 四、业务流程设计 / 五、企业微信方案总览 / 六、智能表格交付设计（字段逐字段展开） / 七、审批与自动化设计 / 八、权限与数据看板设计 / 九、数据来源、系统对接与交付边界 / 十、实施计划、报价口径与变更机制 / 十一、待客户确认问题与签署

**强制要求**：
- 将用户确认的修改意见**全部融入**新版本中
- 保持完整 HTML 结构，华文楷体，右下角有下载 Word 按钮
- 以 `<!DOCTYPE` 或 `<html` 开头，不要有任何其他文字"""

    REGEN_TECH_SYSTEM = "你是一个企业微信定制开发技术方案顾问。请严格根据用户提供的需求数据生成技术文档，直接输出完整 HTML，不要任何前缀解释。"

    ai_result = call_codebuddy(REGEN_TECH_SYSTEM, REGEN_TECH_PROMPT, max_tokens=12000)
    raw = ai_result.get("content", "").strip()

    if raw.startswith("Error:") or not raw:
        return {"success": False, "error": raw or "生成失败"}

    # raw 是完整 HTML，不需要 parse_json_response
    new_html_content = raw

    # 保存文件
    from pathlib import Path
    out_dir = Path("/Users/laixiangjun/Eco-Wecom/outputs") / str(client_id)
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"技术路线及报价方案_{client.get('name','客户')}_{ts}_V{len(json.loads(client.get('step4_technical_versions') or '[]')) + 1}.html"
    filepath = out_dir / filename

    # 注入访客追踪
    tracking_js = (
        '<script>'
        '(function(){'
        'var vid=localStorage.getItem("pa_vid")||(localStorage.setItem("pa_vid","v"+Math.random().toString(36).substr(2,9)+Date.now()),localStorage.getItem("pa_vid"));'
        'var fu=location.pathname;'
        'var cid=' + str(client_id) + ';'
        'var sd=0;'
        'function tk(a,e){var p={visitor_id:vid,file_url:fu,client_id:cid,referer:document.referrer,scroll_depth:sd};if(e)Object.assign(p,e);fetch("/api/track/"+a,{method:"POST",headers:{"Content-Type":"application/json"},body:JSON.stringify(p)}).catch(function(){});}'
        'tk("visit");'
        'window.addEventListener("scroll",function(){var h=document.documentElement,b=document.body,pct=Math.round(100*(h.scrollTop||b.scrollTop)/(h.scrollHeight-h.clientHeight));if(pct>sd)sd=pct;},{passive:true});'
        'setInterval(function(){tk("heartbeat");},30000);'
        '})();'
        '</script>'
    )
    if "</body>" in new_html_content.lower():
        new_html_content = new_html_content.replace("</body>", tracking_js + "</body>")
    else:
        new_html_content += tracking_js

    filepath.write_text(new_html_content, encoding="utf-8")
    file_url = f"/outputs/{client_id}/{filename}"

    # 保存新版本到数据库
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT step4_technical_versions FROM clients WHERE id = ?", (client_id,))
    row = cursor.fetchone()
    existing_versions = []
    if row and row[0]:
        try:
            existing_versions = json.loads(row[0])
        except:
            existing_versions = []
    conn.close()

    new_version = {
        "version": len(existing_versions) + 1,
        "content": {"htmlContent": new_html_content},
        "created_at": datetime.now().isoformat(),
        "file_url": file_url,
        "filename": filename,
        "chat_history": chat_history,
        "confirmed_changes": confirmed_changes,
        "source": "chat_regenerate"
    }
    existing_versions.append(new_version)

    db_update_client(client_id, {"step4_technical_versions": existing_versions})

    return {
        "success": True,
        "version": len(existing_versions),
        "html_content": new_html_content,
        "file_url": file_url,
        "filename": filename,
        "message": f"新版本 V{len(existing_versions)} 已生成并追加到版本列表"
    }


# ==================== 访问追踪 ====================
def _parse_ua(ua_str):
    """从 User-Agent 解析设备/操作系统/浏览器"""
    ua = ua_str or ""
    device = "未知"
    os_type = "未知"
    browser = "未知"
    if "Mobile" in ua or "Android" in ua and "Mobile" in ua:
        device = "手机"
    elif "iPad" in ua or "Tablet" in ua:
        device = "平板"
    elif "Windows" in ua:
        device = "电脑"
    elif "Macintosh" in ua or "Mac OS" in ua:
        device = "电脑"
    elif "Linux" in ua and "Android" not in ua:
        device = "电脑"
    if "Windows NT 10" in ua:
        os_type = "Windows 10/11"
    elif "Windows NT 6.3" in ua:
        os_type = "Windows 8"
    elif "Mac OS X" in ua:
        os_type = "macOS"
    elif "Android" in ua:
        os_type = "Android"
    elif "iPhone" in ua or "iOS" in ua:
        os_type = "iOS"
    elif "Linux" in ua:
        os_type = "Linux"
    if "Chrome/" in ua and "Edg/" not in ua:
        browser = "Chrome"
    elif "Firefox/" in ua:
        browser = "Firefox"
    elif "Safari/" in ua and "Chrome/" not in ua:
        browser = "Safari"
    elif "Edg/" in ua:
        browser = "Edge"
    elif "MicroMessenger/" in ua:
        browser = "微信"
    return device, os_type, browser


def _get_client_ip(request):
    """获取真实 IP（支持代理）"""
    return request.headers.get("x-forwarded-for", "").split(",")[0].strip() or \
           request.headers.get("x-real-ip", "") or \
           "127.0.0.1"


@app.post("/api/track/visit")
async def track_visit(request: Request, body: dict):
    """记录首次访问（访客打开分享链接时）"""
    client_id = body.get("client_id")
    file_url = body.get("file_url", "")
    visitor_id = body.get("visitor_id", "")
    referer = body.get("referer", "")
    scroll_depth = body.get("scroll_depth", 0)

    if not client_id or not file_url:
        return {"success": False}

    conn = get_db()
    cursor = conn.cursor()
    ip = _get_client_ip(request)
    ua = request.headers.get("user-agent", "")
    device, os_type, browser = _parse_ua(ua)

    # 检查是否已有该 visitor_id 的记录
    existing = None
    if visitor_id:
        cursor.execute(
            "SELECT id,visit_count FROM visit_tracking WHERE visitor_id=? AND file_url=? ORDER BY id DESC LIMIT 1",
            (visitor_id, file_url)
        )
        existing = cursor.fetchone()

    now = datetime.now().isoformat()
    if existing:
        # 回访：更新 visit_count / last_visit_at / scroll_depth
        cursor.execute("""
            UPDATE visit_tracking
            SET visit_count=visit_count+1, last_visit_at=?, last_heartbeat_at=?,
                scroll_depth=?, ip_address=?
            WHERE id=?
        """, (now, now, scroll_depth, ip, existing[0]))
    else:
        # 首次访问
        cursor.execute("""
            INSERT INTO visit_tracking
            (client_id,file_url,visitor_id,ip_address,user_agent,device_type,os_type,browser_type,
             referer,is_first_visit,visit_count,first_visit_at,last_visit_at,last_heartbeat_at,stay_duration,scroll_depth)
            VALUES (?,?,?,?,?,?,?,?,?,1,1,?,?,?,0,?)
        """, (client_id, file_url, visitor_id, ip, ua, device, os_type, browser,
              referer, now, now, now, scroll_depth))
    conn.commit()
    return {"success": True}


@app.post("/api/track/heartbeat")
async def track_heartbeat(request: Request, body: dict):
    """页面心跳（每30秒发送一次，更新停留时长）"""
    visitor_id = body.get("visitor_id", "")
    file_url = body.get("file_url", "")
    scroll_depth = body.get("scroll_depth", 0)

    if not visitor_id or not file_url:
        return {"success": False}

    conn = get_db()
    cursor = conn.cursor()
    now = datetime.now().isoformat()
    cursor.execute("""
        UPDATE visit_tracking
        SET last_heartbeat_at=?,
            stay_duration=COALESCE(stay_duration,0)+30,
            scroll_depth=MAX(COALESCE(scroll_depth,0),?)
        WHERE visitor_id=? AND file_url=?
        ORDER BY id DESC LIMIT 1
    """, (now, scroll_depth, visitor_id, file_url))
    conn.commit()
    rows = cursor.rowcount
    return {"success": rows > 0}


@app.get("/api/clients/{client_id}/visits")
async def get_client_visits(client_id: int,
                            page: int = 1,
                            limit: int = 20,
                            device: str = "",
                            date_from: str = "",
                            date_to: str = "",
                            user: dict = Depends(require_auth)):
    offset = (page - 1) * limit
    conn = get_db()
    cursor = conn.cursor()

    # 组合筛选条件
    sql = "SELECT * FROM visit_tracking WHERE client_id=? AND file_url LIKE '/public/s/%'"
    params = [client_id]
    if device:
        sql += " AND device_type=?"
        params.append(device)
    if date_from:
        sql += " AND created_at>=?"
        params.append(date_from)
    if date_to:
        sql += " AND created_at<=?"
        params.append(date_to + " 23:59:59")

    # 总数
    cursor.execute("SELECT COUNT(*) FROM visit_tracking WHERE client_id=? AND file_url LIKE '/public/s/%'", (client_id,))
    total = cursor.fetchone()[0]

    sql += " ORDER BY last_visit_at DESC LIMIT ? OFFSET ?"
    params.extend([limit, offset])
    cursor.execute(sql, params)
    rows = cursor.fetchall()
    cols = [d[0] for d in cursor.description]
    items = [dict(zip(cols, r)) for r in rows]
    return {"success": True, "total": total, "page": page, "limit": limit, "visits": items}


@app.get("/api/clients/{client_id}/visits/by-user")
async def get_visits_by_user(client_id: int,
                             user: dict = Depends(require_auth)):
    """按访客聚合的访问统计"""
    conn = get_db()
    cursor = conn.cursor()
    # 按 visitor_id 聚合
    cursor.execute("""
        SELECT
            visitor_id,
            COUNT(*) as visit_count,
            SUM(visit_count) as total_visits,
            SUM(stay_duration) as total_duration,
            MAX(last_visit_at) as last_visit_at,
            MIN(first_visit_at) as first_visit_at,
            MAX(scroll_depth) as max_scroll_depth,
            device_type,
            os_type,
            browser_type,
            country,
            region,
            city,
            referer
        FROM visit_tracking
        WHERE client_id=? AND file_url LIKE '/public/s/%'
        GROUP BY visitor_id
        ORDER BY total_visits DESC
        LIMIT 50
    """, (client_id,))
    rows = cursor.fetchall()
    cols = [d[0] for d in cursor.description]
    users = [dict(zip(cols, r)) for r in rows]
    # 同时返回各页面的访问统计
    cursor.execute("""
        SELECT file_url, COUNT(*) as visit_count, SUM(stay_duration) as total_duration
        FROM visit_tracking
        WHERE client_id=? AND file_url LIKE '/public/s/%'
        GROUP BY file_url
        ORDER BY visit_count DESC
    """, (client_id,))
    page_rows = cursor.fetchall()
    page_cols = [d[0] for d in cursor.description]
    pages = [dict(zip(page_cols, r)) for r in page_rows]
    return {"success": True, "visitors": users, "pages": pages}


@app.get("/api/clients/{client_id}/visits/{visitor_id}/timeline")
async def get_visitor_timeline(client_id: int, visitor_id: str,
                               user: dict = Depends(require_auth)):
    """某个访客的访问时间线"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT * FROM visit_tracking
        WHERE client_id=? AND visitor_id=?
        ORDER BY last_visit_at DESC
        LIMIT 20
    """, (client_id, visitor_id))
    rows = cursor.fetchall()
    cols = [d[0] for d in cursor.description]
    items = [dict(zip(cols, r)) for r in rows]
    return {"success": True, "visits": items}


@app.post("/api/step5/agent-suggest")
async def step5_agent_suggest(body: dict, user: dict = Depends(require_auth)):
    """生成 Step5 AI 增强建议"""
    client_id = body.get("client_id")
    if not client_id:
        return {"success": False, "error": "缺少 client_id"}

    client = db_get_client(client_id)
    if not client:
        return {"success": False, "error": "客户不存在"}

    # 获取 schema 和需求数据
    schema = client.get("step5_schema")
    if isinstance(schema, str):
        try:
            schema = json.loads(schema)
        except:
            schema = {}

    if not schema:
        return {"success": False, "error": "请先生成智能表格 Schema（Step5）"}

    # 构建 schema 摘要用于 prompt
    sheets = schema.get("sheets") or []
    schema_lines = []
    for s in sheets:
        name = s.get("sheet_name", "未命名")
        fields = s.get("fields") or []
        field_names = [f.get("field_title", "") for f in fields]
        schema_lines.append(f"- {name}：{', '.join(field_names)}")
    schema_summary = "\n".join(schema_lines) if schema_lines else "暂无子表"

    customer_name = client.get("name", "")
    industry = client.get("industry", "")
    initial_demand = client.get("initial_demand", "")

    user_prompt = STEP5_AGENT_PROMPT.format(
        customer_name=customer_name,
        industry=industry,
        initial_demand=initial_demand,
        schema_summary=schema_summary
    )

    ai_result = call_codebuddy(STEP5_AGENT_PROMPT, user_prompt, max_tokens=4000)
    raw = ai_result["content"]
    record_ai_tokens(client_id, ai_result["usage"]["total_tokens"])
    if raw.startswith("Error:"):
        return {"success": False, "error": raw}

    suggestions = parse_json_response(raw)
    if not suggestions:
        return {"success": False, "error": "AI 返回格式异常，请重试"}

    # 保存到后端
    db_update_client(client_id, {"step5_agent_suggestions": suggestions})

    return {"success": True, "suggestions": suggestions}


# ==================== 创建企业微信智能表格（从 smartTableSpec） ====================

# 字段类型映射：我们的规范 → WeCom field_type id
def _build_field_option_map(fields_resp: dict) -> dict:
    """从 smartsheet_get_fields 响应中，构建 {field_title: {option_id: option_text}} 的映射"""
    option_map = {}
    for f in (fields_resp.get("fields") or []):
        title = f.get("field_title", "")
        ft = f.get("field_type", "")
        # 只有单选/多选字段有 options
        prop = f.get("property_single_select") or f.get("property_select") or {}
        opts = prop.get("options") or f.get("options") or []
        if ft in ("FIELD_TYPE_SINGLE_SELECT", "FIELD_TYPE_SELECT") and opts:
            option_map[title] = {opt.get("id", ""): opt.get("text", "") for opt in opts}
    return option_map


def _convert_select_values(record: dict, option_map: dict) -> dict:
    """将单选/多选字段的文本值转换为 option_id；如果 option_id 为空则保留原文字值"""
    for field_title, options in option_map.items():
        if field_title in record:
            val = record[field_title]
            if isinstance(val, list):
                # 多选：转成 [option_id, ...]，option_id 为空则保留文字
                record[field_title] = [oid if oid else otxt for oid, otxt in options.items() if otxt in val] or [v for v in val if v]
            else:
                # 单选：找 option_id，找不到则保留原文字值
                matched = False
                for oid, otxt in options.items():
                    if otxt == val:
                        record[field_title] = oid if oid else val
                        matched = True
                        break
                if not matched:
                    record[field_title] = val  # 保留原文字值，让企微自己匹配
    return record


FIELD_TYPE_MAP = {
    # WeCom API 格式（FIELD_TYPE_ 前缀）
    "文本": "FIELD_TYPE_TEXT",
    "多行文本": "FIELD_TYPE_TEXT",
    "单选": "FIELD_TYPE_SINGLE_SELECT",
    "多选": "FIELD_TYPE_SELECT",
    "数字": "FIELD_TYPE_NUMBER",
    "金额": "FIELD_TYPE_CURRENCY",
    "日期": "FIELD_TYPE_DATE_TIME",
    "日期时间": "FIELD_TYPE_DATE_TIME",
    "人员": "FIELD_TYPE_USER",
    "附件": "FIELD_TYPE_ATTACHMENT",
    "图片": "FIELD_TYPE_IMAGE",
    "关联记录": "FIELD_TYPE_TEXT",
    "公式": "FIELD_TYPE_TEXT",
    "自动编号": "FIELD_TYPE_TEXT",
    "进度": "FIELD_TYPE_PROGRESS",
    "勾选": "FIELD_TYPE_CHECKBOX",
    "URL": "FIELD_TYPE_URL",
    # 别名
    "text": "FIELD_TYPE_TEXT",
    "number": "FIELD_TYPE_NUMBER",
    "currency": "FIELD_TYPE_CURRENCY",
    "date": "FIELD_TYPE_DATE_TIME",
    "datetime": "FIELD_TYPE_DATE_TIME",
    "contact": "FIELD_TYPE_USER",
    "file": "FIELD_TYPE_ATTACHMENT",
    "checkbox": "FIELD_TYPE_CHECKBOX",
    "percent": "FIELD_TYPE_PERCENTAGE",
}


def _map_field_type(ft: str) -> str:
    """将字段类型字符串映射为 WeCom field_type 枚举值"""
    ft = ft.strip()
    return FIELD_TYPE_MAP.get(ft, "FIELD_TYPE_TEXT")  # 默认文本


@app.post("/api/create")
async def create_wecom_sheet(body: dict, user: dict = Depends(require_auth)):
    """
    基于 smartTableSpec 创建企业微信智能表格（包含完整字段和样例数据）。

    body = {
        "client_id": 71,
        "smartTableSpec": {
            "confirmedTables": [...],
            "fieldsByTable": [...],
            "sheets": [...]  // 直接包含 sheets/fields/sample_records 结构
        }
    }
    """
    import re

    client_id = body.get("client_id")
    spec = body.get("smartTableSpec") or {}

    if not client_id:
        return {"success": False, "error": "缺少 client_id"}

    # 支持两种结构：
    # 1. confirmedTables + fieldsByTable（旧结构）
    # 2. sheets 直接包含字段和样例数据（新结构，来自 step5_schema）
    confirmed_tables = spec.get("confirmedTables") or []
    fields_by_table = spec.get("fieldsByTable") or []
    sheets_data = spec.get("sheets") or []  # 直接的 sheets 结构
    fields_map = {f.get("tableName", ""): f.get("fields", []) for f in fields_by_table}

    # ---- 1. 创建智能表格文档 ----
    client = db_get_client(client_id)
    doc_name = (client.get("name", "") if client else "") + " - 需求智能表格"
    r = call_mcp("create_doc", {
        "doc_name": doc_name,
        "doc_type": 10  # 智能表格（整数）
    })
    if not r or (isinstance(r, dict) and r.get("errcode", 0) != 0):
        import logging
        logging.warning(f"[create_wecom_sheet] MCP create_doc resp: {r}")
        return {"success": False, "error": "创建文档失败：" + (r.get("errmsg", str(r)) if isinstance(r, dict) else str(r))}

    docid = r.get("docid", "")
    doc_url = r.get("url", "")
    if not docid:
        return {"success": False, "error": "创建文档失败，未返回 docid"}

    created_sheets = []

    # ---- 2. 获取默认子表的 sheet_id ----
    sr = call_mcp("smartsheet_get_sheet", {"docid": docid})
    if not sr or sr.get("error"):
        return {"success": False, "error": "获取子表失败：" + (sr.get("error") or str(sr))}
    sheets = sr.get("sheet_list", []) or sr.get("sheets", []) or []
    if not sheets:
        return {"success": False, "error": "未找到子表"}
    first_sheet_id = sheets[0].get("sheet_id", "")

    # ---- 2b. 获取默认子表的字段（含默认 field_id 和选项）----
    fr = call_mcp("smartsheet_get_fields", {"docid": docid, "sheet_id": first_sheet_id})
    default_field_id = ""
    if fr and not fr.get("error") and fr.get("fields"):
        default_field_id = fr["fields"][0].get("field_id", "")

    # ---- 3a. 使用 sheets 结构（直接包含字段和样例数据）----
    if sheets_data:
        # 为第一个子表构建 option_map（用于单选/多选值转换）
        first_sheet_option_map = _build_field_option_map(fr) if fr and not fr.get("error") else {}

        for idx, sheet in enumerate(sheets_data):
            sheet_name = sheet.get("sheet_name", f"子表{idx + 1}")
            fields_list = sheet.get("fields") or []
            sample_records = sheet.get("sample_records") or []
            # 每个子表用各自字段构建 option_map（新增子表时从其字段响应获取）
            sheet_option_map = first_sheet_option_map if idx == 0 else {}

            if idx == 0:
                # ---- 用默认子表 ----
                sheet_id = first_sheet_id
                # 重命名默认字段（需要真实的 field_id）
                if fields_list and default_field_id:
                    first_field = fields_list[0]
                    first_ft = _map_field_type(first_field.get("field_type") or first_field.get("fieldType") or "文本")
                    call_mcp("smartsheet_update_fields", {
                        "docid": docid,
                        "sheet_id": sheet_id,
                        "fields": [{"field_title": first_field.get("field_title") or first_field.get("fieldName") or "", "field_id": default_field_id, "field_type": first_ft}]
                    })
                    # 添加其余字段
                    if len(fields_list) > 1:
                        add_fields = [{"field_title": f.get("field_title") or f.get("fieldName") or "", "field_type": _map_field_type(f.get("field_type") or f.get("fieldType") or "文本")} for f in fields_list[1:]]
                        call_mcp("smartsheet_add_fields", {"docid": docid, "sheet_id": sheet_id, "fields": add_fields})
                    # 添加字段后重新获取字段列表（含新字段的 option_id）
                    fr_after = call_mcp("smartsheet_get_fields", {"docid": docid, "sheet_id": sheet_id})
                    import logging
                    logger4 = logging.getLogger("uvicorn")
                    logger4.warning(f"[create_wecom_sheet] fr_after: {fr_after}")
                    first_sheet_option_map = _build_field_option_map(fr_after) if fr_after and not fr_after.get("error") else {}
                    # 如果选项为空，从 sample_records 中提取不重复的单选/多选值，自动创建选项
                    if not first_sheet_option_map and sample_records:
                        all_option_values = set()
                        for rec in sample_records:
                            for k, v in rec.items():
                                if isinstance(v, str):
                                    all_option_values.add(v)
                        if all_option_values:
                            # 为每个有选项需求的字段创建选项
                            select_fields_to_options = {}
                            for f in fields_list:
                                ft = _map_field_type(f.get("field_type") or f.get("fieldType") or "文本")
                                if ft in ("FIELD_TYPE_SINGLE_SELECT", "FIELD_TYPE_SELECT"):
                                    # 从所有记录中提取该字段出现过的值作为选项
                                    opts = []
                                    for rec in sample_records:
                                        val = rec.get(f.get("field_title") or f.get("fieldName", ""))
                                        if val:
                                            opts.append({"text": val, "color": 0})
                                    if opts:
                                        select_fields_to_options[f.get("field_title") or f.get("fieldName", "")] = opts
                            logger4.warning(f"[create_wecom_sheet] auto_create_options: {select_fields_to_options}")
                            # 更新这些字段的选项，并直接从更新响应中提取 option_id
                            for field_title, options in select_fields_to_options.items():
                                # 找到该字段的 field_id 和原始 field_type
                                field_id = None
                                orig_ft = "FIELD_TYPE_TEXT"
                                for f in (fr_after.get("fields") or []):
                                    if f.get("field_title") == field_title:
                                        field_id = f.get("field_id")
                                        orig_ft = f.get("field_type", "FIELD_TYPE_TEXT")
                                        break
                                if field_id:
                                    upd = call_mcp("smartsheet_update_fields", {
                                        "docid": docid,
                                        "sheet_id": sheet_id,
                                        "fields": [{"field_id": field_id, "field_title": field_title, "field_type": orig_ft, "property_single_select": {"is_multiple": False, "options": options}}]
                                    })
                                    logger4.warning(f"[create_wecom_sheet] set_options resp for '{field_title}': {upd}")
                                    # 直接从更新响应中提取 option_id 映射（不依赖 get_fields 重新获取）
                                    if upd and upd.get("fields"):
                                        for updated_field in upd["fields"]:
                                            ps = updated_field.get("property_single_select") or updated_field.get("property_select") or {}
                                            opts = ps.get("options") or []
                                            if opts:
                                                first_sheet_option_map[field_title] = {opt.get("id", ""): opt.get("text", "") for opt in opts}
                                    sheet_option_map = first_sheet_option_map
                            # 不再重新 get_fields，直接用 update 响应中的 option_id
                    logger4.warning(f"[create_wecom_sheet] first_sheet_option_map: {first_sheet_option_map}")
                    sheet_option_map = first_sheet_option_map
                # 重命名子表
                call_mcp("smartsheet_update_sheet", {"docid": docid, "properties": {"sheet_id": sheet_id, "title": sheet_name}})
                # 添加样例数据（需先转换单选/多选字段值）
                if sample_records:
                    # 转换单选/多选字段的文本值为 option_id
                    import logging
                    logger3 = logging.getLogger("uvicorn")
                    converted_records = [_convert_select_values(dict(rec), sheet_option_map) for rec in sample_records]
                    logger3.warning(f"[create_wecom_sheet] converted_records: {converted_records}")
                    logger3.warning(f"[create_wecom_sheet] sheet_option_map: {sheet_option_map}")
                    records_formatted = [{"values": rec} for rec in converted_records]
                    logger3.warning(f"[create_wecom_sheet] records_formatted: {records_formatted}")
                    add_records_resp = call_mcp("smartsheet_add_records", {"docid": docid, "sheet_id": sheet_id, "records": records_formatted})
                    logger3.warning(f"[create_wecom_sheet] add_records resp: {add_records_resp}")
                    if add_records_resp and add_records_resp.get("errcode", 0) != 0:
                        logger3.error(f"[create_wecom_sheet] add_records failed: {add_records_resp}")
            else:
                # ---- 新增子表 ----
                add_resp = call_mcp("smartsheet_add_sheet", {"docid": docid})
                if not add_resp or add_resp.get("errcode", 0) != 0:
                    continue
                # 云函数版: {"errcode": 0, "sheet_id": "xxx"} 或 {"errcode": 0, "properties": {"sheet_id": "xxx"}}
                new_sheet_id = add_resp.get("sheet_id") or (add_resp.get("properties", {}) or {}).get("sheet_id", "") or ""
                if not new_sheet_id:
                    continue
                sheet_id = new_sheet_id
                # 获取新子表的默认 field_id 和字段选项
                new_fr = call_mcp("smartsheet_get_fields", {"docid": docid, "sheet_id": sheet_id})
                new_default_field_id = ""
                if new_fr and not new_fr.get("error") and new_fr.get("fields"):
                    new_default_field_id = new_fr["fields"][0].get("field_id", "")
                    # 为新子表构建 option_map
                    sheet_option_map = _build_field_option_map(new_fr)
                else:
                    sheet_option_map = {}
                # 重命名 + 设字段
                if fields_list and new_default_field_id:
                    first_field = fields_list[0]
                    first_ft = _map_field_type(first_field.get("field_type") or first_field.get("fieldType") or "文本")
                    call_mcp("smartsheet_update_fields", {
                        "docid": docid,
                        "sheet_id": sheet_id,
                        "fields": [{"field_title": first_field.get("field_title") or first_field.get("fieldName") or "", "field_id": new_default_field_id, "field_type": first_ft}]
                    })
                    if len(fields_list) > 1:
                        add_fields = [{"field_title": f.get("field_title") or f.get("fieldName") or "", "field_type": _map_field_type(f.get("field_type") or f.get("fieldType") or "文本")} for f in fields_list[1:]]
                        call_mcp("smartsheet_add_fields", {"docid": docid, "sheet_id": sheet_id, "fields": add_fields})
                # 重命名子表
                call_mcp("smartsheet_update_sheet", {"docid": docid, "properties": {"sheet_id": sheet_id, "title": sheet_name}})
                # 添加样例数据（需先转换单选/多选字段值）
                if sample_records:
                    converted_records = [_convert_select_values(dict(rec), sheet_option_map) for rec in sample_records]
                    records_formatted = [{"values": rec} for rec in converted_records]
                    add_records_resp = call_mcp("smartsheet_add_records", {"docid": docid, "sheet_id": sheet_id, "records": records_formatted})
                    import logging
                    logger3 = logging.getLogger("uvicorn")
                    logger3.warning(f"[create_wecom_sheet] add_records resp: {add_records_resp}")
                    if add_records_resp and add_records_resp.get("errcode", 0) != 0:
                        logger3.error(f"[create_wecom_sheet] add_records failed: {add_records_resp}")

            created_sheets.append(sheet_name)

    # ---- 3b. 使用 confirmedTables + fieldsByTable 结构（兼容旧）----
    elif confirmed_tables:
        for idx, table in enumerate(confirmed_tables):
            table_name = table.get("tableName", f"子表{idx + 1}")
            phase = table.get("phase", "一期")
            if phase != "一期":
                continue

            fields_def = fields_map.get(table.get("tableName", ""), [])

            if idx == 0:
                sheet_id = first_sheet_id
                if fields_def:
                    first_field = fields_def[0]
                    first_ft = _map_field_type(first_field.get("fieldType", "文本"))
                    call_mcp("smartsheet_update_fields", {
                        "docid": docid,
                        "sheet_id": sheet_id,
                        "fields": [{"field_title": first_field.get("fieldName", ""), "field_id": default_field_id, "field_type": first_ft}]
                    })
                    if len(fields_def) > 1:
                        add_fields = [{"field_title": f.get("fieldName", ""), "field_type": _map_field_type(f.get("fieldType", "文本"))} for f in fields_def[1:]]
                        call_mcp("smartsheet_add_fields", {"docid": docid, "sheet_id": sheet_id, "fields": add_fields})
                call_mcp("smartsheet_update_sheet", {"docid": docid, "properties": {"sheet_id": sheet_id, "title": table_name}})
            else:
                add_resp = call_mcp("smartsheet_add_sheet", {"docid": docid})
                if not add_resp or add_resp.get("errcode", 0) != 0:
                    continue
                new_sheet_id = add_resp.get("sheet_id") or (add_resp.get("properties", {}) or {}).get("sheet_id", "") or ""
                if not new_sheet_id:
                    continue
                sheet_id = new_sheet_id
                # 获取新子表的默认 field_id
                new_fr = call_mcp("smartsheet_get_fields", {"docid": docid, "sheet_id": sheet_id})
                new_default_field_id = ""
                if new_fr and not new_fr.get("error") and new_fr.get("fields"):
                    new_default_field_id = new_fr["fields"][0].get("field_id", "")
                if fields_def and new_default_field_id:
                    first_field = fields_def[0]
                    first_ft = _map_field_type(first_field.get("fieldType", "文本"))
                    call_mcp("smartsheet_update_fields", {
                        "docid": docid,
                        "sheet_id": sheet_id,
                        "fields": [{"field_title": first_field.get("fieldName", ""), "field_id": new_default_field_id, "field_type": first_ft}]
                    })
                    if len(fields_def) > 1:
                        add_fields = [{"field_title": f.get("fieldName", ""), "field_type": _map_field_type(f.get("fieldType", "文本"))} for f in fields_def[1:]]
                        call_mcp("smartsheet_add_fields", {"docid": docid, "sheet_id": sheet_id, "fields": add_fields})
                call_mcp("smartsheet_update_sheet", {"docid": docid, "properties": {"sheet_id": sheet_id, "title": table_name}})

            created_sheets.append(table_name)

    return {
        "success": True,
        "docid": docid,
        "url": doc_url,
        "sheets": created_sheets
    }


@app.post("/api/create_doc")
async def create_wecom_doc(body: dict, user: dict = Depends(require_auth)):
    """
    创建企业微信智能文档（smartpage）并写入 Markdown 内容。
    body = {
        "client_id": 71,
        "doc_name": "文档标题",
        "content": "# Markdown 内容..."
    }
    """
    client_id = body.get("client_id")
    doc_name = body.get("doc_name", "未命名文档")
    content = body.get("content", "")

    if not client_id:
        return {"success": False, "error": "缺少 client_id"}

    # 使用 smartpage_create 创建智能文档（不是 create_doc，那是智能表格）
    pages = [{"page_title": "内容" if not doc_name else doc_name, "page_content": content, "content_type": 1}]
    create_resp = call_mcp("smartpage_create", {"title": doc_name, "pages": pages})
    if create_resp.get("error"):
        return {"success": False, "error": "创建文档失败：" + create_resp["error"]}
    docid = create_resp.get("docid", "")
    doc_url = create_resp.get("url", "")
    if not docid:
        return {"success": False, "error": "创建文档失败，未返回 docid"}

    return {"success": True, "docid": docid, "url": doc_url}


# ==================== Step5 Agent Demo H5 生成 ====================

AGENT_DEMO_SYSTEM_PROMPT = """你是一个专业的 H5 页面生成助手。根据客户需求分析报告，生成一个独立的、可直接在浏览器中运行的 HTML5 页面。

## 输出要求
1. 生成完整的 HTML 文件，包含所有 CSS/JS 内联代码
2. 页面必须是响应式的，支持手机和 PC
3. 页面内容要基于客户真实需求定制
4. 直接输出 HTML 代码，不要 markdown 代码块包裹

## 页面结构要求
- 顶部：客户名称 + Logo
- 主要内容区：根据需求分析展示关键信息（角色卡片、流程步骤、风险点等）
- 对话模拟区：模拟 AI 助手与用户的对话场景
- 底部：服务商信息

## 技术要求
- 使用纯 HTML + CSS + JavaScript（无外部依赖）
- 使用 CSS 变量管理颜色主题
- 页面加载后有基础的动画效果
- 支持滚动和基础交互"""


@app.post("/api/agent-demo/create")
async def create_agent_demo(body: dict, user: dict = Depends(require_auth)):
    """生成 Agent Demo H5 页面"""
    client_data = body.get("client_data", {})

    # 构建用户 prompt
    client_name = client_data.get("name", "未知客户")
    industry = client_data.get("industry", "")
    step4_report = client_data.get("step4_report", {})
    step1_result = client_data.get("step1_result", {})

    user_prompt = f"""## 客户信息
- 客户名称：{client_name}
- 行业：{industry}

## 需求分析报告摘要
{json.dumps(step4_report, ensure_ascii=False, indent=2) if step4_report else '暂无'}

## 客户画像摘要
{json.dumps(step1_result, ensure_ascii=False, indent=2) if step1_result else '暂无'}

请基于以上信息，生成一个展示 AI 售前助手能力的 H5 页面。"""

    result = call_minimax(AGENT_DEMO_SYSTEM_PROMPT, user_prompt, max_tokens=8000)
    record_ai_tokens(client_id, result["usage"]["total_tokens"])
    result = result["content"]

    # 保存到 public 目录
    import uuid, os
    from pathlib import Path

    # 确保 public 目录存在
    public_dir = Path(__file__).parent / "public"
    public_dir.mkdir(exist_ok=True)

    # 生成文件名
    filename = f"agent_demo_{client_name}_{uuid.uuid4().hex[:8]}.html"
    filepath = public_dir / filename

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(result)

    # 返回访问 URL
    url = f"/public/{filename}"

    return {"success": True, "url": url, "filename": filename}


PROFILE_GENERATE_PROMPT = """你是一个售前顾问。根据客户基本信息，生成一份结构化的客户画像摘要。

直接返回 JSON（不要 markdown 代码块），格式：
{
  "summary": "50字以内的客户画像一句话描述",
  "background": "公司背景描述，80字以内",
  "pain_points": ["核心痛点1", "核心痛点2", "核心痛点3"],
  "scale": "公司规模判断，如：200-1000人中大型企业",
  "tags": ["标签1", "标签2", "标签3"]
}"""

@app.post("/api/profile/generate")
async def generate_profile(body: dict, user: dict = Depends(require_auth)):
    """生成客户画像摘要"""
    company_name = body.get("company_name", "")
    industry = body.get("industry", "")
    initial_demand = body.get("initial_demand", "")
    if not company_name:
        return {"error": "缺少客户名称"}
    user_prompt = f"客户名称：{company_name}\n行业：{industry or '未指定'}\n原始需求：{initial_demand or '暂无'}\n请生成客户画像 JSON。"
    ai_result = call_minimax(PROFILE_GENERATE_PROMPT, user_prompt, max_tokens=1000)
    raw = ai_result["content"]
    if raw.startswith("Error:"):
        return {"error": raw}
    result = parse_json_response(raw)
    if not result:
        return {"error": "AI 返回格式异常，请重试"}
    return result


# ==================== 健康检查 ====================

COMPANY_SEARCH_PROMPT = """你是一个企业信息分析助手。根据客户名称、行业以及搜索结果，生成公司简介、主要客户群体、可能关注点。

参考搜索结果（来自互联网实时搜索）：
${search_results}

直接返回 JSON（不要 markdown 代码块），格式：
{
  "company_type": "公司类型描述，如：制造业龙头民营企业",
  "main_customers": "主要客户群体描述，如：大型三甲医院、政府机构",
  "possible_focus": "可能关注点，用/分隔，如：提升审批效率/降低运营成本/数据打通",
  "company_intro": "20字以内的公司简介"
}"""

@app.post("/api/company_search")
async def company_search(body: dict, user: dict = Depends(require_auth)):
    """AI 智搜：根据客户名称和行业生成公司简介（基于真实搜索结果）"""
    company_name = body.get("company_name", "")
    industry = body.get("industry", "")
    if not company_name:
        return {"error": "缺少公司名称"}

    # 第一步：用 Tavily 搜索获取真实信息
    # 搜索词加上"公司"和行业关键词，提高相关性
    search_terms = [company_name, "公司"]
    if industry and industry != "未指定":
        search_terms.append(industry)
    search_query = " ".join(search_terms)
    search_result = tavily_search(search_query, max_results=8)

    # 构建搜索结果摘要给 AI
    if search_result.get("success") and search_result.get("results"):
        results_text = "\n".join([
            f"- {r['title']}: {r['content'][:200]}..."
            for r in search_result["results"][:5]
        ])
    else:
        results_text = "（搜索失败，使用默认分析）"

    # 第二步：用 AI 分析生成 JSON
    user_prompt = f"客户名称：{company_name}\n行业：{industry or '未指定'}"
    raw = call_codebuddy(
        COMPANY_SEARCH_PROMPT.replace("${search_results}", results_text),
        user_prompt,
        max_tokens=3000
    )
    raw = raw["content"] if isinstance(raw, dict) else raw
    if raw.startswith("Error:"):
        return {"error": raw}
    result = parse_json_response(raw)
    if not result:
        return {"error": "AI 返回格式异常，请重试"}
    return result


@app.post("/api/skill/company_search")
async def skill_company_search(data: dict, request: Request):
    """
    Skill 企查查接口：搜索公司信息，结合新闻输出客户理解
    免登录，不需要 DeepSeek，只用 Tavily 搜索
    """
    company_name = data.get("company_name", "")
    if not company_name:
        raise HTTPException(status_code=400, detail="缺少公司名称")

    # 并行搜索：公司信息 + 新闻 + 行业动态
    queries = [
        (f"{company_name} 公司 简介 工商信息", 3),
        (f"{company_name} 企业 业务 经营", 2),
        (f"{company_name} 企业微信 数字化转型", 2),
    ]

    all_results = []
    for query, max_r in queries:
        result = tavily_search(query, max_results=max_r)
        if result.get("success") and result.get("results"):
            all_results.extend(result["results"])

    if all_results:
        # 去重并格式化为客户理解
        seen = set()
        lines = [f"关于【{company_name}】的搜索结果："]
        for item in all_results[:6]:
            title = item.get("title", "").strip()
            content = item.get("content", "")[:100].strip()
            key = title[:20]  # 简单去重
            if title and content and key not in seen:
                seen.add(key)
                lines.append(f"\n📌 {title}")
                lines.append(f"   {content}...")
    else:
        lines = [f"未找到 {company_name} 的公开信息。"]

    return {
        "success": True,
        "company_name": company_name,
        "description": "\n".join(lines)
    }


@app.get("/api/health")
async def health_check():
    """健康检查"""
    return {"status": "ok", "timestamp": datetime.now().isoformat()}

@app.get("/")
async def root():
    """根路径"""
    return {"message": "Provider Assist API", "version": "1.0.0"}

# ==================== Skill API（供 Work Buddy Skill 调用）====================

import uuid
import hashlib

# Skill API Key 认证
async def require_skill_auth(request: Request):
    """验证 Skill API Key"""
    api_key = request.headers.get("X-API-Key") or request.query_params.get("api_key")
    if not api_key:
        raise HTTPException(status_code=401, detail="缺少 API Key")

    conn = get_db()
    cursor = conn.cursor()

    # 解析 API Key：{受邀码}:{用户名}:{user_id}
    parts = api_key.split(":")
    if len(parts) == 3:
        invitation_code, username, user_id = parts
        try:
            user_id = int(user_id)
        except ValueError:
            raise HTTPException(status_code=401, detail="无效的 API Key")

        # 先尝试有受邀码的方式
        cursor.execute("""
            SELECT u.id, u.username, u.provider_name
            FROM users u
            JOIN invitation_codes ic ON u.provider_name = ic.provider_name
            WHERE u.id = ? AND u.username = ? AND ic.code = ?
        """, (user_id, username, invitation_code))
        row = cursor.fetchone()

        if not row:
            # 尝试没有受邀码的方式（测试用户）
            cursor.execute("""
                SELECT u.id, u.username, u.provider_name
                FROM users u
                WHERE u.id = ? AND u.username = ?
            """, (user_id, username))
            row = cursor.fetchone()
    else:
        # 旧格式或无效格式，只用 api_key 当用户名查
        cursor.execute("SELECT id, username, provider_name FROM users WHERE username = ?", (api_key,))
        row = cursor.fetchone()

    conn.close()

    if not row:
        raise HTTPException(status_code=401, detail="无效的 API Key")

    return {"user_id": row["id"], "username": row["username"], "provider_name": row["provider_name"]}


@app.post("/api/skill/login")
async def skill_login(data: dict):
    """
    Skill 登录接口
    API Key 格式：{受邀码}:{用户名}:{user_id}
    返回：access_token（与 /api/auth/login 一致）
    """
    api_key = data.get("api_key", "")
    password = data.get("password", "")

    if not api_key:
        raise HTTPException(status_code=400, detail="API Key 不能为空")
    if not password:
        raise HTTPException(status_code=400, detail="密码不能为空")

    # 解析 API Key：{受邀码}:{用户名}:{user_id} 或 {用户名}:{user_id}（测试用户）
    parts = api_key.split(":")
    invitation_code = ""
    if len(parts) == 3:
        invitation_code, username, user_id = parts
    elif len(parts) == 2:
        username, user_id = parts
    else:
        raise HTTPException(status_code=400, detail="API Key 格式错误")

    # 验证用户ID和用户名匹配
    try:
        user_id = int(user_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="API Key 格式错误")

    conn = get_db()
    cursor = conn.cursor()

    # 验证用户存在（先不用受邀码验证，支持没有受邀码的测试用户）
    cursor.execute("""
        SELECT u.id, u.username, u.password_hash, u.provider_name
        FROM users u
        WHERE u.id = ? AND u.username = ?
    """, (user_id, username))
    row = cursor.fetchone()

    if not row:
        raise HTTPException(status_code=401, detail="API Key 无效")

    # 如果 API Key 里有受邀码，额外验证受邀码匹配（兼容有受邀码的用户）
    if invitation_code:
        cursor.execute("""
            SELECT 1 FROM users u
            JOIN invitation_codes ic ON u.provider_name = ic.provider_name
            WHERE u.id = ? AND ic.code = ?
        """, (user_id, invitation_code))
        if not cursor.fetchone():
            raise HTTPException(status_code=401, detail="API Key 无效")

    # 验证密码
    if not verify_password(password, row["password_hash"]):
        raise HTTPException(status_code=401, detail="密码错误")

    # 生成 access_token
    token = create_access_token({"sub": row["username"], "user_id": row["id"]})

    return {
        "success": True,
        "access_token": token,
        "token_type": "bearer",
        "user": {
            "id": row["id"],
            "username": row["username"],
            "provider_name": row["provider_name"] or ""
        }
    }


@app.post("/api/skill/render")
async def skill_render_report(data: dict, request: Request):
    """
    Skill 报告渲染接口
    AI 生成 JSON 数据 → 后端用模板渲染 HTML
    """
    try:
        auth = await require_skill_auth(request)
    except HTTPException:
        raise HTTPException(status_code=401, detail="无效的 API Key")

    report_type = data.get("type", "sales")  # sales / tech / quote
    report_data = data.get("data", {})

    # 加载模板
    from jinja2 import Environment, FileSystemLoader
    templates_dir = Path(__file__).parent / "templates" / "reports"
    env = Environment(loader=FileSystemLoader(str(templates_dir)))
    template_file = f"{report_type}.html"

    try:
        template = env.get_template(template_file)
    except Exception:
        raise HTTPException(status_code=400, detail=f"不支持的报告类型: {report_type}")

    # 渲染
    from datetime import datetime
    report_data["generate_date"] = datetime.now().strftime("%Y-%m-%d")

    html_content = template.render(**report_data)

    return {
        "success": True,
        "html": html_content,
        "type": report_type
    }


@app.post("/api/skill/submit")
async def skill_submit(data: dict, request: Request):
    """WB Skill 提交完整售前数据（必须通过 X-API-Key 或 Bearer Token 鉴权）"""
    # 1) 先尝试 X-API-Key / ?api_key= 鉴权
    auth = None
    skill_auth_err = None
    try:
        auth = await require_skill_auth(request)
    except HTTPException as e:
        skill_auth_err = e

    # 2) 若 X-API-Key 缺失/失败，再尝试 Bearer Token（兼容旧客户端）
    if auth is None:
        bearer = request.headers.get("Authorization", "")
        if bearer.startswith("Bearer "):
            try:
                from auth import decode_token
                payload = decode_token(bearer[7:])
                if payload and "sub" in payload and "user_id" in payload:
                    conn_tmp = get_db()
                    cur_tmp = conn_tmp.cursor()
                    cur_tmp.execute(
                        "SELECT id, username, provider_name FROM users WHERE id = ? AND username = ?",
                        (payload["user_id"], payload["sub"]),
                    )
                    row = cur_tmp.fetchone()
                    conn_tmp.close()
                    if row:
                        auth = {
                            "user_id": row["id"],
                            "username": row["username"],
                            "provider_name": row["provider_name"] or "",
                        }
            except Exception:
                auth = None

    if auth is None:
        # 优先把 Skill Auth 的原始 401 信息抛给调用方
        if skill_auth_err is not None:
            raise skill_auth_err
        raise HTTPException(status_code=401, detail="缺少有效的 API Key 或 Bearer Token")

    conn = get_db()
    cursor = conn.cursor()

    # 解析数据
    client_name = data.get("client_name", "")
    industry = data.get("industry", "")
    scale = data.get("scale", "")
    tags = data.get("tags", [])
    initial_demand = data.get("initial_demand", "")
    profile_json = data.get("profile_json", {})
    profile_text = data.get("profile_text", "")
    visit_outline = data.get("visit_outline", "")
    meeting_notes = data.get("meeting_notes", [])
    md_outline = data.get("md_outline", "")
    reports = data.get("reports", [])
    token_estimate = data.get("token_estimate", 0)
    cost_estimate = data.get("cost_estimate", 0)

    # 创建或更新客户（SaaS 兼容的数据模型）
    user_id = auth.get("user_id")

    # 构建 step4_input_draft_json（SaaS API 需要的数据结构）
    step4_draft = {
        "confirmedNeeds": [{"title": initial_demand or "待确认", "description": ""}],
        "painPoints": profile_json.get("part1", {}).get("pain_points", []) if profile_json else [],
        "involvedRoles": [],
        "phaseOneScope": [],
        "phaseTwoScope": [],
        "pendingQuestions": profile_json.get("part2", {}).get("gaps", []) if profile_json else []
    }

    if user_id:
        cursor.execute("SELECT id FROM clients WHERE user_id = ? AND name = ?", (user_id, client_name))
        existing = cursor.fetchone()
        if existing:
            client_id = existing["id"]
            cursor.execute("""
                UPDATE clients SET
                    industry = ?, scale = ?, tags = ?, initial_demand = ?,
                    step1_result = ?, step2_report = ?, step3_summary_full = ?,
                    step4_input_draft_json = ?
                WHERE id = ?
            """, (industry, scale, json.dumps(tags), initial_demand,
                  json.dumps(profile_json), profile_text, md_outline,
                  json.dumps(step4_draft, ensure_ascii=False), client_id))
        else:
            cursor.execute(
                "INSERT INTO clients (user_id, name, industry, scale, tags, initial_demand, step1_result, step2_report, step3_summary_full, step4_input_draft_json) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (user_id, client_name, industry, scale, json.dumps(tags), initial_demand,
                 json.dumps(profile_json), profile_text, md_outline,
                 json.dumps(step4_draft, ensure_ascii=False))
            )
            client_id = cursor.lastrowid
    else:
        # 无用户ID，使用默认值0（游客记录）
        cursor.execute(
            "INSERT INTO clients (user_id, name, industry, scale, tags, initial_demand, step1_result, step2_report, step3_summary_full, step4_input_draft_json) VALUES (0, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (client_name, industry, scale, json.dumps(tags), initial_demand,
             json.dumps(profile_json), profile_text, md_outline,
             json.dumps(step4_draft, ensure_ascii=False))
        )
        client_id = cursor.lastrowid

    # 存储报告信息
    reports_json = json.dumps(reports, ensure_ascii=False)
    cursor.execute("UPDATE clients SET step4_presales_versions = ? WHERE id = ?", (reports_json, client_id))

    # 更新 Token 消耗
    cursor.execute("UPDATE clients SET token_count = COALESCE(token_count, 0) + ? WHERE id = ?", (token_estimate, client_id))

    conn.commit()
    conn.close()

    return {"success": True, "client_id": client_id, "message": "数据已保存"}


@app.post("/api/skill/reports")
async def skill_upload_report(request: Request, data: dict = None):
    """接收 WB Skill 生成的 HTML 报告，写入 sining.cloud/reports/"""
    try:
        auth = await require_skill_auth(request)
    except HTTPException:
        raise HTTPException(status_code=401, detail="无效的 API Key")

    if data is None:
        body = await request.json()
    else:
        body = data

    html_content = body.get("html", "")
    report_type = body.get("type", "sales")  # sales / tech / quote
    client_name = body.get("client_name", "unknown")

    if not html_content:
        raise HTTPException(status_code=400, detail="HTML 内容不能为空")

    # 生成唯一 ID
    report_id = f"rpt_{uuid.uuid4().hex[:12]}"
    filename = f"{report_id}_{report_type}.html"

    # 报告存储目录
    reports_dir = Path(__file__).parent.parent / "data" / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)

    # 写入文件
    filepath = reports_dir / filename
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html_content)

    # 构建访问 URL
    base_url = os.environ.get("BASE_URL", "https://sining.cloud")
    report_url = f"{base_url}/reports/{filename}"

    return {
        "success": True,
        "id": report_id,
        "url": report_url,
        "filename": filename
    }


@app.get("/api/skill/reports/{report_id}")
async def skill_get_report(report_id: str, request: Request):
    """获取报告信息"""
    try:
        auth = await require_skill_auth(request)
    except HTTPException:
        raise HTTPException(status_code=401, detail="无效的 API Key")

    reports_dir = Path(__file__).parent.parent / "data" / "reports"
    # 查找匹配的报告文件
    for f in reports_dir.glob(f"{report_id}_*.html"):
        return {
            "id": report_id,
            "filename": f.name,
            "url": f"/reports/{f.name}"
        }

    raise HTTPException(status_code=404, detail="报告不存在")


@app.post("/api/skill/knowledge")
async def skill_save_knowledge(request: Request, data: dict = None):
    """保存到服务商知识库"""
    try:
        auth = await require_skill_auth(request)
    except HTTPException:
        raise HTTPException(status_code=401, detail="无效的 API Key")

    if data is None:
        body = await request.json()
    else:
        body = data

    kb_type = body.get("type", "case")  # case / template / fragment
    title = body.get("title", "")
    content = body.get("content", "")
    industry = body.get("industry", "")
    tags = body.get("tags", [])

    conn = get_db()
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO provider_knowledge (user_id, category, title, content, industry, tags) VALUES (?, ?, ?, ?, ?, ?)",
        (auth["user_id"], kb_type, title, content, industry, json.dumps(tags))
    )
    conn.commit()
    kb_id = cursor.lastrowid
    conn.close()

    return {"success": True, "id": kb_id, "message": "已保存到知识库"}


@app.get("/api/skill/knowledge")
async def skill_list_knowledge(
    request: Request,
    industry: str = "",
    kb_type: str = "",
    user: dict = Depends(require_auth)
):
    """查询服务商知识库"""
    conn = get_db()
    cursor = conn.cursor()

    query = "SELECT * FROM provider_knowledge WHERE user_id = ?"
    params = [user["user_id"]]

    if industry:
        query += " AND industry = ?"
        params.append(industry)
    if kb_type:
        query += " AND category = ?"
        params.append(kb_type)

    query += " ORDER BY created_at DESC"

    cursor.execute(query, params)
    items = [dict(row) for row in cursor.fetchall()]
    conn.close()

    return items


@app.get("/api/skill/prompts/{step_name}")
async def skill_get_prompt(step_name: str, request: Request):
    """获取指定 Step 的 Prompt 内容"""
    try:
        auth = await require_skill_auth(request)
    except HTTPException:
        raise HTTPException(status_code=401, detail="无效的 API Key")

    # Prompt 文件路径
    skill_dir = Path(__file__).parent.parent / "skill" / "prompts"
    prompt_map = {
        "step1_prep": "step1_prep.md",
        "step3_notes": "step3_notes.md",
        "step4_sales": "step4_sales.md",
        "step4_tech": "step4_tech.md",
        "step4_quote": "step4_quote.md",
        "step5_kb": "step5_kb.md",
    }

    filename = prompt_map.get(step_name)
    if not filename:
        raise HTTPException(status_code=404, detail=f"未找到 Prompt: {step_name}")

    prompt_file = skill_dir / filename
    if not prompt_file.exists():
        raise HTTPException(status_code=404, detail=f"Prompt 文件不存在: {step_name}")

    content = prompt_file.read_text(encoding="utf-8")

    return {
        "step": step_name,
        "content": content
    }


@app.get("/api/skill/clients")
async def skill_list_clients(request: Request, user: dict = Depends(require_auth)):
    """获取服务商的所有客户（Skill 用）"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT id, name, industry, scale, tags, initial_demand,
               step1_result, step2_report, step3_summary, step4_presales_versions,
               created_at, updated_at
        FROM clients WHERE user_id = ? ORDER BY updated_at DESC
    """, (user["user_id"],))

    clients = []
    for row in cursor.fetchall():
        d = dict(row)
        # 解析 JSON 字段
        for field in ["step1_result", "step4_presales_versions"]:
            if d.get(field) and isinstance(d[field], str):
                try:
                    d[field] = json.loads(d[field])
                except:
                    pass
        clients.append(d)

    conn.close()
    return clients


@app.get("/api/skill/clients/{client_id}")
async def skill_get_client(client_id: int, request: Request, user: dict = Depends(require_auth)):
    """获取单个客户的完整数据（Skill 用）"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT * FROM clients WHERE id = ? AND user_id = ?
    """, (client_id, user["user_id"]))
    row = cursor.fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="客户不存在")

    result = dict(row)
    # 解析 JSON 字段
    for field in ["step1_result", "step2_report", "step3_summary", "step4_presales_versions", "step5_schema"]:
        if result.get(field) and isinstance(result[field], str):
            try:
                result[field] = json.loads(result[field])
            except:
                pass

    return result


# ==================== 访问日志 ====================

@app.post("/api/visits")
async def log_visit(request: Request, data: dict):
    """记录报告访问日志"""
    client_id = data.get("client_id")
    report_type = data.get("report_type", "")
    ip = request.client.host if request.client else ""

    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO visit_tracking (client_id, file_url, ip_address, last_visit_at)
        VALUES (?, ?, ?, datetime('now'))
    """, (client_id, report_type, ip))
    conn.commit()
    conn.close()

    return {"success": True}


@app.get("/api/visits")
async def list_visits(request: Request, user: dict = Depends(require_auth)):
    """获取当前服务商的访问记录"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT vt.*, c.name as client_name
        FROM visit_tracking vt
        LEFT JOIN clients c ON vt.client_id = c.id
        WHERE c.user_id = ?
        ORDER BY vt.last_visit_at DESC
        LIMIT 200
    """, (user["user_id"],))
    rows = [dict(r) for r in cursor.fetchall()]
    conn.close()
    return {"visits": rows}


@app.get("/api/provider-knowledge/{kb_id}")
async def get_knowledge_item(kb_id: int, user: dict = Depends(require_auth)):
    """获取知识库单条内容"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT * FROM provider_knowledge WHERE id = ? AND user_id = ?",
        (kb_id, user["user_id"])
    )
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="不存在")
    return dict(row)


# ==================== Skill 安装接口 ====================

@app.get("/api/skill/skill.md")
async def get_skill_manifest(request: Request):
    """返回 Work Buddy 可安装的 skill.md manifest"""
    skill_md = Path(__file__).parent.parent / "skill" / "skill.md"
    if not skill_md.exists():
        raise HTTPException(status_code=404, detail="Skill not found")
    content = skill_md.read_text(encoding="utf-8")
    from starlette.responses import Response
    return Response(content=content, media_type="text/markdown; charset=utf-8")


@app.get("/api/skill/manifest")
async def get_skill_manifest_json(request: Request):
    """返回 Skill 元信息"""
    base_url = os.environ.get("BASE_URL", "https://sining.cloud")
    return {
        "name": "provider-assist",
        "version": "2.0.0",
        "description": "服务商售前助手 - 客户调研、沟通纪要、方案生成",
        "trigger": "/xiaoqiu",
        "install_url": f"{base_url}/api/skill/skill.md",
        "author": "Provider Assist Team",
        "commands": [
            {"name": "/clean", "description": "清理当前客户会话，开始新客户"},
            {"name": "/memory", "description": "查看当前客户已收集的信息"},
            {"name": "/my", "description": "查看我的客户列表、访问记录、知识库"},
            {"name": "/help", "description": "显示帮助"}
        ]
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
